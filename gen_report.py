from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = 'SimSun'
font.size = Pt(11)

# 标题样式
for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = 'SimHei'
    heading_style.font.color.rgb = RGBColor(0, 0, 0)

# 标题
title = doc.add_heading('', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('关于中国发展阶段认定的立场报告')
run.font.size = Pt(22)
run.font.name = 'SimHei'

doc.add_paragraph('')

# 一、问题的提出
doc.add_heading('一、问题的提出', level=1)
doc.add_paragraph(
    '近年来，美国持续推动在国际组织中将中国从"发展中国家"重新归类为"发达国家"。'
    '这一提议涉及国际经济秩序、贸易规则、气候治理等多个领域的深层博弈，'
    '对中国的国际地位和发展权益具有重大影响。本报告旨在梳理各方标准与立场，分析中国应如何被归类。'
)

# 二、国际组织定义与标准
doc.add_heading('二、国际组织对"发达国家"的定义与标准', level=1)

doc.add_heading('1. 联合国', level=2)
doc.add_paragraph(
    '联合国并未对"发达国家"给出严格的量化定义。联合国开发计划署（UNDP）主要通过'
    '"人类发展指数（HDI）"来衡量各国发展水平，将HDI≥0.800归类为"极高人类发展水平"。'
    '2024年数据显示，中国HDI约为0.788，处于"高人类发展水平"区间（0.700-0.799），'
    '尚未达到"极高"门槛。'
)

doc.add_heading('2. 世界银行', level=2)
doc.add_paragraph(
    '世界银行按人均国民总收入（GNI）将经济体分为四组：低收入、中低收入、中高收入和高收入。'
    '2024年标准中，高收入经济体门槛为人均GNI≥14,005美元。'
    '中国2024年人均GNI约为13,400美元，尚未达到高收入国家标准。'
)

doc.add_heading('3. 国际货币基金组织（IMF）', level=2)
doc.add_paragraph(
    'IMF将成员国分为"先进经济体"和"新兴市场与发展中经济体"两大类。'
    '中国的分类为"新兴市场与发展中经济体"。IMF的分类更多基于经济结构特征'
    '（金融开放度、贸易依存度等），而非单一经济指标。'
)

doc.add_heading('4. 世界贸易组织（WTO）', level=2)
doc.add_paragraph(
    'WTO没有正式的"发展中国家"定义清单，而是实行"自我认定"原则，由各成员国自行声明。'
    'WTO中的特殊与差别待遇（SDT）条款允许发展中成员在关税减让、过渡期等方面享受优惠。'
    '美国多次批评中国利用发展中国家身份获取不正当利益。'
)

# 三、中国当前关键发展指标
doc.add_heading('三、中国当前关键发展指标', level=1)

table = doc.add_table(rows=6, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['指标', '中国数据', '发达国家参考标准']
data = [
    ['人均GDP', '约12,800美元（2024）', '通常>25,000美元'],
    ['人均GNI', '约13,400美元（2024）', '>14,005美元（世行高收入门槛）'],
    ['HDI人类发展指数', '0.788（2024）', '≥0.800（极高人类发展水平）'],
    ['城镇化率', '约67%（2024）', '通常>80%'],
    ['基尼系数', '约0.47（2024）', '多数发达国家<0.40'],
]

for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

for row_idx, row_data in enumerate(data, 1):
    for col_idx, cell_text in enumerate(row_data):
        table.rows[row_idx].cells[col_idx].text = cell_text

doc.add_paragraph('')
doc.add_paragraph(
    '从上述数据可见，中国在多项核心指标上仍与发达国家存在显著差距。'
)

# 四、美国的行动与理由
doc.add_heading('四、美国的行动与理由', level=1)

doc.add_heading('主要行动', level=2)
actions = [
    '2020年：特朗普政府正式向WTO提交提案，反对中国继续使用发展中国家身份。该提案未获通过。',
    '2022-2023年：拜登政府继续在多个场合强调中国应承担"与其实力相匹配"的国际责任。',
    '2023年以来：美国在WTO改革谈判、G20峰会、APEC等场合反复提出此议题。',
    '气候领域：美国在联合国气候变化框架公约（UNFCCC）下质疑中国作为发展中国家获取气候资金的资格。',
]
for action in actions:
    doc.add_paragraph(action, style='List Bullet')

doc.add_heading('美国的核心论据', level=2)
reasons = [
    '经济体量：中国GDP总量超过18万亿美元，居全球第二，是全球最大的货物贸易国。',
    '军事实力：中国拥有世界第二大军费开支。',
    '科技水平：在5G、人工智能、航天、量子计算等领域具有全球竞争力。',
    '对外投资与贷款：中国通过"一带一路"等倡议进行大规模海外投资，已成为全球最大双边债权国之一。',
    '碳排放：中国是全球最大的碳排放国，年排放量超过部分发达国家总和。',
]
for reason in reasons:
    doc.add_paragraph(reason, style='List Bullet')

# 五、国际反应与各方立场
doc.add_heading('五、国际反应与各方立场', level=1)

doc.add_heading('支持重新归类的力量', level=2)
support = [
    '美国：最积极的推动者，核心关切是削弱中国在多边体系中的竞争优势。',
    '日本：支持逐步调整中国的发展中国家地位，关切中国在亚太地区的经济竞争。',
    '欧盟部分国家：在贸易谈判中倾向于不将中国视为发展中国家，但在气候等议题上态度较为灵活。',
    '澳大利亚：在贸易争端背景下支持重新审视中国的分类。',
]
for s in support:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('反对重新归类的力量', level=2)
oppose = [
    '中国：坚决反对，强调人均指标仍处于发展中国家水平。',
    'G77+中国集团：广大发展中国家普遍支持中国的立场，担心多米诺效应导致自身地位也受到挑战。',
    '基础四国（巴西、印度、南非）：与中国立场一致，反对以总量指标取代人均指标来认定发展状态。',
    '联合国贸发会议（UNCTAD）：多数成员国反对单方面重新分类。',
]
for o in oppose:
    doc.add_paragraph(o, style='List Bullet')

doc.add_heading('各方利益关切', level=2)
doc.add_paragraph(
    '支持方核心关切在于：重新归类将迫使中国在贸易减让、市场准入、气候资金等方面承担更多义务，'
    '有利于缩小中美之间的竞争差距。'
)
doc.add_paragraph(
    '反对方核心关切在于：如果中国被强制归类，发展中国家在国际体系中的特殊待遇将被系统性削弱，'
    '全球南方国家的利益将受到损害。'
)

# 六、潜在影响分析
doc.add_heading('六、潜在影响分析', level=1)
doc.add_paragraph('若中国被强制认定为"发达国家"，可能失去以下待遇：')

impacts = [
    ('WTO特殊与差别待遇（SDT）',
     '中国将在关税减让谈判中失去过渡期和灵活安排，需在更多领域实行零关税或大幅降低关税。'
     '在贸易争端解决中，中国将不再享有更长的合规期限。'),
    ('气候资金援助',
     '中国将失去获得绿色气候基金（GCF）、全球环境基金（GEF）等气候融资的资格。'
     '发达国家被要求成为气候资金的出资方而非受益方。'),
    ('世界银行贷款优惠',
     '中国将无法以优惠利率从世界银行国际开发协会（IDA）获取贷款，'
     '只能从国际复兴开发银行（IBRD）获得商业化条件的贷款。'),
    ('联合国发展援助',
     '中国可能被排除在联合国开发计划署（UNDP）等机构的优先援助对象之外。'),
    ('其他领域',
     '包括但不限于：国际农业发展基金（IFAD）优惠贷款、联合国人口基金项目、'
     '部分国际科研合作项目的优先资格等。'),
]

for title_text, content in impacts:
    doc.add_heading(title_text, level=2)
    doc.add_paragraph(content)

# 七、中国的立场与回应
doc.add_heading('七、中国的立场与回应', level=1)

doc.add_heading('官方立场', level=2)
positions = [
    '人均指标优先：中国坚持应以人均发展水平而非总量来衡量发展阶段。中国人均GDP仅为美国的六分之一、日本的四分之一。',
    '最大发展中国家的定位：习近平主席多次在国际场合指出，"中国仍然是世界上最大的发展中国家"，这一基本国情没有改变。',
    '反对政治化操作：中国认为美国的提议本质上是将分类标准政治化，旨在遏制中国发展。',
    '主动承担大国责任：中国强调，虽然坚持发展中国家定位，但愿在力所能及范围内为全球发展事业作出更大贡献。',
]
for p in positions:
    doc.add_paragraph(p, style='List Bullet')

doc.add_heading('关键表态', level=2)
expressions = [
    '习近平在金砖国家峰会、G20峰会等场合多次重申中国的发展中国家地位。',
    '外交部发言人多次回应称，"中国的定位不是由某个国家说了算"，应由国际社会共同决定。',
    '在气候变化领域，中国承诺在2030年前碳达峰、2060年前碳中和，体现大国担当。',
]
for e in expressions:
    doc.add_paragraph(e, style='List Bullet')

# 八、结论与立场
doc.add_heading('八、结论与立场', level=1)

doc.add_heading('核心判断：强行认定中国为"发达国家"不合适', level=2)

conclusions = [
    '第一，从客观标准看，中国的发展中国家地位有充分的数据支撑。人均GDP、HDI、城镇化率、基尼系数等核心指标均表明，中国仍处于从中等收入向高收入过渡的阶段。世界银行本身尚未将中国列为高收入经济体，联合国的HDI分类亦然。',
    '第二，从程序正义看，单方面强制重新归类缺乏合法性。国际组织的分类标准应由多边机制决定，而非由某一国家推动。WTO的自我认定原则是各成员国长期协商的结果。',
    '第三，从实际影响看，此举将严重损害全球南方国家的集体利益。如果中国的定位可以被单方面改变，其他发展中国家的特殊待遇同样面临风险，国际秩序的公平性将进一步被削弱。',
    '第四，从道义角度看，14亿人口的发展权不容剥夺。中国仍有多地区发展不平衡、大量人口收入偏低等现实问题，不应因经济总量的增长而忽视人均水平和发展质量的差距。',
]
for c in conclusions:
    doc.add_paragraph(c, style='List Bullet')

doc.add_heading('建议', level=2)
suggestions = [
    '坚持以人均指标为核心的发展阶段评估体系。',
    '在国际组织中团结广大发展中国家，维护多边主义和公平分类原则。',
    '主动参与国际分类标准的讨论与制定，推动建立更加科学、公正的评估框架。',
    '在坚持发展中国家定位的同时，主动承担与自身能力相适应的国际责任，展现负责任大国形象。',
]
for i, s in enumerate(suggestions, 1):
    doc.add_paragraph(f'{i}. {s}')

doc.add_paragraph('')
doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('报告生成日期：2025年6月')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(128, 128, 128)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('说明：本报告基于公开资料整理分析，所引用数据截至2024年底或2025年初。')
run2.font.size = Pt(9)
run2.font.color.rgb = RGBColor(128, 128, 128)

doc.save(r'C:\Users\fy\Desktop\关于中国发展阶段认定的立场报告.docx')
print('文档已保存到桌面')
