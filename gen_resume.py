from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# 页面边距
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ========== 辅助函数 ==========
def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_section_title(doc, text):
    """添加带蓝色竖线的章节标题"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    # 添加蓝色竖线标记（用蓝色文字模拟）
    run = p.add_run("▎")
    run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
    run.font.size = Pt(14)
    run.bold = True
    # 添加标题文字
    run2 = p.add_run(text)
    run2.font.size = Pt(14)
    run2.bold = True
    run2.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
    run2.font.name = "微软雅黑"
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_normal_text(doc, text, bold=False, indent=False):
    """添加正文段落"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(20)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if bold:
        run.bold = True
    return p

def add_experience_item(doc, title, subtitle, bullets):
    """添加经历条目"""
    # 标题行
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.font.size = Pt(11)
    run.bold = True
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run2 = p.add_run(f"    {subtitle}")
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run2.font.name = "宋体"
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # 要点
    for bullet in bullets:
        bp = doc.add_paragraph()
        bp.paragraph_format.space_before = Pt(1)
        bp.paragraph_format.space_after = Pt(1)
        bp.paragraph_format.left_indent = Cm(0.5)
        run = bp.add_run(f"• {bullet}")
        run.font.size = Pt(10.5)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ========== 标题 ==========
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(4)
run = title.add_run("个 人 简 历")
run.font.size = Pt(22)
run.bold = True
run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)
run.font.name = "微软雅黑"
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 分隔线
line = doc.add_paragraph()
line.alignment = WD_ALIGN_PARAGRAPH.CENTER
line.paragraph_format.space_before = Pt(0)
line.paragraph_format.space_after = Pt(6)
run = line.add_run("━" * 50)
run.font.size = Pt(6)
run.font.color.rgb = RGBColor(0x2B, 0x57, 0x9A)

# ========== 一、基本信息 ==========
add_section_title(doc, "基本信息")

info_table = doc.add_table(rows=4, cols=4)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_table.style = 'Table Grid'

info_data = [
    ["姓    名", "", "性    别", ""],
    ["学    历", "本科", "出生年月", ""],
    ["毕业院校", "郑州西亚斯学院", "所学专业", "人工智能"],
    ["联系电话", "", "电子邮箱", ""],
]

for i, row_data in enumerate(info_data):
    row = info_table.rows[i]
    for j, cell_text in enumerate(row_data):
        cell = row.cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(cell_text)
        run.font.size = Pt(10.5)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # 标签列加粗 + 浅蓝背景
        if j % 2 == 0:
            run.bold = True
            set_cell_shading(cell, "E8F0FE")

# 设置列宽
for row in info_table.rows:
    row.cells[0].width = Cm(2.5)
    row.cells[1].width = Cm(5)
    row.cells[2].width = Cm(2.5)
    row.cells[3].width = Cm(5)

# 求职意向单独一行
p意向 = doc.add_paragraph()
p意向.paragraph_format.space_before = Pt(6)
p意向.paragraph_format.space_after = Pt(2)
run = p意向.add_run("求职意向：")
run.bold = True
run.font.size = Pt(11)
run.font.name = "宋体"
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run2 = p意向.add_run("数据分析师 / AI算法工程师")
run2.font.size = Pt(11)
run2.font.name = "宋体"
run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ========== 二、教育背景 ==========
add_section_title(doc, "教育背景")

add_experience_item(doc,
    "郑州西亚斯学院",
    "2023.09 — 2026.06（预计）",
    [
        "专业：人工智能（本科）",
        "主修课程：深度学习、机器学习、Python程序设计、操作系统、计算机组成原理",
    ]
)

# ========== 三、项目经历 ==========
add_section_title(doc, "项目经历")

add_experience_item(doc,
    "AI视频识别信号平台（独立/团队项目）",
    "2025.03 — 至今",
    [
        "项目简介：面向园区围栏、仓库、码头等安防场景的AI视频识别与告警闭环平台",
        "技术栈：Python / FastAPI / YOLOv8 / OpenCV / SQLite / Docker / HTML+JS",
        "负责AI视觉检测模块：基于YOLOv8实现人员目标实时检测，集成IoU贪心跟踪器实现目标轨迹追踪",
        "参与前后端开发：使用FastAPI构建RESTful API，HTML/CSS/JavaScript实现监控大屏与告警管理页面",
        "参与规则引擎开发：实现围栏翻越检测、区域滞留检测等安防规则，支持配置驱动的场景化部署",
        "完成Docker容器化部署与Windows守护脚本，系统经44项后端单元测试与16项黑盒测试均通过",
    ]
)

add_experience_item(doc,
    "全国大学生统计建模大赛（河南赛区）",
    "2025.04 — 2025.08",
    [
        "参赛作品：《中国老龄化背景下农村家庭结构变迁与医疗支出趋势分析》",
        "负责数据采集与清洗，运用统计建模方法对农村老龄化医疗支出进行定量分析",
        "基于Python完成数据可视化与回归模型构建，产出完整分析报告",
        "荣获本科生组省级二等奖（河南赛区）",
    ]
)

# ========== 四、荣誉与技能 ==========
add_section_title(doc, "荣誉与技能")

# 荣誉
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
run = p.add_run("荣誉奖项")
run.bold = True
run.font.size = Pt(11)
run.font.name = "宋体"
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

honors = [
    "2025年  全国大学生统计建模大赛 河南赛区 本科生组二等奖",
    "2024年  全国大学英语四级证书（CET-4）",
    "2024年  普通话水平测试等级证书",
]
for h in honors:
    bp = doc.add_paragraph()
    bp.paragraph_format.space_before = Pt(1)
    bp.paragraph_format.space_after = Pt(1)
    bp.paragraph_format.left_indent = Cm(0.5)
    run = bp.add_run(f"• {h}")
    run.font.size = Pt(10.5)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 专业技能
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run = p.add_run("专业技能")
run.bold = True
run.font.size = Pt(11)
run.font.name = "宋体"
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

skills = [
    "编程语言：Python（熟练）、SQL（基础）",
    "AI视觉：熟悉YOLO系列目标检测模型，了解IoU跟踪算法与OpenCV视频处理",
    "后端开发：掌握FastAPI框架，具备RESTful API设计与SQLite数据库开发经验",
    "数据分析：掌握NumPy、Pandas、Matplotlib等数据分析与可视化工具",
    "运维部署：了解Docker容器化部署、Nginx反向代理与基本Linux操作",
    "办公技能：熟练使用WPS/Office办公套件",
]
for s in skills:
    bp = doc.add_paragraph()
    bp.paragraph_format.space_before = Pt(1)
    bp.paragraph_format.space_after = Pt(1)
    bp.paragraph_format.left_indent = Cm(0.5)
    run = bp.add_run(f"• {s}")
    run.font.size = Pt(10.5)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ========== 五、自我评价 ==========
add_section_title(doc, "自我评价")

evaluations = [
    "具备AI视觉检测与全栈开发能力，独立完成基于YOLOv8的视频识别平台开发，涵盖目标检测、规则引擎、前后端页面与容器化部署",
    "学习能力强，能够快速掌握新技术和新工具，在项目中自学YOLO模型部署、FastAPI框架与Docker容器化等技术",
    "具有良好的团队协作精神，在竞赛与项目中与团队成员紧密配合，高效完成分工任务",
]
for e in evaluations:
    bp = doc.add_paragraph()
    bp.paragraph_format.space_before = Pt(2)
    bp.paragraph_format.space_after = Pt(2)
    bp.paragraph_format.left_indent = Cm(0.5)
    run = bp.add_run(f"• {e}")
    run.font.size = Pt(10.5)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ========== 保存 ==========
doc.save(r"D:\Project\个人简历.docx")
print("简历已生成：D:\\Project\\个人简历.docx")
