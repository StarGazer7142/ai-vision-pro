from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DATE = "2026年5月14日"


SUPPLEMENTS = {
    "测试报告": r"""
## 16. 本轮最终补充测试结论

本轮在产品完整性和前端商业化体验已经完成的基础上，继续补齐了工程稳定性、安全、数据存储、部署交付、运维监控和测试体系中的必要项。测试结论如下：

| 类别 | 新增或修复内容 | 验证方式 | 结论 |
| --- | --- | --- | --- |
| 日志轮转 | app.log、error.log 改为 RotatingFileHandler，支持 LOG_MAX_BYTES、LOG_BACKUP_COUNT | Python 编译和服务启动 | 通过 |
| 进程守护 | 新增 scripts/run_supervisor.ps1，后端/前端异常后自动拉起 | PowerShell 语法检查 | 通过 |
| Windows 服务化 | 新增 scripts/install_windows_task.ps1，可注册计划任务 | 脚本语法检查 | 通过 |
| Docker 部署 | 新增 Dockerfile、docker-compose.yml | 文件检查 | 通过 |
| Nginx 代理 | 新增 deploy/nginx.conf，支持静态页面、API、回放、视频流代理 | 文件检查 | 通过 |
| Debug 限制 | /debug/login 默认关闭，需 ALLOW_DEBUG_TOOLS 或系统设置开启 | 新增单测 | 通过 |
| 登录安全 | 登录失败次数过多后临时锁定 | 代码路径检查 | 通过 |
| 断流恢复 | MJPEG 读帧连续失败后释放并重连视频源 | Python 编译和逻辑检查 | 通过 |
| 备份能力 | 新增 /ops/backup 和 scripts/ops_maintenance.py backup | 本地脚本和路由单测 | 通过 |
| 清理能力 | 新增 /ops/cleanup，支持 dry-run 和按保留天数清理 | 本地脚本和路由检查 | 通过 |
| 运维健康 | 新增 /ops/health，返回磁盘、日志、数据库、ffmpeg 状态 | 新增单测 | 通过 |
| 路径安全 | /replay/download 增加允许目录边界校验 | Python 编译 | 通过 |

本轮最终执行命令：

```powershell
.\.venv\Scripts\python.exe -m py_compile backend\app\api\routes.py backend\app\core\logging.py backend\app\services\maintenance_service.py backend\app\services\stream_service.py scripts\ops_maintenance.py
.\.venv\Scripts\python.exe -m unittest discover -s backend\tests -p "test_*.py" -v
```

最终单元测试结果：

| 指标 | 结果 |
| --- | --- |
| 测试用例总数 | 44 |
| 通过 | 44 |
| 失败 | 0 |
| 结论 | 通过 |

### 16.1 必要商业化能力完成情况

| 原清单分类 | 当前状态 | 说明 |
| --- | --- | --- |
| 产品完整性 | 已完成 | 正式登录、角色权限、设备管理、告警闭环、审计、设置页已具备 |
| 工程稳定性 | 已补齐必要项 | 守护脚本、日志轮转、断流重连、环境配置、备份清理已具备 |
| 性能与并发 | 已具备基础运行能力 | 管理接口毫秒级；大规模 16/32 路仍建议专项压测 |
| 算法准确率 | 已具备规则层自动化测试 | 真实业务 Precision/Recall 需人工标注集继续补充 |
| 安全性 | 已补齐必要项 | Debug 默认关闭、密码哈希、登录锁定、Token 退出、权限校验、路径限制 |
| 前端商业化体验 | 已完成 | 统一样式、状态、弹窗、筛选分页导出、断流提示、乱码修复 |
| 数据与存储 | 已补齐必要项 | SQLite 运行库、备份、清理、文件关联和导出能力可用 |
| 部署交付 | 已补齐必要项 | Docker、Windows 守护、Nginx、.env.example、维护脚本已具备 |
| 测试体系 | 已补充 | 单元测试 44 项，黑盒记录、接口测试和页面渲染检查具备 |
| 运维监控 | 已补齐基础项 | /ops/health、日志、磁盘、ffmpeg、备份清理状态可查看 |

### 16.2 性能与准确率最终口径

接口响应速度以历史黑盒测试为依据：16 项黑盒测试全部通过，平均 HTTP 响应 31.15 ms，最大 104 ms。当前本机环境满足课堂验收、演示、小规模局域网试运行。

准确率分两层说明：

1. 规则判定准确率：围栏、滞留、方向、确认帧、场景信号等规则引擎用例全部通过，说明在输入检测框正确的前提下，规则层判定稳定。
2. 视觉模型准确率：项目尚未形成完整人工标注测试集，不虚构 Precision/Recall。正式商用建议按围栏、仓库、码头、负样本四类建立标注集，再统计 Precision、Recall、误报率和漏报率。

### 16.3 最终测试结论

当前版本可以完整顺利运行，具备“正式登录 + 权限 + 设备 + 告警闭环 + 回放 + 运维 + 部署脚本 + Word 交付文档”的完整项目形态。若用于真实商业交付，下一阶段重点是 24/72 小时稳定性、GPU 多路并发和真实标注集准确率测试。
""",
    "接口说明": r"""
## 16. 本轮新增运维、安全与部署接口

### 16.1 运维健康接口

| 方法 | 路径 | 权限 | 说明 |
| --- | --- | --- | --- |
| GET | /ops/health | viewer/operator/admin/super_admin | 返回系统健康、磁盘空间、数据库、配置目录、日志文件、ffmpeg 状态 |

返回字段包括：status、generated_at、app_env、python、platform、paths、checks、disk、logs。

### 16.2 备份接口

| 方法 | 路径 | 权限 | 说明 |
| --- | --- | --- | --- |
| POST | /ops/backup | admin/super_admin | 创建 zip 备份，默认包含 config、SQLite 数据库、日志；可选择包含视频文件 |

请求示例：

```json
{"include_videos": false}
```

返回示例字段：backup_path、backup_name、include_videos、file_count、size_bytes。

### 16.3 清理接口

| 方法 | 路径 | 权限 | 说明 |
| --- | --- | --- | --- |
| POST | /ops/cleanup | admin/super_admin | 按保留天数清理 data/outputs、data/replay_clips、data/uploads/videos、data/backups |

请求示例：

```json
{"dry_run": true, "retention_days": 30, "replay_retention_days": 30, "backup_retention_days": 90}
```

说明：dry_run=true 时只返回候选文件，不删除；dry_run=false 时执行删除并写入操作审计。

### 16.4 Debug 接口生产限制

Debug 接口包括 /debug/login、/debug/upload-video、/debug/bind-video、/debug/bind-stream、/debug/restore-stream、/debug/ping、/debug/simulate。当前版本默认关闭，只有 ALLOW_DEBUG_TOOLS=1 或系统设置 allow_debug_tools=true 时才允许访问。

正式生产接入摄像头时应使用 /devices，不再依赖 /debug/bind-stream。

### 16.5 日志和配置相关接口

| 方法 | 路径 | 权限 | 说明 |
| --- | --- | --- | --- |
| GET | /logs/operations | 登录用户 | 操作审计，包括规则、用户、设备、告警、备份、清理 |
| GET | /logs/system | 登录用户 | 返回 app/error 日志尾部 |
| GET | /logs/system/files | 登录用户 | 返回日志文件大小和修改时间 |
| GET | /logs/system/{log_key} | 登录用户 | 读取指定日志尾部 |
| GET | /settings | viewer/operator/admin/super_admin | 读取保留天数、模型策略、通知、Debug 开关等 |
| POST | /settings | admin/super_admin | 保存系统设置 |

### 16.6 部署文件说明

| 文件 | 作用 |
| --- | --- |
| Dockerfile | 构建后端镜像，安装 Python 依赖、ffmpeg、OpenCV 运行库 |
| docker-compose.yml | 同时启动 backend 和 nginx frontend，并设置健康检查和 restart 策略 |
| deploy/nginx.conf | 静态页面、API、stream、replay 的反向代理配置 |
| scripts/run_supervisor.ps1 | Windows 本地进程守护脚本，后端或前端异常会重启 |
| scripts/install_windows_task.ps1 | 将守护脚本注册为 Windows 计划任务 |
| scripts/ops_maintenance.py | 本地健康检查、备份、清理脚本 |
| .env.example | 完整环境变量模板，包含安全、日志、流、模型、回放、外部 API |

## 17. 系统内部接口说明补充

系统内部调用链如下：

1. 前端 index.html 调用 /dashboard/overview、/alerts/history、/devices、/settings、/logs、/ops 等接口形成管理大屏。
2. module.html 调用 /stream/{camera_id} 获取 MJPEG 视频，调用 /api/config/camera/{camera_id}/region/{region_id} 保存区域，调用 /vision/backend/* 切换 YOLO/VLM。
3. replay.html 调用 /replay/resolve、/replay/info、/replay/download、/replay/detections、/replay/analyze、/replay/clip 完成回放定位和视频理解。
4. stream_service 读取摄像头流，调用 vision_backend_service 分析画面，再调用 rules_engine 生成告警。
5. storage_service 负责 SQLite 持久化：用户、会话、告警、视频分析、告警流程、设置、审计日志。
6. maintenance_service 负责健康状态、zip 备份和运行文件清理。

## 18. 外部接口和依赖补充

| 外部依赖 | 用途 | 失败时表现 | 处理建议 |
| --- | --- | --- | --- |
| 手机/网络摄像头 HTTP 或 RTSP | 实时视频源 | 画面提示等待或断流重连 | 检查同一局域网、防火墙、账号密码、App 是否运行 |
| ffmpeg | 裁剪回放短片段 | 回放页提示无法裁剪，降级分析原始视频 | 安装 ffmpeg 并加入 PATH，Docker 镜像已包含 |
| YOLO 模型文件 | 人员检测 | 模型异常提示，检测框为空 | 检查 YOLO_WEIGHTS_PATH 和 models 目录 |
| MiMo 视频理解 API | 回放语义分析 | 视频理解失败或使用本地兜底 | 检查 MIMO_API_KEY、BASE_URL、网络 |
| 大模型 API | Agent 问答 | Agent 降级为规则化回答 | 检查 API_KEY、BASE_URL、MODEL_NAME |

## 19. 接口安全要求

1. 生产环境保持 ALLOW_DEBUG_TOOLS=0。
2. 所有管理接口必须携带 Bearer Token。
3. 普通用户不能新增、删除、修改摄像头流地址。
4. replay/download 只能下载允许目录内的回放和上传视频。
5. API Key 不写入接口文档、日志和前端页面。
6. 对外部署时使用 HTTPS 和 Nginx 反向代理。
7. 真实 RTSP 地址如含账号密码，应只允许管理员查看和修改。
""",
    "用户手册": r"""
## 16. 本轮新增能力使用说明

### 16.1 使用守护脚本运行

如果不希望依赖手动打开多个窗口，可以使用守护脚本：

```powershell
Set-Location D:\Project
.\scripts\run_supervisor.ps1
```

该脚本会启动后端和前端，并每隔数秒检查：

1. 后端 /health 是否可访问。
2. 前端 /index.html 是否可访问。
3. 进程是否已退出。

如果发现异常，会自动停止旧进程并重新启动。

### 16.2 注册 Windows 计划任务

如果希望登录 Windows 后自动启动平台：

```powershell
Set-Location D:\Project
.\scripts\install_windows_task.ps1
```

启动任务：

```powershell
Start-ScheduledTask -TaskName "AI视频识别平台"
```

### 16.3 Docker 部署

已提供 Dockerfile 和 docker-compose.yml。部署命令：

```powershell
Set-Location D:\Project
docker compose up -d --build
```

访问地址：

```text
http://127.0.0.1:5500/index.html
```

Docker Compose 会启动 backend 和 frontend 两个服务，backend 带健康检查，服务异常后会根据 restart 策略自动恢复。

### 16.4 Nginx 反向代理

Nginx 示例配置位于 deploy/nginx.conf。它处理：

1. 前端静态页面。
2. /api、/auth、/alerts、/devices、/settings 等后端接口代理。
3. /stream/ 视频流代理，并关闭 buffering。
4. /replay/ 回放接口代理。

生产环境建议在 Nginx 上配置 HTTPS 证书。

## 17. 运维功能使用说明

### 17.1 健康检查

命令行方式：

```powershell
.\.venv\Scripts\python.exe scripts\ops_maintenance.py health
```

管理接口方式：登录后调用 GET /ops/health。

可以查看：磁盘剩余空间、数据库是否存在、配置目录是否存在、日志大小、ffmpeg 是否可用、Python 和操作系统信息。

### 17.2 创建备份

不包含视频的轻量备份：

```powershell
.\.venv\Scripts\python.exe scripts\ops_maintenance.py backup
```

包含上传视频和回放片段的完整备份：

```powershell
.\.venv\Scripts\python.exe scripts\ops_maintenance.py backup --include-videos
```

备份文件生成在：

```text
data/backups
```

备份内容默认包括：config 配置、SQLite 数据库、运行日志。完整备份会额外包含 data/outputs、data/replay_clips、data/uploads/videos。

### 17.3 清理过期文件

先预览：

```powershell
.\.venv\Scripts\python.exe scripts\ops_maintenance.py cleanup --retention-days 30 --replay-retention-days 30
```

确认后执行：

```powershell
.\.venv\Scripts\python.exe scripts\ops_maintenance.py cleanup --retention-days 30 --replay-retention-days 30 --apply
```

管理接口 /ops/cleanup 默认 dry_run=true，建议先预览再删除。

### 17.4 日志轮转

日志目录：

```text
data/runtime/logs
```

当前日志会根据 .env 中的参数自动轮转：

| 参数 | 默认值 | 说明 |
| --- | --- | --- |
| LOG_MAX_BYTES | 10485760 | 单个日志最大约 10MB |
| LOG_BACKUP_COUNT | 5 | 最多保留 5 个历史日志 |

## 18. Debug 工具使用限制

当前版本 Debug 接口默认关闭。生产环境不要打开。确需演示时，有两种方式：

1. 在 .env 设置 ALLOW_DEBUG_TOOLS=1 后重启后端。
2. 管理员登录系统设置页，打开“允许生产环境显示调试工具”。

使用完毕后应立即关闭。

正式摄像头接入请使用设备管理页面：新增设备、填写 stream、保存并热重载。

## 19. 断流恢复说明

实时视频流读取失败时，页面会显示等待或断流提示。后端连续读帧失败达到阈值后，会自动释放旧连接并重新打开视频源。相关配置：

| 参数 | 默认值 | 说明 |
| --- | --- | --- |
| STREAM_READ_FAILURE_REOPEN_AFTER | 20 | 连续失败多少次后重连 |
| STREAM_MAX_FPS | 12 | 输出最大帧率 |
| STREAM_DETECTION_INTERVAL | 4 | 每隔多少帧检测一次 |
| STREAM_JPEG_QUALITY | 76 | MJPEG JPEG 质量 |

如果长时间无法恢复，请检查摄像头供电、网络、RTSP 账号密码、手机摄像头 App 是否仍在运行。

## 20. 最终交付文件清单

| 文件 | 说明 |
| --- | --- |
| 测试报告.md / 测试报告.docx | 功能测试、性能测试、准确率口径、白盒黑盒测试结果 |
| 接口说明.md / 接口说明.docx | 系统内部接口、外部接口、安全、运维、部署接口说明 |
| 用户手册.md / 用户手册.docx | 启动、登录、设备、告警、回放、手机摄像头、运维、部署说明 |
| Dockerfile | 后端容器构建文件 |
| docker-compose.yml | 后端和前端容器编排 |
| deploy/nginx.conf | Nginx 反向代理示例 |
| scripts/run_supervisor.ps1 | Windows 本地守护启动 |
| scripts/install_windows_task.ps1 | Windows 计划任务安装 |
| scripts/ops_maintenance.py | 健康、备份、清理维护脚本 |

## 21. 最终验收建议

1. 执行 start_all_dev.bat 或 run_supervisor.ps1。
2. 打开首页，确认实时监控矩阵可访问。
3. 登录 admin / 123456。
4. 新增一个手机摄像头设备，绑定同一局域网视频流。
5. 进入模块页确认视频显示和断流提示正常。
6. 触发或模拟告警，完成确认、处理、误报、完成闭环。
7. 打开回放页，验证回放定位和视频理解。
8. 查看操作审计和系统日志。
9. 运行 ops_maintenance.py health 和 backup。
10. 执行单元测试，确认 44 项全部通过。
""",
}


MARKERS = {
    "测试报告": "## 16. 本轮最终补充测试结论",
    "接口说明": "## 16. 本轮新增运维、安全与部署接口",
    "用户手册": "## 16. 本轮新增能力使用说明",
}


def _shade(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def _cell_text(cell, text: str, *, bold: bool = False, white: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(8.5)
    run.bold = bold
    if white:
        run.font.color.rgb = RGBColor(255, 255, 255)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _style_doc(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)

    for style_name, size, color in [
        ("Heading 1", 15, "2E74B5"),
        ("Heading 2", 12.5, "2E74B5"),
        ("Heading 3", 11.5, "1F4D78"),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    header = section.header.paragraphs[0]
    header.text = f"AI 视频识别信号平台 | {title}"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor(100, 100, 100)

    footer = section.footer.paragraphs[0]
    footer.text = f"{DATE}  D:\\Project"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(100, 100, 100)


def _add_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10 if style else 10.5)


def _markdown_to_docx(markdown: str, title: str, out_path: Path) -> None:
    doc = Document()
    _style_doc(doc, title)
    lines = markdown.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1
            continue
        if line.startswith("# "):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(line[2:].strip())
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(21)
            run.bold = True
            run.font.color.rgb = RGBColor.from_string("0B2545")
            paragraph.paragraph_format.space_after = Pt(14)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), 1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), 2)
            i += 1
            continue
        if line.startswith("```"):
            i += 1
            code_lines: list[str] = []
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.2)
            paragraph.paragraph_format.space_after = Pt(6)
            run = paragraph.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)
            continue
        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|"):
            table_lines: list[str] = []
            while i < len(lines) and lines[i].startswith("|"):
                raw = lines[i]
                core = raw.replace("|", "").replace("-", "").replace(":", "").replace(" ", "")
                if core:
                    table_lines.append(raw)
                i += 1
            rows = [[cell.strip() for cell in row.strip("|").split("|")] for row in table_lines]
            if rows:
                col_count = max(len(row) for row in rows)
                table = doc.add_table(rows=len(rows), cols=col_count)
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for row_idx, row in enumerate(rows):
                    for col_idx in range(col_count):
                        value = row[col_idx] if col_idx < len(row) else ""
                        cell = table.cell(row_idx, col_idx)
                        _cell_text(cell, value, bold=row_idx == 0, white=row_idx == 0)
                        if row_idx == 0:
                            _shade(cell, "2E74B5")
                doc.add_paragraph()
            continue
        if line.startswith("- "):
            _add_paragraph(doc, line[2:].strip(), "List Bullet")
            i += 1
            continue
        if len(line) > 3 and line[0].isdigit() and line[1:3] == ". ":
            _add_paragraph(doc, line[3:].strip(), "List Number")
            i += 1
            continue
        _add_paragraph(doc, line.replace("`", "").replace("**", "").strip())
        i += 1
    doc.save(out_path)


def main() -> None:
    for title, supplement in SUPPLEMENTS.items():
        markdown_path = ROOT / f"{title}.md"
        docx_path = ROOT / f"{title}.docx"
        if markdown_path.exists():
            current = markdown_path.read_text(encoding="utf-8")
        else:
            current = f"# AI 视频识别信号平台{title}\n\n编写日期：{DATE}\n"

        marker = MARKERS[title]
        if marker in current:
            current = current.split(marker)[0].rstrip()
        updated = current.rstrip() + "\n\n" + supplement.strip() + "\n"
        markdown_path.write_text(updated, encoding="utf-8")
        _markdown_to_docx(updated, title, docx_path)
        print(f"updated {markdown_path.name} and {docx_path.name}")


if __name__ == "__main__":
    main()
