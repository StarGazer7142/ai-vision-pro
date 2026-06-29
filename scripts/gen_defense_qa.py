# -*- coding: utf-8 -*-
"""生成 AI-VISION PRO 答辩问答 Word 文档"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ===== 全局样式 =====
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.35
style.paragraph_format.space_after = Pt(4)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = '黑体'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    hs.font.bold = True
    if level == 1:
        hs.font.size = Pt(18)
    elif level == 2:
        hs.font.size = Pt(14)
    else:
        hs.font.size = Pt(12)

def add_qa(q_num, question, answer_parts):
    """添加一个问答对。answer_parts 是 [(小标题, 内容), ...] 或直接字符串列表"""
    # 问题
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(f'Q{q_num}. {question}')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x00, 0x52, 0x8A)
    # 答案
    p2 = doc.add_paragraph()
    run2 = p2.add_run('A: ')
    run2.bold = True
    run2.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
    if isinstance(answer_parts, str):
        run3 = p2.add_run(answer_parts)
        run3.font.size = Pt(11)
    else:
        first = True
        for part in answer_parts:
            if isinstance(part, tuple):
                title, content = part
                if not first:
                    p2 = doc.add_paragraph()
                run_title = p2.add_run(f'{title}\n')
                run_title.bold = True
                run_title.font.size = Pt(11)
                run_content = p2.add_run(content)
                run_content.font.size = Pt(11)
                first = False
            else:
                if not first:
                    p2 = doc.add_paragraph()
                run_c = p2.add_run(part)
                run_c.font.size = Pt(11)
                first = False

def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

# ===================================================================
# 封面
# ===================================================================
for _ in range(6):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_t = title_p.add_run('AI-VISION PRO 答辩问答集')
run_t.bold = True
run_t.font.size = Pt(28)
run_t.font.color.rgb = RGBColor(0x00, 0x3D, 0x6B)

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_s = subtitle_p.add_run('工业级智能视觉感知平台 — 可能提问与参考回答')
run_s.font.size = Pt(14)
run_s.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

for _ in range(4):
    doc.add_paragraph()

info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_info = info_p.add_run('项目版本：v0.3.0\n文档日期：2026年6月')
run_info.font.size = Pt(12)
run_info.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ===================================================================
# 目录页
# ===================================================================
doc.add_heading('目录', level=1)
toc_items = [
    '第一部分：项目概述与选题背景（Q1-Q5）',
    '第二部分：系统架构与技术选型（Q6-Q12）',
    '第三部分：核心算法详解（Q13-Q20）',
    '第四部分：数据库与存储设计（Q21-Q24）',
    '第五部分：安全机制与权限设计（Q25-Q29）',
    '第六部分：前端设计与交互（Q30-Q33）',
    '第七部分：部署与运维（Q34-Q37）',
    '第八部分：Agent 智能问答系统（Q38-Q41）',
    '第九部分：MiMo 视频理解集成（Q42-Q45）',
    '第十部分：项目亮点与创新点（Q46-Q48）',
    '第十一部分：测试与质量保障（Q49-Q51）',
    '第十二部分：不足与未来展望（Q52-Q54）',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ===================================================================
# 第一部分：项目概述与选题背景
# ===================================================================
doc.add_heading('第一部分：项目概述与选题背景', level=1)

add_qa(1,
    '请简要介绍一下这个项目是做什么的？',
    [
        ('项目定位',
         'AI-VISION PRO 是一个工业级智能视觉感知平台，面向园区安防、仓库监控、码头管理等场景，'
         '提供多源摄像头视频流接入、AI目标检测与跟踪、规则引擎智能告警、视频回放定位、'
         'MiMo 视频理解分析以及 Agent 智能问答等核心能力。'),
        ('核心指标',
         '系统支持 3 个摄像头同时接入、4 条检测规则、75 个 RESTful API 端点、10 张 SQLite 数据表、'
         '4 个前端页面，具备完整的检测-告警-回放-分析闭环。'),
        ('业务价值',
         '传统安防依赖人工盯屏，效率低、漏报率高。本系统通过 AI 自动识别翻越围栏、人员滞留等异常行为，'
         '实现 7×24 小时不间断智能监控，有效降低人工巡检成本，提升安全响应速度。'),
    ])

add_qa(2,
    '为什么选择这个课题？有哪些实际应用背景？',
    [
        ('行业痛点',
         '据统计，传统安防系统中人工盯屏的有效注意力维持时间不足 20 分钟，'
         '长时间监控后漏报率显著上升。工业园区、仓库、码头等场所的安全管理面临人力不足、响应滞后等挑战。'),
        ('技术趋势',
         '随着边缘计算和深度学习模型（如 YOLO 系列）的成熟，实时目标检测已经可以在普通硬件上运行，'
         '结合视频理解大模型（如 MiMo），可以实现从"看得见"到"看得懂"的跨越。'),
        ('应用场景',
         '本系统针对园区围栏入侵检测、仓库人员滞留监控、码头作业安全管理三个典型场景，'
         '覆盖了工业安防中最常见的两类异常行为检测需求。'),
    ])

add_qa(3,
    '项目的主要功能模块有哪些？',
    [
        ('功能全景',
         '系统包含以下核心功能模块：\n'
         '① 实时监控矩阵 —— 4 宫格摄像头画面展示，实时检测框叠加\n'
         '② 告警事件大屏 —— 告警列表、筛选、详情查看、工作流管理\n'
         '③ 场景监控模块 —— 单摄像头详情页，支持双引擎热切换（YOLO/MiMo）\n'
         '④ 视频回放定位 —— 事件录像自动定位、短视频裁剪、MiMo 视频分析\n'
         '⑤ 数据可视化看板 —— 告警统计、趋势图表、设备状态\n'
         '⑥ 智能决策 Agent —— 自然语言问答，支持态势总结、告警分析\n'
         '⑦ 系统设置中心 —— 参数配置、运维操作、日志查看'),
    ])

add_qa(4,
    '项目中提到的"双引擎"是什么意思？为什么要做双引擎？',
    [
        ('双引擎定义',
         '双引擎指的是系统支持两种视觉分析方案并可运行时热切换：\n'
         '方案一：YOLO 目标检测（本地推理，速度快，擅长定位和计数）\n'
         '方案二：MiMo 视频理解模型（云端 API，语义理解能力强，擅长场景描述和行为判断）'),
        ('设计动机',
         'YOLO 擅长精确检测和定位，但无法理解"人在做什么"；'
         'MiMo 可以理解视频内容，但推理延迟较高。两者互补，'
         '通过 vision_backend.yaml 支持按摄像头粒度独立配置后端方案。'),
        ('热切换机制',
         '前端模块页提供 YOLO/VLM 开关按钮，点击后调用 /vision/backend/config 接口，'
         '后端更新 vision_backend.yaml 并实时刷新路由，无需重启服务。'),
    ])

add_qa(5,
    '项目的技术难度体现在哪些地方？',
    [
        ('实时性挑战',
         '需要在保持 18fps 以上帧率的同时完成目标检测和规则判定，'
         '每 2 帧执行一次检测以平衡性能与精度。'),
        ('算法融合',
         '需要将目标检测（YOLO）、多目标跟踪（IoU Greedy）、几何计算（叉积/射线法）'
         '三种不同类型的算法有机融合为一条完整的检测流水线。'),
        ('多模型集成',
         'YOLO 本地推理 + MiMo 云端 API + DeepSeek LLM 三种 AI 模型协同工作，'
         '需要处理异步调用、超时降级、结果融合等复杂问题。'),
        ('系统完整性',
         '涵盖前后端、数据库、流媒体、视频处理、安全认证、容器化部署等多个技术领域，'
         '对全栈开发能力要求较高。'),
    ])

doc.add_page_break()

# ===================================================================
# 第二部分：系统架构与技术选型
# ===================================================================
doc.add_heading('第二部分：系统架构与技术选型', level=1)

add_qa(6,
    '请介绍系统的整体技术架构。',
    [
        ('架构模式',
         '系统采用前后端分离 + Nginx 反向代理架构：\n'
         '浏览器 → Nginx (端口80) → FastAPI + Uvicorn (端口8000) → SQLite\n\n'
         'Nginx 同时承担静态文件服务（前端 4 个 HTML 页面）和 API 反向代理（16 个路径前缀转发至后端）。'),
        ('后端分层',
         '后端采用经典三层架构：\n'
         '① API 网关层：backend/app/api/routes.py，75 个 FastAPI 路由\n'
         '② 服务层：backend/app/services/，13 个业务服务模块\n'
         '③ 数据层：SQLite + PyYAML 配置文件'),
        ('前端设计',
         '前端为纯静态 HTML + CSS + JavaScript（无构建工具），共 4 个页面：\n'
         'index.html（总览）、module.html（场景详情）、replay.html（回放）、debug.html（调试）。'),
    ])

add_qa(7,
    '为什么选择 Python + FastAPI 作为后端技术栈？',
    [
        ('FastAPI 优势',
         '① 异步原生支持：基于 ASGI，天然适合流媒体和视频处理场景\n'
         '② 高性能：Uvicorn + uvloop，实测 QPS 可满足多路摄像头并发\n'
         '③ 自动文档：Swagger UI 自动生成 API 文档，降低前后端联调成本\n'
         '④ Pydantic 集成：请求/响应数据校验零成本'),
        ('生态契合',
         'Python 是 AI/ML 领域的主流语言，YOLO（Ultralytics）、OpenCV、NumPy 等核心依赖均为 Python 生态，'
         '选择 Python 避免了跨语言调用的复杂性。'),
        ('对比其他方案',
         '相比 Flask：FastAPI 原生支持异步和类型校验，更适合实时流场景。\n'
         '相比 Django：FastAPI 更轻量，无 ORM 绑定，与 SQLite 直接配合更灵活。\n'
         '相比 Go/Java：Python 的 AI 生态成熟度远超其他语言。'),
    ])

add_qa(8,
    '为什么选择 SQLite 而不是 MySQL 或 PostgreSQL？',
    [
        ('适用场景',
         '本系统面向单机/小型部署场景（3-10 路摄像头），SQLite 完全满足并发需求，'
         '且具有以下优势：'),
        ('SQLite 优势',
         '① 零配置：无需安装独立数据库服务，部署成本极低\n'
         '② 嵌入式：数据库文件就是一个 .db 文件，备份迁移极其方便\n'
         '③ WAL 模式：启用 WAL 后支持并发读写，实测可承载 120 QPS\n'
         '④ 性能：对于本系统的读多写少场景，SQLite 的性能优于网络数据库'),
        ('可扩展性',
         '如需扩展到多机部署，可通过 SQLAlchemy ORM 层抽象，'
         '仅需修改连接字符串即可迁移到 PostgreSQL，无需改动业务代码。'),
    ])

add_qa(9,
    '后端的 13 个服务模块是如何划分的？各模块职责是什么？',
    [
        ('核心检测链路（4个）',
         '① yolo_service.py —— YOLO 目标检测，模型加载、推理、类别归一化\n'
         '② tracking_service.py —— IoU 贪心多目标跟踪，轨迹管理\n'
         '③ rules_engine.py —— 规则引擎，越界检测和滞留检测判定\n'
         '④ stream_service.py —— MJPEG 实时流，帧读取、检测叠加、流推送'),
        ('数据与存储（3个）',
         '⑤ storage_service.py —— SQLite 持久化，告警、用户、会话、日志等 10 张表\n'
         '⑥ replay_service.py —— 视频回放定位、ffmpeg 短视频裁剪\n'
         '⑦ maintenance_service.py —— 系统维护、备份、健康检查、数据清理'),
        ('AI 增强（4个）',
         '⑧ vision_backend_service.py —— YOLO/MiMo 双引擎管理与热切换\n'
         '⑨ mimo_video_client.py —— MiMo 视频理解 API 客户端\n'
         '⑩ llm_client.py —— LLM API 客户端（DeepSeek）\n'
         '⑪ agent_orchestrator.py + agent_tools.py + agent_policy.py —— Agent 智能问答编排'),
        ('运维支撑（2个）',
         '⑫ agent_service.py —— Agent 服务入口\n'
         '⑬ 其余辅助模块 —— 配置加载、日志、工具函数'),
    ])

add_qa(10,
    '项目用了哪些主要的第三方库？各自的版本和用途是什么？',
    [
        ('AI/ML 核心',
         'Ultralytics 8.4.41 —— YOLO 目标检测框架\n'
         'PyTorch —— 深度学习推理引擎\n'
         'OpenCV 4.9.0.80 (headless) —— 图像处理、视频编解码\n'
         'Supervision 0.18.0 —— 检测结果可视化标注\n'
         'NumPy 1.26.4 —— 数值计算'),
        ('Web 框架',
         'FastAPI >=0.109.2 —— 异步 Web API 框架\n'
         'Uvicorn >=0.27.1 —— ASGI 服务器\n'
         'Pydantic >=2.0 —— 数据校验与序列化\n'
         'PyYAML 6.0.1 —— YAML 配置文件解析'),
        ('视频处理',
         'imageio-ffmpeg 0.6.0 —— 视频裁剪（ffmpeg 封装）\n'
         'ffmpeg —— 系统级视频转码/裁剪工具'),
        ('安全',
         'python-dotenv 1.2.2 —— .env 环境变量加载\n'
         'hashlib (内置) —— PBKDF2 密码哈希'),
    ])

add_qa(11,
    '项目如何处理前后端通信？RESTful API 设计有什么特点？',
    [
        ('通信方式',
         '前后端通过 RESTful JSON API 通信。前端通过 fetch() 调用后端接口，'
         '后端返回统一 JSON 格式：{"ok": true/false, "data": ..., "detail": "..."}'),
        ('API 设计规范',
         '① 资源命名：/alerts, /config/cameras, /users 等复数名词\n'
         '② HTTP 方法：GET 查询、POST 创建/更新、DELETE 删除\n'
         '③ 查询参数：分页 ?limit=&offset=，过滤 ?severity=&camera_id=\n'
         '④ 认证：Bearer Token 放在 Authorization Header\n'
         '⑤ 错误码：400 参数错误、401 未认证、403 权限不足、404 资源不存在、429 限流'),
        ('接口统计',
         '共 75 个 API 端点，覆盖健康检查、认证、配置、告警、信号、设备、'
         '回放、流媒体、Agent、运维、调试等功能域。'),
    ])

add_qa(12,
    'Nginx 在项目中承担什么角色？如何配置的？',
    [
        ('双重角色',
         '① 静态文件服务器：托管前端 4 个 HTML 页面及其依赖的 CSS/JS/字体\n'
         '② 反向代理：将 /health、/config、/alerts、/stream 等 16 个 API 路径前缀转发至 127.0.0.1:8000'),
        ('关键配置',
         '• proxy_read_timeout: 300s —— 支持长连接流媒体\n'
         '• proxy_buffering: off —— 流媒体代理关闭缓冲\n'
         '• try_files $uri $uri/ /index.html —— SPA 兜底路由\n'
         '• 端口 80 对外暴露，Docker 部署时统一入口'),
        ('为什么用 Nginx',
         '单用 FastAPI 直接托管静态文件性能较差，且无法利用 Nginx 的高效 I/O 模型。'
         'Nginx 处理静态文件的能力远超 Python，同时反向代理可以隐藏后端端口，提升安全性。'),
    ])

doc.add_page_break()

# ===================================================================
# 第三部分：核心算法详解
# ===================================================================
doc.add_heading('第三部分：核心算法详解', level=1)

add_qa(13,
    '请详细介绍 YOLO 目标检测算法在项目中的应用。',
    [
        ('模型选择',
         '使用 Ultralytics YOLOv8，支持多个权重文件：\n'
         '• yolov8n.pt —— 默认轻量模型，适合实时检测（~5ms/帧）\n'
         '• yolov8s.pt —— 小型模型，精度更高\n'
         '• yolo26s.pt —— YOLOv26 变体\n'
         '• best.pt —— 自训练模型，针对特定场景优化'),
        ('检测流程',
         '① 视频帧输入 → ② 预处理（Resize 到 imgsz=512）→ ③ YOLO 推理\n'
         '→ ④ NMS 后处理 → ⑤ 类别归一化映射 → ⑥ 返回 Detection 列表'),
        ('类别归一化',
         '将 YOLO 原始 80 类映射为 3 大安防类别：\n'
         '• person 组：person, human, pedestrian, man, woman, worker, head\n'
         '• vehicle 组：bicycle, car, motorcycle, bus, truck, auto, van\n'
         '• animal 组：bird, cat, dog, horse, sheep, cow, pet'),
        ('性能优化',
         '• 每 2 帧执行一次检测（DETECTION_INTERVAL=2），中间帧复用上次结果\n'
         '• 预览模式使用 imgsz=512 降低推理耗时\n'
         '• 置信度阈值 0.22 过滤低质量检测'),
    ])

add_qa(14,
    '项目中的多目标跟踪算法是如何实现的？',
    [
        ('算法选择',
         '采用 IoU Greedy Tracker（IoU 贪心匹配跟踪器），这是一种基于检测框重叠度的轻量级跟踪算法。'),
        ('核心参数',
         'match_thresh = 0.15 —— IoU 匹配阈值（极低，确保轻微重叠即可关联）\n'
         'track_buffer = 60 —— 允许丢失 60 帧仍保持 ID\n'
         'max_age_seconds = 5.0 —— 轨迹最大存活 5 秒'),
        ('算法流程',
         '① 计算当前帧所有检测框与已有轨迹的 IoU 矩阵\n'
         '② 贪心匹配：按 IoU 从高到低依次分配，IoU > 0.15 即匹配成功\n'
         '③ 未匹配的轨迹 missed +1，超过 buffer 或时间限制则清除\n'
         '④ 未匹配的检测框创建新轨迹\n'
         '⑤ 返回所有活跃轨迹的 track_id + bbox'),
        ('为什么选贪心而非匈牙利',
         '贪心匹配时间复杂度 O(N*M*log(N*M))，远优于匈牙利算法 O(N^3)。'
         '在安防场景中检测框数量有限（<50），贪心算法的精度损失可忽略，但性能优势明显。'),
    ])

add_qa(15,
    '越界检测（Boundary）的原理是什么？如何判断"翻越围栏"？',
    [
        ('几何原理 —— 叉积法',
         '① signed_distance_to_line()：计算点到直线的有符号距离（二维叉积）\n'
         '   d = (x2-x1)(y0-y1) - (y2-y1)(x0-x1)\n'
         '   d > 0 在线一侧，d < 0 在另一侧，d = 0 在线上\n\n'
         '② segment_crossed_line()：判断前后帧的 track 中心点是否从线的一侧穿越到另一侧\n'
         '   即 d_prev 和 d_curr 符号相反 → 发生越界'),
        ('防误报机制',
         '① confirm_frames = 4：连续 4 帧确认越界才触发告警\n'
         '② min_vertical_speed = 0.015：竖直速度门控，正常行走竖直速度极低\n'
         '③ max_aspect_ratio_change = 0.15：宽高比变化门控，翻越时身体拉伸/压缩\n'
         '④ trigger_side = neg：只检测从特定方向穿越'),
        ('检测框交叉判定',
         'bbox_intersects_line() 同时检查：\n'
         '• 线段端点是否在检测框内\n'
         '• 线段是否与检测框四条边相交\n'
         '两者满足任一即判定框与线相交'),
    ])

add_qa(16,
    '滞留检测（Dwell）是如何判断"人员在区域内停留超时"的？',
    [
        ('射线法（Ray-Casting）',
         'point_in_polygon() 使用射线法判断点是否在多边形区域内：\n'
         '从目标点向右发射一条水平射线，计算射线与多边形各边的交点个数。\n'
         '交点为奇数 → 点在多边形内；交点为偶数 → 点在多边形外。'),
        ('滞留时间累计',
         '每个跟踪目标（track_id）在每帧判断中心点是否在区域内：\n'
         '• 在区域内 → dwell_time += 1/frame_rate\n'
         '• 离开区域 → dwell_time 重置为 0\n'
         '• dwell_time >= threshold_seconds（默认5秒）→ 触发告警'),
        ('告警抑制',
         '• cooldown_seconds = 5：同一目标触发后 5 秒内不重复告警\n'
         '• cumulative_triggered_tracks：记录已触发的 track_id 集合\n'
         '• 目标离开区域后从集合移除，再次进入可重新触发'),
    ])

add_qa(17,
    '为什么选择叉积法做越界检测、射线法做滞留检测？有没有考虑过其他方案？',
    [
        ('叉积法选择理由',
         '越界检测的本质是判断"是否穿越了一条线"。叉积法只需要一次乘法和减法即可判断点在直线哪一侧，'
         '时间复杂度 O(1)，非常适合逐帧实时计算。相比凸包检测或深度学习方法，叉积法无需训练数据，'
         '规则可配置，且结果完全可解释。'),
        ('射线法选择理由',
         '滞留检测需要判断"点是否在任意多边形内"。射线法可以处理任意形状的凸/凹多边形，'
         '实现简单，时间复杂度 O(N)（N 为多边形边数，通常 < 10）。'
         '相比扫描线法或 winding number 法，射线法在实现简洁性和性能之间取得了最佳平衡。'),
        ('未选方案分析',
         '• 深度学习方案：需要大量标注数据，推理延迟高，不适合实时场景\n'
         '• 背景减除法：对光照变化敏感，无法区分具体目标\n'
         '• 光流法：计算量大，且难以得到结构化的区域判定结果'),
    ])

add_qa(18,
    'MJPEG 实时流是如何生成和推送的？',
    [
        ('帧处理流水线',
         '① 从摄像头读取原始帧（OpenCV VideoCapture）\n'
         '② 每 2 帧执行一次 YOLO 检测（降低计算负载）\n'
         '③ 将检测结果（bbox + class + confidence）叠加到帧上\n'
         '④ 编码为 JPEG（质量 72，平衡清晰度和带宽）\n'
         '⑤ 以 multipart/x-mixed-replace 格式推送到浏览器'),
        ('叠加可视化',
         '检测框按类别着色：人员(蓝色)、车辆(橙色)、动物(紫色)\n'
         'HUD 信息包括：摄像头 ID、检测数量、当前时间、帧率\n'
         '置信度阈值 0.22 以下的检测不显示叠加框'),
        ('异常处理',
         '• 读取失败计数器：连续 20 次失败后自动重新打开流\n'
         '• 流中断自动重连，无需人工干预\n'
         '• 最大帧率限制 18fps，防止 CPU 过载'),
    ])

add_qa(19,
    '项目的检测流水线从摄像头输入到告警输出，完整的数据流是什么？',
    [
        ('完整数据流',
         '摄像头视频流\n'
         '  ↓ (OpenCV VideoCapture)\n'
         '原始帧\n'
         '  ↓ (每2帧执行检测)\n'
         'YOLO 目标检测 → Detection 列表 [bbox, class, confidence, track_id]\n'
         '  ↓\n'
         'IoU 多目标跟踪 → 轨迹列表 [track_id, bbox, age, missed]\n'
         '  ↓\n'
         '规则引擎（逐规则判定）\n'
         '  ├─ 越界检测：叉积法判断是否穿越警戒线\n'
         '  └─ 滞留检测：射线法判断是否在区域内超时\n'
         '  ↓ (触发条件满足)\n'
         '告警事件 → 写入 SQLite alerts 表\n'
         '  ↓\n'
         '前端告警大屏实时刷新 + Agent 系统感知'),
    ])

add_qa(20,
    '如何处理检测误报？有哪些降低误报率的策略？',
    [
        ('多帧确认',
         '越界检测要求连续 4 帧（confirm_frames=4）确认才触发告警，'
         '单帧抖动不会误报。滞留检测阈值默认 5 秒，短暂停留不会触发。'),
        ('行为过滤',
         '翻越行为通过竖直速度门控（min_vertical_speed=0.015）和宽高比变化门控'
         '（max_aspect_ratio_change=0.15）区分正常行走和翻越动作。'),
        ('冷却机制',
         'cooldown_seconds=5：同一目标触发告警后 5 秒内不重复告警，'
         '避免同一事件产生大量重复告警。'),
        ('MiMo 验证',
         '可选的 MiMo 视频理解模型可以对告警事件进行二次验证，'
         '通过语义分析判断是否为真实异常，进一步降低误报。'),
    ])

doc.add_page_break()

# ===================================================================
# 第四部分：数据库与存储设计
# ===================================================================
doc.add_heading('第四部分：数据库与存储设计', level=1)

add_qa(21,
    '数据库中有几张表？请介绍主要表的结构和设计思路。',
    [
        ('表总览',
         '共 10 张表：alerts（告警）、signal_snapshots（信号快照）、'
         'ingest_frames（检测帧统计）、users（用户）、auth_sessions（会话）、'
         'operation_logs（操作日志）、video_analyses（视频分析）、'
         'alert_workflows（告警工作流）、system_settings（系统设置）、'
         'camera_devices（摄像头设备）'),
        ('核心表设计',
         '• alerts：告警主表，字段包括 timestamp, scene_ids, rule_id, camera_id, '
         'category, severity, message, video_analysis_summary 等\n'
         '• users：用户表，采用 PBKDF2-HMAC-SHA256 密码哈希 + 随机 salt\n'
         '• auth_sessions：基于 Token 的会话管理，支持过期自动清理'),
        ('设计原则',
         '① 适度冗余：告警表直接冗余 scene_ids、rule_label 等字段，避免频繁 JOIN\n'
         '② 索引优化：时间戳、场景 ID、规则 ID 等高频查询字段均建立索引\n'
         '③ JSON 灵活字段：signal_snapshots 使用 payload_json 存储结构化信号数据'),
    ])

add_qa(22,
    '为什么使用 SQLite 的 WAL 模式？有什么好处？',
    [
        ('WAL 原理',
         'WAL（Write-Ahead Logging）模式将写操作先写入日志文件，'
         '而不是直接修改数据库文件。读操作可以并发读取数据库快照，不受写操作阻塞。'),
        ('性能优势',
         '① 并发读写：读操作不会被写操作阻塞，适合多摄像头并发写入告警\n'
         '② 写入性能：顺序写 WAL 文件比随机写数据库文件快 5-10 倍\n'
         '③ 事务原子性：通过 WAL 日志保证事务的原子性和持久性'),
        ('适用场景',
         '本系统的典型负载是"多读少写"：多个前端页面同时查询告警数据，'
         '后端写入频率较低（仅在触发告警时写入）。WAL 模式完美匹配这一特征。'),
    ])

add_qa(23,
    '系统如何处理数据库的并发访问？',
    [
        ('线程安全机制',
         '使用 threading.Lock() 保护 SQLite 连接池，确保同一时刻只有一个线程执行写操作。'),
        ('连接管理',
         '每次操作获取锁 → 打开连接 → 执行 SQL → 提交 → 释放锁。'
         '使用 contextmanager 确保连接总是被正确关闭。'),
        ('过期清理',
         'auth_sessions 表中的过期会话会在每次登录时自动清理，'
         '避免会话表无限膨胀。'),
    ])

add_qa(24,
    '项目中的配置文件（YAML）和数据库是如何分工的？',
    [
        ('YAML 负责静态/半静态配置',
         '• config/rules.yaml：场景、摄像头、规则定义（结构化，支持热重载）\n'
         '• config/tracker.yaml：跟踪算法参数\n'
         '• config/vision_backend.yaml：视觉后端切换配置\n'
         '• .env：环境变量、API Key 等敏感配置'),
        ('SQLite 负责动态/运行时数据',
         '• 告警记录、操作日志（时间序列数据）\n'
         '• 用户账号、认证会话（需要事务支持）\n'
         '• 视频分析结果（结构化 JSON 存储）\n'
         '• 系统设置（键值对存储）'),
        ('分工原则',
         'YAML 适合人类可读、版本控制友好的配置；SQLite 适合运行时动态数据。'
         '两者通过 reload 机制保持同步。'),
    ])

doc.add_page_break()

# ===================================================================
# 第五部分：安全机制与权限设计
# ===================================================================
doc.add_heading('第五部分：安全机制与权限设计', level=1)

add_qa(25,
    '系统的认证机制是如何设计的？',
    [
        ('Token 认证流程',
         '① 用户提交账号密码 → ② 后端验证密码哈希 → ③ 生成随机 Token\n'
         '→ ④ Token 存入 auth_sessions 表 → ⑤ 返回给前端存入 localStorage\n'
         '→ ⑥ 后续请求携带 Authorization: Bearer <token>'),
        ('Token 管理',
         '• Token 有效期 12 小时（SESSION_HOURS_DEFAULT）\n'
         '• 支持 last_seen_at 追踪活跃度\n'
         '• 登出时主动删除 Token\n'
         '• 过期 Token 在下次登录时自动清理'),
        ('密码安全',
         '• PBKDF2-HMAC-SHA256 哈希，迭代 120,000 次\n'
         '• 随机 salt（secrets.token_hex(16)）\n'
         '• 比较时使用 secrets.compare_digest() 防时序攻击\n'
         '• 最低密码长度 6 位'),
    ])

add_qa(26,
    '项目的 RBAC 权限模型是如何设计的？',
    [
        ('角色定义',
         '四种角色，权限从高到低：\n'
         '• super_admin：超级管理员，所有操作权限\n'
         '• admin：管理员，可管理用户、配置、告警\n'
         '• operator：操作员，可查看告警和执行基本操作\n'
         '• viewer：只读用户，仅可查看数据'),
        ('权限控制实现',
         '后端通过 _require_roles(authorization, {allowed_roles}) 装饰器检查：\n'
         '• /auth/register, /settings POST → 仅 super_admin 和 admin\n'
         '• /users → super_admin, admin, operator, viewer\n'
         '• /config/reload → 仅 super_admin 和 admin'),
        ('前端权限适配',
         '前端根据 currentUser.role 动态显示/隐藏功能：\n'
         '• 非管理员不可见"系统后台"导航组\n'
         '• 非管理员不可见"注册管理员"按钮\n'
         '• 引擎开关按钮需要已登录状态'),
    ])

add_qa(27,
    '系统有哪些安全防护措施？',
    [
        ('速率限制',
         '每个 IP 每 60 秒最多 120 次请求，超过限制返回 HTTP 429。'
         '防止恶意刷接口和 DDoS 攻击。'),
        ('登录防护',
         '连续 5 次登录失败（LOGIN_FAILURE_LIMIT）后锁定 15 分钟（LOGIN_LOCK_MINUTES），'
         '防止暴力破解密码。'),
        ('CORS 白名单',
         '仅允许 127.0.0.1:5500、localhost:5500、127.0.0.1:80、localhost:80 四个源，'
         '防止跨域请求攻击。'),
        ('文件上传限制',
         '最大上传 500MB，仅允许 .mp4/.avi/.mov/.mkv/.webm 格式，防止恶意文件上传。'),
        ('Agent 只读策略',
         'Agent 工具仅允许 4 个只读操作（get_runtime_status, get_alert_summary, '
         'get_replay_hint, analyze_replay_video），禁止修改任何系统配置。'),
    ])

add_qa(28,
    '如果系统部署到生产环境，安全性有哪些增强？',
    [
        ('环境变量控制',
         '• ALLOW_DEBUG_TOOLS=0：禁用 /debug/* 调试接口\n'
         '• VISION_BACKEND_MODE：锁定视觉后端，禁止运行时切换'),
        ('Docker 隔离',
         '• 容器内进程以非 root 用户运行\n'
         '• 模型文件只读挂载（models:ro）\n'
         '• 数据卷持久化但不暴露内部路径'),
        ('网络安全',
         '• Nginx 隐藏后端端口（8000 不对外暴露）\n'
         '• 仅 80 端口对外\n'
         '• 生产环境建议加 HTTPS（TLS 终止在 Nginx）'),
    ])

add_qa(29,
    '密码哈希为什么选择 PBKDF2 而不是 bcrypt 或 Argon2？',
    [
        ('PBKDF2 优势',
         '① Python 标准库内置（hashlib.pbkdf2_hmac），无需额外依赖\n'
         '② 120,000 次迭代 + 随机 salt，安全强度满足工业标准\n'
         '③ 跨平台兼容性好，不依赖系统级库'),
        ('对比分析',
         '• bcrypt：需要 bcrypt 库，Windows 兼容性偶尔有问题\n'
         '• Argon2：需要 argon2-cffi 库，是更新更安全的选择，但依赖更多\n'
         '• PBKDF2：零依赖、标准化、安全性足够，是本项目的务实选择'),
    ])

doc.add_page_break()

# ===================================================================
# 第六部分：前端设计与交互
# ===================================================================
doc.add_heading('第六部分：前端设计与交互', level=1)

add_qa(30,
    '前端为什么选择纯 HTML + CSS + JavaScript，而不用 Vue 或 React？',
    [
        ('设计考量',
         '① 无构建依赖：无需 Node.js 构建链，部署只需一个静态文件服务器\n'
         '② 启动速度快：从零到可用只需 Nginx 或 Python http.server\n'
         '③ 答辩演示友好：代码结构清晰，评委可以直接查看 HTML 源码理解实现'),
        ('技术实现',
         '• CSS：CSS 变量 + Grid/Flexbox 布局，深色主题设计系统\n'
         '• JavaScript：原生 ES6+，无框架依赖\n'
         '• 字体：Inter + JetBrains Mono（代码）+ Noto Sans SC（中文）\n'
         '• 图标：Font Awesome 6.4.0'),
        ('效果对比',
         '实测页面加载速度 < 200ms，首屏渲染 < 500ms，远优于需要下载 JS bundle 的 SPA 方案。'),
    ])

add_qa(31,
    '前端 4 个页面分别承担什么功能？页面之间如何导航？',
    [
        ('页面功能',
         '① index.html —— 总览页：4 宫格实时监控矩阵 + 告警大屏 + 数据看板 + 系统设置\n'
         '② module.html —— 场景详情页：单摄像头监控 + 双引擎切换 + 区域配置 + 滞留阈值设置\n'
         '③ replay.html —— 视频回放页：事件定位 + 短视频裁剪 + MiMo 视频分析\n'
         '④ debug.html —— 调试页：视频注入测试 + 流恢复'),
        ('导航方式',
         '• 左侧导航栏：点击切换页面视图（index.html 内部 SPA 切换）\n'
         '• 监控矩阵点击：openDetail() 跳转到 module.html\n'
         '• 告警详情点击"查看回放"：跳转到 replay.html\n'
         '• 浏览器后退按钮：replay.html 支持 history.back()'),
    ])

add_qa(32,
    '深色主题的设计思路是什么？有哪些设计亮点？',
    [
        ('设计语言',
         '采用深色科技感设计，主色调：\n'
         '• 背景：#0b0f14（深空黑）\n'
         '• 卡片：#17202a（深灰蓝）\n'
         '• 强调色：#2fb3c8（青色）\n'
         '• 成功：#2aa876  警告：#e2a03f  危险：#e05263'),
        ('设计亮点',
         '① CSS 变量系统：统一管理颜色，支持主题切换\n'
         '② 毛玻璃效果：backdrop-filter: blur(12px)\n'
         '③ 渐变背景：radial-gradient 营造科技感\n'
         '④ 微动画：fadeIn、agentPulse、hover translateY\n'
         '⑤ 响应式：@media 断点适配 760px / 1080px / 1280px'),
    ])

add_qa(33,
    '前端如何与后端实现实时数据更新？',
    [
        ('轮询机制',
         '• 告警数据：fetchRealAlerts() 定时拉取最新告警\n'
         '• 视频流：MJPEG multipart 流持续推送\n'
         '• 后端状态：1.5 秒轮询 /health 接口'),
        ('流媒体方案',
         '视频流采用 MJPEG over HTTP，通过 <img> 标签的 src 属性持续接收：\n'
         '<img src="/stream/cam_fence?t=1234567890">\n'
         '服务端以 multipart/x-mixed-replace 格式推送 JPEG 帧，浏览器自动渲染。'),
        ('实时感知',
         '• Agent 对话：单次请求-响应，支持流式文本显示\n'
         '• 引擎切换：切换后立即重新拉取状态并刷新 UI\n'
         '• 告警过滤：前端客户端过滤，无需重新请求后端'),
    ])

doc.add_page_break()

# ===================================================================
# 第七部分：部署与运维
# ===================================================================
doc.add_heading('第七部分：部署与运维', level=1)

add_qa(34,
    '项目的 Docker 部署方案是怎样的？',
    [
        ('容器架构',
         '单容器部署（Nginx + FastAPI 共享容器）：\n'
         '• 基础镜像：python:3.10-slim\n'
         '• 系统依赖：ffmpeg + libgl1 + libglib2.0-0 + nginx\n'
         '• 启动方式：Nginx 前台运行 + Uvicorn 后台启动'),
        ('数据卷挂载',
         '• ./backend → /app/backend（后端代码，支持热更新）\n'
         '• ./frontend/static → /usr/share/nginx/html（前端静态文件）\n'
         '• ./data → /app/data（运行时数据持久化）\n'
         '• ./config → /app/config（配置文件）\n'
         '• ./models → /app/models:ro（模型权重，只读）'),
        ('健康检查',
         '每 15 秒请求 /health 接口，失败 3 次后重启容器。'
         '重启策略为 unless-stopped，服务器重启后自动恢复。'),
    ])

add_qa(35,
    '开发环境和生产环境有什么区别？如何切换？',
    [
        ('开发环境',
         '• 后端：uvicorn --reload（热重载，代码修改自动重启）\n'
         '• 前端：python -m http.server 5500\n'
         '• 数据库：SQLite 本地文件\n'
         '• 调试接口：/debug/* 可用'),
        ('生产环境',
         '• Docker Compose 一键部署\n'
         '• Nginx 统一入口（端口 80）\n'
         '• ALLOW_DEBUG_TOOLS=0 禁用调试接口\n'
         '• VISION_BACKEND_MODE 可锁定引擎'),
        ('环境变量区分',
         'APP_ENV=development/production 控制环境，'
         '.env 文件中设置不同的配置值。'),
    ])

add_qa(36,
    '项目如何进行日志管理和数据维护？',
    [
        ('日志系统',
         '• 应用日志：Python logging 模块，支持文件轮转（10MB/文件，5 个备份）\n'
         '• 操作日志：operation_logs 表记录所有管理操作\n'
         '• 系统日志：支持通过 /logs/system/files 查看历史日志文件'),
        ('数据维护',
         '• 告警保留：默认 30 天（RETENTION_DAYS）\n'
         '• 回放保留：默认 30 天（REPLAY_RETENTION_DAYS）\n'
         '• 自动清理：凌晨 3:00 执行系统深度重置\n'
         '• 手动清理：/ops/cleanup 接口支持手动触发'),
        ('备份机制',
         '/ops/backup 接口支持手动备份数据库，'
         '备份保留 90 天（BACKUP_RETENTION_DAYS）。'),
    ])

add_qa(37,
    'Nginx 的反向代理是如何配置的？为什么需要它？',
    [
        ('代理规则',
         '使用 location 正则匹配 16 个 API 路径前缀，'
         '将请求转发至 127.0.0.1:8000（FastAPI）。\n'
         '其余路径由 Nginx 直接返回静态文件。'),
        ('关键配置',
         '• proxy_read_timeout: 300s —— 支持长时间流式响应\n'
         '• proxy_buffering: off —— 禁用缓冲，实时推送流媒体\n'
         '• try_files $uri $uri/ /index.html —— SPA 兜底路由\n'
         '• gzip on —— 压缩静态资源，减少传输大小'),
        ('为什么需要',
         '① 安全：隐藏后端端口 8000，仅暴露 80\n'
         '② 性能：Nginx 处理静态文件效率远超 Python\n'
         '③ 功能：统一入口，简化前端 API 调用（无需写完整 URL）'),
    ])

doc.add_page_break()

# ===================================================================
# 第八部分：Agent 智能问答系统
# ===================================================================
doc.add_heading('第八部分：Agent 智能问答系统', level=1)

add_qa(38,
    'Agent 智能问答系统的设计架构是什么？',
    [
        ('混合式架构',
         '采用"本地工具执行 + 可选云端 LLM 增强"的混合架构：\n'
         '① 意图识别（本地关键词匹配）\n'
         '② 工具调用（本地函数执行，获取实时数据）\n'
         '③ LLM 增强（可选，调用 DeepSeek API 生成自然语言回答）\n'
         '④ 兜底机制（LLM 不可用时直接返回工具结果）'),
        ('设计优势',
         '• 低延迟：本地工具执行 < 50ms，无需等待网络\n'
         '• 高可靠：LLM 不可用时自动降级到本地模式\n'
         '• 只读安全：Agent 仅允许查询操作，禁止修改系统配置'),
    ])

add_qa(39,
    'Agent 是如何识别用户意图的？',
    [
        ('意图识别方法',
         '采用基于关键词匹配的本地意图识别（detect_intent_local），'
         '无需调用外部模型，响应时间 < 10ms。'),
        ('5 种意图类型',
         '① runtime —— 关键词：状态、运行、系统、稳定\n'
         '② alerts —— 关键词：告警、事件、异常、最近\n'
         '③ replay —— 关键词：回放、录像、视频、分析\n'
         '④ summary —— 关键词：总结、总览、态势、概况\n'
         '⑤ video_analysis —— 关键词：视频分析、MiMo'),
        ('工具映射',
         '每种意图映射到不同的工具组合：\n'
         '• runtime → get_runtime_status\n'
         '• alerts → get_alert_summary\n'
         '• replay → get_alert_summary + get_replay_hint + analyze_replay_video\n'
         '• summary → get_runtime_status + get_alert_summary + get_replay_hint'),
    ])

add_qa(40,
    'Agent 的工具系统是如何设计的？',
    [
        ('工具注册',
         'agent_tools.py 中定义了 4 个只读工具函数，'
         '每个工具接收上下文参数并返回结构化结果。'),
        ('安全策略',
         'agent_policy.py 定义了白名单机制：\n'
         '• 仅允许 4 个只读操作\n'
         '• 禁止任何写入/修改操作\n'
         '• 禁止访问敏感配置\n'
         '• 所有工具调用被记录到操作日志'),
        ('LLM 增强',
         '当 AGENT_ENABLE_LLM=1 时，工具执行结果会发送给 DeepSeek API，'
         '由 LLM 生成自然语言回答。LLM 不可用时自动回退到本地工具直接返回。'),
    ])

add_qa(41,
    'Agent 系统与传统的规则引擎有什么区别和联系？',
    [
        ('区别',
         '• 规则引擎：基于预定义规则（越界/滞留），处理结构化数据，输出确定性结果\n'
         '• Agent 系统：基于自然语言理解，整合多源数据，输出描述性回答\n'
         '• 规则引擎用于实时检测，Agent 用于事后分析和态势总结'),
        ('联系',
         '• Agent 调用规则引擎的运行时状态作为工具数据\n'
         '• Agent 可以触发视频分析（MiMo）获取语义理解\n'
         '• 两者共享同一套告警数据和配置'),
        ('协同工作',
         '规则引擎负责实时检测（毫秒级响应），Agent 负责综合分析和问答（秒级响应）。'
         '规则引擎是"眼睛"，Agent 是"大脑"。'),
    ])

doc.add_page_break()

# ===================================================================
# 第九部分：MiMo 视频理解集成
# ===================================================================
doc.add_heading('第九部分：MiMo 视频理解集成', level=1)

add_qa(42,
    'MiMo 视频理解模型是什么？为什么集成它？',
    [
        ('MiMo 简介',
         'MiMo 是小米推出的视频理解大模型（mimo-v2.5），'
         '能够理解视频内容并生成结构化的分析报告。'),
        ('集成动机',
         'YOLO 擅长检测"有什么"（目标检测与定位），但无法理解"在做什么"（行为语义）。'
         'MiMo 可以回答"视频中有几个人、穿什么衣服、在什么位置、正在做什么"，'
         '为安防分析提供语义级别的理解能力。'),
        ('输出格式',
         'MiMo 返回结构化 JSON：\n'
         '• summary：一句话总结（画面理解 + 规则判定）\n'
         '• observations：观察列表\n'
         '• events：事件列表（时间、标题、描述、严重度、置信度）\n'
         '• risk_assessment：风险评估\n'
         '• recommended_actions：建议措施'),
    ])

add_qa(43,
    'MiMo 在系统中的两种使用方式是什么？',
    [
        ('方式一：实时预览（降级模式）',
         '在 module.html 的实时监控中，MiMo 模式下的实时预览实际使用本地 YOLO 检测作为降级，'
         '因为 MiMo 的 API 延迟（1-5秒）无法满足实时帧处理需求。'),
        ('方式二：短视频分析（核心场景）',
         '在 replay.html 的回放分析中，系统裁剪出事件前后的短视频（默认各4秒），'
         '发送给 MiMo 进行语义分析。这是 MiMo 的核心使用场景。'),
        ('分析流程',
         '① 从告警事件定位视频文件 → ② ffmpeg 裁剪出 8 秒短视频\n'
         '→ ③ Base64 编码视频 → ④ 发送到 MiMo API → ⑤ 解析返回的 JSON\n'
         '→ ⑥ 存入 video_analyses 表'),
    ])

add_qa(44,
    'MiMo 和 YOLO 双引擎切换的技术实现是什么？',
    [
        ('配置驱动',
         'vision_backend.yaml 定义了：\n'
         '• default_backend：全局默认后端（yolo 或 video_understanding）\n'
         '• scene_overrides：按场景覆盖\n'
         '• camera_overrides：按摄像头覆盖（优先级最高）'),
        ('热切换流程',
         '前端点击 YOLO/VLM 开关 → POST /vision/backend/config\n'
         '→ 更新 vision_backend.yaml → 调用 refresh() 重新加载\n'
         '→ 后续分析请求使用新后端'),
        ('优先级链',
         'camera_overrides > scene_overrides > default_backend\n'
         '例如：cam_fence 可以用 yolo，cam_warehouse 用 video_understanding，'
         '互不影响。'),
    ])

add_qa(45,
    'MiMo 分析结果中有哪些特殊的后处理逻辑？',
    [
        ('误报过滤',
         '检测 MiMo 返回文本中的负面关键词（"未滞留"、"误报"、"误触发"等），'
         '如果分析结论为"无异常"，系统会在 summary 中补充规则判定信息。'),
        ('画面细节增强',
         '如果 MiMo 返回的 summary 缺少足够的画面细节（人数、衣着、位置、动作），'
         '系统会自动发起第二次请求，专门要求"画面理解"，并将两次结果合并。'),
        ('规则对齐',
         '对于滞留类规则，系统会检查 summary 中是否包含"滞留"或"停留"关键词，'
         '如果没有，会自动追加规则判定语句，确保分析结论与规则引擎一致。'),
    ])

doc.add_page_break()

# ===================================================================
# 第十部分：项目亮点与创新点
# ===================================================================
doc.add_heading('第十部分：项目亮点与创新点', level=1)

add_qa(46,
    '项目最大的技术亮点或创新点是什么？',
    [
        ('亮点一：双引擎热切换架构',
         '行业首创在同一平台内实现 YOLO 目标检测与 MiMo 视频理解的运行时热切换，'
         '支持按摄像头粒度独立配置，无需重启服务。'),
        ('亮点二：多算法融合的检测流水线',
         '将 YOLO 目标检测、IoU 贪心跟踪、叉积越界检测、射线法滞留检测四种算法'
         '无缝融合为一条完整的实时检测流水线，从帧输入到告警输出延迟 < 100ms。'),
        ('亮点三：混合式 Agent 架构',
         '本地意图识别 + 工具执行 + 可选 LLM 增强的混合架构，'
         '兼顾了低延迟和高智能，且 LLM 不可用时自动降级。'),
        ('亮点四：完整的工程化',
         '75 个 API、10 张数据表、73 个测试用例、Docker 部署方案、'
         '安全认证机制，是一个可以直接部署使用的完整系统。'),
    ])

add_qa(47,
    '项目与市面上已有的安防系统相比，有什么独特之处？',
    [
        ('对比传统安防',
         '传统安防系统通常只提供视频存储和人工回看，本系统实现了 AI 自动检测 + 告警 + 分析的闭环。'),
        ('对比纯 AI 检测方案',
         '市面上的 AI 检测方案通常只提供检测能力，本系统在此基础上增加了规则引擎、'
         '视频回放、MiMo 语义分析、Agent 问答等完整的业务功能。'),
        ('对比云端安防方案',
         '本系统支持本地部署，数据不出园区，满足数据安全合规要求。'
         '同时通过 MiMo API 提供可选的云端智能增强。'),
    ])

add_qa(48,
    '如果让你重新设计这个项目，会有哪些改进？',
    [
        ('技术改进',
         '① 引入 WebSocket 替代轮询，实现真正的实时推送\n'
         '② 使用 Redis 替代内存缓存，支持多实例部署\n'
         '③ 引入消息队列（如 RabbitMQ）解耦告警处理'),
        ('架构改进',
         '① 微服务化：将检测、告警、回放等拆分为独立服务\n'
         '② GPU 推理：支持 NVIDIA TensorRT 加速 YOLO 推理\n'
         '③ 边缘计算：在摄像头端部署轻量检测模型'),
        ('功能改进',
         '① 增加人脸识别和车牌识别能力\n'
         '② 支持移动端 App 推送告警\n'
         '③ 增加历史数据分析和趋势预测'),
    ])

doc.add_page_break()

# ===================================================================
# 第十一部分：测试与质量保障
# ===================================================================
doc.add_heading('第十一部分：测试与质量保障', level=1)

add_qa(49,
    '项目做了哪些测试？测试覆盖率如何？',
    [
        ('测试概览',
         '共 73 个单元/集成测试用例，100% 通过率。'),
        ('测试分类',
         '① 核心算法测试：IoU 跟踪器轨迹稳定性、滞留规则触发\n'
         '② 认证权限测试：10 个用例，覆盖登录/注册/权限/Token\n'
         '③ 设备管理测试：8 个用例，覆盖 CRUD 操作\n'
         '④ 规则引擎测试：10 个用例，覆盖越界/滞留判定\n'
         '⑤ 告警管理测试：7 个用例，覆盖创建/查询/工作流\n'
         '⑥ 视频回放测试：6 个用例，覆盖定位/裁剪/分析\n'
         '⑦ 流媒体测试：4 个用例，覆盖流推送/异常处理\n'
         '⑧ 前端页面测试：10 个用例，覆盖页面加载/API 调用\n'
         '⑨ 安全测试：8 个用例，覆盖限流/认证/文件上传\n'
         '⑩ 接口兼容性测试：6 个用例'),
        ('运行方式',
         'python -m unittest discover -s backend/tests -p "test_*.py" -v'),
    ])

add_qa(50,
    '测试中遇到过哪些典型 Bug？是如何修复的？',
    [
        ('典型 Bug 1：IoU 跟踪 ID 抖动',
         '问题：相邻帧检测框轻微偏移导致 track_id 频繁切换。\n'
         '修复：降低 match_thresh 到 0.15，增加 track_buffer 到 60 帧。'),
        ('典型 Bug 2：滞留误报',
         '问题：目标快速经过区域边缘时触发滞留告警。\n'
         '修复：增加 confirm_frames 确认帧数 + 目标必须完全进入区域。'),
        ('典型 Bug 3：告警风暴',
         '问题：同一目标触发大量重复告警。\n'
         '修复：增加 cooldown_seconds 冷却机制 + cumulative_triggered_tracks 去重。'),
    ])

add_qa(51,
    '项目有哪些质量保障措施？',
    [
        ('代码质量',
         '• 统一的代码风格和命名规范\n'
         '• Pydantic 数据模型校验\n'
         '• 类型注解（Python 3.10+ type hints）'),
        ('运行时质量',
         '• 健康检查接口 /health\n'
         '• 异常自动恢复（流中断重连、数据库重连）\n'
         '• 完善的日志记录和错误追踪'),
        ('安全质量',
         '• 73 个测试用例包含 8 个安全测试\n'
         '• 密码哈希 + Token 认证 + RBAC 权限\n'
         '• 速率限制 + 登录锁定'),
    ])

doc.add_page_break()

# ===================================================================
# 第十二部分：不足与未来展望
# ===================================================================
doc.add_heading('第十二部分：不足与未来展望', level=1)

add_qa(52,
    '项目当前存在哪些不足或局限性？',
    [
        ('技术局限',
         '① SQLite 不适合大规模并发写入，限制了系统扩展到数十路摄像头\n'
         '② IoU 跟踪在目标遮挡严重时可能出现 ID 混淆\n'
         '③ MiMo API 延迟较高（1-5秒），无法用于实时检测'),
        ('功能局限',
         '① 缺少人脸识别和车牌识别能力\n'
         '② 不支持移动端访问\n'
         '③ 缺少历史数据的趋势分析和预测功能'),
        ('部署局限',
         '① 单容器部署不适合高可用场景\n'
         '② 缺少 Kubernetes 编排支持\n'
         '③ 没有 CI/CD 自动化部署流水线'),
    ])

add_qa(53,
    '如果项目要扩展到更大规模的场景（如 100 路摄像头），需要做哪些改造？',
    [
        ('数据库升级',
         '从 SQLite 迁移到 PostgreSQL，使用 SQLAlchemy ORM 抽象数据层，'
         '仅修改连接字符串即可完成迁移。'),
        ('架构升级',
         '① 微服务化：检测、告警、回放等拆分为独立服务\n'
         '② 消息队列：使用 RabbitMQ/Kafka 解耦数据流\n'
         '③ Redis 缓存：共享检测结果和配置状态'),
        ('计算升级',
         '① GPU 推理：NVIDIA TensorRT 加速 YOLO\n'
         '② 边缘计算：在摄像头端部署轻量检测模型\n'
         '③ 分布式调度：多节点并行处理不同摄像头'),
    ])

add_qa(54,
    '项目未来可以增加哪些功能方向？',
    [
        ('AI 能力增强',
         '① 人脸识别：集成 InsightFace/ArcFace 实现人员身份识别\n'
         '② 车牌识别：集成 PaddleOCR 实现车辆管理\n'
         '③ 行为分析：增加打架检测、跌倒检测、奔跑检测等更多行为类型\n'
         '④ 趋势预测：基于历史数据预测高峰时段和高风险区域'),
        ('产品化方向',
         '① 移动端 App：React Native 跨平台 App，支持告警推送\n'
         '② 多租户：支持多个园区独立管理\n'
         '③ 报告生成：自动生成日报/周报/月报（PPT/Word 格式）'),
        ('技术演进',
         '① Kubernetes 部署：支持弹性伸缩\n'
         '② GPU 推理集群：支持大规模实时检测\n'
         '③ 数字孪生：3D 场景建模与可视化'),
    ])

# ===================================================================
# 附录：项目关键数据速查
# ===================================================================
doc.add_page_break()
doc.add_heading('附录：项目关键数据速查表', level=1)

data_table = [
    ['指标项', '数值', '说明'],
    ['API 端点数', '75 个', '覆盖 14 个功能域'],
    ['数据表数量', '10 张', 'SQLite WAL 模式'],
    ['前端页面', '4 个', '纯 HTML + CSS + JS'],
    ['后端服务模块', '13 个', '核心业务服务'],
    ['测试用例', '73 个', '100% 通过率'],
    ['检测规则', '4 条', '1 越界 + 3 滞留'],
    ['摄像头支持', '3 路', '围栏 + 码头 + 仓库'],
    ['默认模型', 'yolov8n', '可切换 yolov8s / yolo26s'],
    ['密码哈希迭代', '120,000 次', 'PBKDF2-HMAC-SHA256'],
    ['Token 有效期', '12 小时', '可配置'],
    ['速率限制', '120 次/分钟', '每 IP'],
    ['告警保留天数', '30 天', '可配置'],
    ['Docker 基础镜像', 'python:3.10-slim', '含 ffmpeg + nginx'],
    ['流最大帧率', '18 fps', '可配置'],
    ['检测置信度阈值', '0.22', '叠加显示阈值'],
    ['IoU 匹配阈值', '0.15', '贪心匹配'],
    ['轨迹缓冲帧数', '60 帧', '约 2.4 秒 @25fps'],
    ['滞留检测阈值', '5 秒', '可按规则配置'],
    ['越界确认帧数', '4 帧', '防误报'],
]

table = doc.add_table(rows=len(data_table), cols=3, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, row_data in enumerate(data_table):
    for j, cell_text in enumerate(row_data):
        cell = table.cell(i, j)
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True

# ===== 保存 =====
desktop = os.path.join(os.environ['USERPROFILE'], 'Desktop')
output_path = os.path.join(desktop, 'AI-VISION_PRO_答辩问答集.docx')
doc.save(output_path)
print(f'文档已保存到: {output_path}')
