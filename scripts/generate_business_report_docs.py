from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DATE = "2026年5月14日"


DOCS = {
    "商业汇报文稿": r"""# AI 视频识别信号平台商业汇报文稿

汇报主题：面向园区、仓库、码头的 AI 视频识别与告警闭环平台

## 一、项目背景

传统视频监控主要依赖人工值守，存在三个明显问题：第一，监控画面多、人员注意力有限，异常事件容易漏看；第二，视频事后追溯成本高，告警、回放、处置记录割裂；第三，普通摄像头只能“看见”，不能形成可查询、可追踪、可闭环的数据资产。

本项目围绕园区围栏、仓库、码头等典型安防场景，建设了一套 AI 视频识别信号平台。系统通过摄像头或局域网手机视频流接入实时画面，使用本地视觉模型识别人员目标，再结合规则引擎判断围栏翻越、人员滞留等风险事件，最终形成告警、回放、视频理解、处置记录和审计日志。

## 二、建设目标

本项目的目标不是单纯做一个检测 Demo，而是把“视频识别、规则判断、告警处置、回放追溯、系统运维”串成完整业务闭环。

核心目标包括：

1. 实时发现：对围栏翻越、仓库/码头滞留等异常行为进行实时检测。
2. 减少漏看：将人工盯屏转为系统主动告警，提高值守效率。
3. 可追溯：告警关联摄像头、规则、时间、回放和视频理解结果。
4. 可处置：支持确认、处理中、完成、误报标记、备注和指派。
5. 可运营：具备登录权限、设备管理、系统设置、日志审计、备份清理和部署脚本。

## 三、系统能力

当前系统已经具备以下能力：

| 模块 | 能力说明 |
| --- | --- |
| 实时监控矩阵 | 多摄像头统一展示，视频断流和等待状态有明显提示 |
| 设备接入管理 | 支持新增、编辑、删除、分组、启停、流地址配置和状态检测 |
| 手机摄像头接入 | 同一局域网下可把手机作为实时摄像头接入平台 |
| 告警事件大屏 | 支持筛选、搜索、分页、导出、处理状态管理 |
| 告警闭环 | 支持确认、处理中、完成、误报、备注和指派 |
| 区域规则配置 | 支持围栏线、轨迹通道、滞留区绘制和阈值配置 |
| 回放中心 | 支持按告警定位录像、播放、裁剪片段、检测框叠加 |
| 视频理解 | 支持对回放片段进行语义分析，辅助说明现场发生了什么 |
| Agent 查询 | 支持用自然语言查询告警、回放和运行状态 |
| 权限体系 | 支持超级管理员、管理员、值班人员、只读用户 |
| 运维能力 | 支持健康检查、日志轮转、备份、清理、Docker、Windows 守护启动 |

## 四、典型业务流程

业务流程可以概括为七步：

1. 管理员登录系统，进入设备接入管理。
2. 新增或编辑摄像头，配置本机摄像头、本地视频、RTSP 地址或手机 MJPEG 地址。
3. 系统读取视频流，实时检测人员目标。
4. 规则引擎根据围栏线、滞留区和阈值判断异常事件。
5. 告警进入事件大屏，值班人员确认并处理。
6. 处理人员打开回放中心，查看录像、短片段和视频理解结果。
7. 处理完成后写入状态、备注、指派人和操作审计。

这个流程形成了从“看见异常”到“处理异常”再到“复盘异常”的闭环。

## 五、商业化改造亮点

相比最初演示版本，当前版本已经完成多项商业化增强：

1. 正式登录体系：不再依赖 Debug 入口，Debug 接口默认关闭。
2. 角色权限体系：不同角色拥有不同操作权限。
3. 设备管理页面：摄像头接入从调试页迁移到正式页面。
4. 告警闭环：告警不是只展示，而是可以被确认、指派、完成和误报标记。
5. 操作审计：关键操作都有记录，便于追责和验收。
6. 前端体验：统一按钮、表格、状态标签、弹窗、加载态、空状态、错误态。
7. 数据导出：告警大屏支持筛选、搜索、分页和导出。
8. 工程稳定性：日志轮转、断流重连、进程守护、备份清理、Docker 和 Nginx 配置均已补齐。

## 六、技术架构

系统采用轻量、可交付的架构：

| 层级 | 技术/模块 |
| --- | --- |
| 前端展示层 | HTML/CSS/JavaScript 静态页面 |
| 后端接口层 | FastAPI + Uvicorn |
| 视觉检测层 | YOLO 本地模型 |
| 规则判断层 | 自研 rules_engine，支持围栏、滞留和场景信号 |
| 数据存储层 | SQLite 运行数据库 |
| 回放分析层 | replay_service + ffmpeg + 视频理解接口 |
| 运维部署层 | PowerShell 守护脚本、Docker Compose、Nginx |

这种架构适合课程项目、实验室部署、园区轻量试点，也便于后续迁移到服务器和容器化环境。

## 七、测试与结果

当前版本已经完成后端单元测试和黑盒功能测试：

| 测试类型 | 结果 |
| --- | --- |
| 后端单元测试 | 44 项通过，0 失败 |
| 黑盒接口测试 | 16 项通过，0 失败 |
| 产品完整性测试 | 7 项通过，0 失败 |
| 前端脚本语法检查 | index、module、replay 页面全部通过 |
| 后端健康检查 | /health 返回 200 |
| 前端访问检查 | index.html 返回 200 |

性能口径：历史黑盒测试平均 HTTP 响应约 31.15 ms，最大 104 ms，满足本机演示和小规模局域网试运行要求。

准确率口径：规则引擎层面的围栏、滞留、方向、确认帧等自动化测试全部通过。视觉模型真实 Precision/Recall 需要基于后续人工标注测试集统计，本项目已预留误报标记和后续优化闭环。

## 八、部署与运维

项目支持三类运行方式：

1. 开发演示：start_all_dev.bat。
2. Windows 守护运行：scripts/run_supervisor.ps1，异常自动重启。
3. Docker 部署：docker compose up -d --build。

同时提供：

1. 日志轮转，避免日志无限增大。
2. /ops/health 运维健康检查。
3. /ops/backup 数据备份。
4. /ops/cleanup 过期文件清理。
5. deploy/nginx.conf 反向代理配置。
6. .env.example 完整环境变量模板。

## 九、项目价值

本项目的价值主要体现在三个方面：

第一，提升安防效率。系统把人工盯屏转为自动识别和主动告警，降低值守压力。

第二，提升处置闭环能力。告警、回放、备注、指派、误报标记和审计日志形成完整链条，方便复盘和管理。

第三，具备低成本试点能力。系统支持普通摄像头、手机摄像头和网络流接入，可以快速在园区、仓库、码头等小场景落地试运行。

## 十、后续规划

后续可以继续完善：

1. 进行 24 小时和 72 小时稳定性测试。
2. 建立真实标注测试集，统计 Precision、Recall、误报率和漏报率。
3. 引入 WebRTC/HLS，优化多路视频并发。
4. 增加短信、邮件、企业微信、钉钉通知。
5. 完成 HTTPS、域名和正式服务器部署。
6. 将误报反馈数据用于模型和阈值优化。

## 十一、汇报结论

当前项目已经从单纯演示系统升级为具备商业项目形态的 AI 视频识别平台。它不仅能展示实时检测效果，还具备正式登录、设备管理、告警闭环、回放追溯、操作审计、运维部署和交付文档，能够支撑课程验收、项目汇报和小规模试点演示。
""",
    "AI_PPT生成文字稿": r"""# AI PPT 生成文字稿：AI 视频识别信号平台商业汇报

## PPT 风格要求

请生成一份偏商业汇报风格的 PPT，适合课程项目答辩和项目验收。整体风格应专业、科技、克制，颜色以深色安防大屏风格为主，辅以青蓝色高亮。不要做花哨动画，重点突出系统完整性、业务闭环和商业化改造。

## 第 1 页：封面

标题：AI 视频识别信号平台

副标题：面向园区、仓库、码头的实时检测、告警闭环与回放追溯系统

汇报人：项目组

日期：2026年5月

画面建议：深色背景，加入监控大屏、摄像头、告警信号、数据流等视觉元素。

## 第 2 页：项目背景

标题：为什么需要 AI 视频识别平台

内容：

传统视频监控依赖人工值守，存在漏看、慢处理、难追溯的问题。园区围栏、仓库、码头等场景中，人员翻越、长时间滞留等事件具有突发性和安全风险。项目目标是让摄像头从“只录像”升级为“能识别、会告警、可追溯、能闭环”的智能安防入口。

要点：

1. 监控画面多，人工注意力有限。
2. 异常事件发现不及时。
3. 告警、回放、处理记录割裂。
4. 缺少可审计、可复盘的数据闭环。

## 第 3 页：项目目标

标题：从检测 Demo 到完整业务闭环

内容：

本项目不只做目标检测展示，而是把视频接入、AI 检测、规则判断、告警处置、回放分析、审计运维串成完整流程。

四个关键词：

1. 实时发现：自动识别人员和异常行为。
2. 快速处置：告警进入大屏并支持指派处理。
3. 证据追溯：关联回放、截图、视频片段和分析结果。
4. 稳定交付：具备登录权限、备份清理、部署脚本和文档。

## 第 4 页：系统总体架构

标题：系统架构

内容：

前端展示层：实时监控矩阵、告警大屏、设备管理、回放中心、系统设置。

后端接口层：FastAPI 提供认证、设备、告警、回放、日志和运维接口。

视觉识别层：本地 YOLO 模型识别人员目标。

规则引擎层：围栏翻越、区域滞留、方向和确认帧判断。

数据存储层：SQLite 保存用户、会话、告警、视频分析、设置和审计日志。

运维部署层：日志轮转、备份清理、Docker、Nginx、Windows 守护脚本。

画面建议：用分层架构图展示。

## 第 5 页：核心业务场景

标题：覆盖园区、仓库、码头三类高频安防场景

内容：

场景一：园区内部围栏检测，识别人员翻越围栏、围栏区域异常滞留。

场景二：仓库区域滞留检测，识别人员长时间停留、异常聚集。

场景三：码头区域滞留检测，识别人员在作业区域异常停留。

补充：系统支持同一局域网下把手机作为实时摄像头，适合快速演示和低成本试点。

## 第 6 页：设备接入管理

标题：正式设备管理替代 Debug 接入

内容：

设备接入管理支持摄像头新增、编辑、删除、分组、启停、流地址配置和状态检测。支持本机摄像头、本地视频、手机 MJPEG、RTSP 网络摄像头等多种来源。

操作流程：

1. 管理员登录。
2. 进入设备接入管理。
3. 新增摄像头 ID、名称、分组、视频流地址。
4. 选择启用状态和关联场景。
5. 保存后系统热重载配置。
6. 回到监控矩阵查看实时画面。

## 第 7 页：实时识别与规则判断

标题：AI 检测 + 规则引擎

内容：

系统先通过 YOLO 检测人员，再由规则引擎根据区域配置和时间阈值判断事件。围栏规则关注穿越边界，滞留规则关注人员在区域内停留时间。规则引擎支持确认帧机制，降低瞬时误报。

亮点：

1. 检测框实时叠加。
2. 支持围栏线和滞留区绘制。
3. 支持滞留阈值调整。
4. 支持断流后自动重连。

## 第 8 页：告警闭环

标题：告警不是展示，而是可处理的业务事件

内容：

告警大屏支持按严重级别、摄像头、视频理解状态、处理状态筛选，也支持搜索、分页和导出。每条告警都可以确认、处理中、完成或标记误报，并记录备注和指派人员。

告警状态：

1. 新告警。
2. 已确认。
3. 处理中。
4. 已完成。
5. 误报。

价值：让告警从“看到了”变成“处理了、记录了、可追责”。

## 第 9 页：回放追溯与视频理解

标题：告警自动关联证据链

内容：

回放中心根据摄像头、时间、场景和规则自动定位录像文件，支持播放、下载、生成短片段、显示事件偏移点和检测框叠加。系统还支持调用视频理解能力，对回放片段生成文字说明，辅助值班人员快速理解现场情况。

价值：

1. 减少人工找录像时间。
2. 告警与视频证据稳定关联。
3. 支持事后复盘和报告撰写。

## 第 10 页：权限、安全与审计

标题：从演示入口升级为正式管理体系

内容：

系统支持超级管理员、普通管理员、值班人员、只读用户。Debug 接口默认关闭，登录失败会临时锁定账号，退出登录后 Token 失效。关键操作写入审计日志，包括规则修改、设备删除、用户管理、告警处理、备份清理等。

安全能力：

1. 密码哈希存储。
2. Token 会话管理。
3. API 权限校验。
4. Debug 默认关闭。
5. 操作审计。
6. 回放下载路径限制。

## 第 11 页：工程稳定性与部署

标题：具备可运行、可守护、可交付能力

内容：

项目补齐了商业交付常见基础能力：日志轮转、视频断流重连、健康检查、备份、清理、Docker 部署、Nginx 反向代理、Windows 守护脚本和计划任务。

交付文件：

1. Dockerfile。
2. docker-compose.yml。
3. deploy/nginx.conf。
4. scripts/run_supervisor.ps1。
5. scripts/ops_maintenance.py。
6. .env.example。

## 第 12 页：测试结果

标题：功能、性能、准确率测试口径

内容：

后端单元测试 44 项全部通过，黑盒接口测试 16 项全部通过，产品完整性测试 7 项全部通过。历史黑盒测试平均 HTTP 响应约 31.15 ms，最大 104 ms，满足本机演示和小规模局域网试运行。

准确率说明：

规则引擎层面的围栏、滞留、方向、确认帧测试全部通过。视觉模型真实 Precision/Recall 需要后续基于人工标注集统计，系统已提供误报标记机制用于持续优化。

## 第 13 页：项目价值

标题：项目价值总结

内容：

第一，提升安防效率：系统主动识别异常并告警，降低人工盯屏压力。

第二，提升处置闭环：告警、回放、备注、指派、误报、完成状态形成闭环。

第三，降低试点成本：支持普通电脑、普通摄像头和手机局域网摄像头快速接入。

第四，具备交付基础：系统已经具备权限、安全、运维、部署和文档能力。

## 第 14 页：后续规划

标题：下一阶段优化方向

内容：

1. 进行 24 小时和 72 小时稳定性测试。
2. 建立真实场景标注测试集，统计 Precision、Recall、误报率和漏报率。
3. 引入 WebRTC/HLS，提高多路视频并发能力。
4. 增加短信、邮件、企业微信、钉钉通知。
5. 完成 HTTPS、域名和服务器部署。
6. 根据误报反馈优化模型和阈值。

## 第 15 页：结论

标题：项目结论

内容：

当前 AI 视频识别信号平台已经从单一检测 Demo 升级为完整业务系统。它具备实时监控、设备管理、告警闭环、回放追溯、视频理解、权限审计、运维部署和文档交付能力，能够支撑课程验收、项目汇报和小规模试点演示。

结束语：

让摄像头从被动录像设备，升级为主动发现风险、辅助处置决策的智能安防节点。
""",
}


def shade(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def cell_text(cell, text: str, *, bold: bool = False, white: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(8.5)
    run.bold = bold
    if white:
        run.font.color.rgb = RGBColor(255, 255, 255)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_doc(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    for name, size, color in [("Heading 1", 15, "2E74B5"), ("Heading 2", 12.5, "2E74B5")]:
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
    header = section.header.paragraphs[0]
    header.text = f"AI 视频识别信号平台 | {title}"
    header.runs[0].font.size = Pt(9)
    footer = section.footer.paragraphs[0]
    footer.text = f"{DATE}  D:\\Project"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.runs[0].font.size = Pt(9)


def add_p(doc: Document, text: str, style: str | None = None) -> None:
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10 if style else 10.5)


def markdown_to_docx(markdown: str, title: str, path: Path) -> None:
    doc = Document()
    style_doc(doc, title)
    lines = markdown.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line[2:].strip())
            r.font.name = "Microsoft YaHei"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            r.font.size = Pt(21)
            r.bold = True
            r.font.color.rgb = RGBColor.from_string("0B2545")
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), 1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), 2)
            i += 1
            continue
        if line.startswith("```"):
            i += 1
            code = []
            while i < len(lines) and not lines[i].startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            r = p.add_run("\n".join(code))
            r.font.name = "Consolas"
            r.font.size = Pt(8.5)
            continue
        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                core = lines[i].replace("|", "").replace("-", "").replace(":", "").replace(" ", "")
                if core:
                    table_lines.append(lines[i])
                i += 1
            rows = [[c.strip() for c in row.strip("|").split("|")] for row in table_lines]
            if rows:
                cols = max(len(row) for row in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for ri, row in enumerate(rows):
                    for ci in range(cols):
                        cell = table.cell(ri, ci)
                        cell_text(cell, row[ci] if ci < len(row) else "", bold=ri == 0, white=ri == 0)
                        if ri == 0:
                            shade(cell, "2E74B5")
                doc.add_paragraph()
            continue
        if line.startswith("- "):
            add_p(doc, line[2:].strip(), "List Bullet")
        elif len(line) > 3 and line[0].isdigit() and line[1:3] == ". ":
            add_p(doc, line[3:].strip(), "List Number")
        else:
            add_p(doc, line.replace("`", "").replace("**", "").strip())
        i += 1
    doc.save(path)


def main() -> None:
    for title, markdown in DOCS.items():
        md_path = ROOT / f"{title}.md"
        docx_path = ROOT / f"{title}.docx"
        md_path.write_text(markdown.strip() + "\n", encoding="utf-8")
        markdown_to_docx(markdown, title, docx_path)
        print(f"generated {md_path.name} and {docx_path.name}")


if __name__ == "__main__":
    main()
