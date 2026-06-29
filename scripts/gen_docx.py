#!/usr/bin/env python3
"""Generate the project documentation as a .docx file."""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── Global style defaults ──────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Microsoft YaHei"
style.font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.35
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Microsoft YaHei"
    hs.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    hs.font.color.rgb = RGBColor(0x1A, 0x47, 0x7A)
    if level == 1:
        hs.font.size = Pt(18)
    elif level == 2:
        hs.font.size = Pt(14)
    else:
        hs.font.size = Pt(12)

# ── Helper functions ────────────────────────────────────────────────────
def add_para(text, bold=False, size=None, color=None, align=None, space_before=0, space_after=4):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if align:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10.5)
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.8)
    return p


def add_code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    shading_elm = run.element.get_or_add_rPr()
    from docx.oxml import OxmlElement
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    shading_elm.append(shd)
    return p


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = "Microsoft YaHei"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        from docx.oxml import OxmlElement
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "2B579A")
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)
        cell._tc.get_or_add_tcPr().append(shd)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = "Microsoft YaHei"
            run.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(10)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
            if r_idx % 2 == 0:
                from docx.oxml import OxmlElement
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "EDF2F9")
                cell._tc.get_or_add_tcPr().append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def add_separator():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    from docx.oxml import OxmlElement
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    p._p.get_or_add_pPr().append(pBdr)


# ═══════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

add_para("AI 视频识别信号平台", bold=True, size=28, color=(0x1A, 0x47, 0x7A),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
add_para("项目讲解书", bold=True, size=20, color=(0x33, 0x66, 0x99),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

add_separator()

add_para("版本：v0.3.0", size=11, color=(0x66, 0x66, 0x66),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=24)
add_para("生成日期：2026年5月25日", size=11, color=(0x66, 0x66, 0x66),
         align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("目录", level=1)
toc_items = [
    "一、项目定位",
    "二、业务范围",
    "三、技术栈一览",
    "四、整体架构",
    "五、核心模块详解",
    "    5.1 YOLO 目标检测模块",
    "    5.2 IoU 跟踪器",
    "    5.3 规则引擎",
    "    5.4 实时视频流模块",
    "    5.5 Agent 智能问答模块",
    "    5.6 存储层",
    "    5.7 配置驱动机制",
    "六、完整数据流",
    "七、关键 API 接口",
    "八、关键算法总结",
    "九、定时维护任务",
    "十、启动方式与部署",
    "十一、自动化测试",
]
for item in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.name = "Microsoft YaHei"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(11)
    if not item.startswith("    "):
        run.bold = True
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0) if not item.startswith("    ") else Cm(1.5)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════
# 一、项目定位
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("一、项目定位", level=1)
add_para(
    "本项目是一个园区安防智能监控系统。核心功能是：摄像头拍到画面后，AI 自动识别画面中的人员、车辆、动物，"
    "并根据预设规则判断是否存在「翻越围栏」或「区域滞留」等违规行为。一旦触发违规，系统立刻生成标准信号"
    "（JSON 格式），供前端大屏、Agent 智能问答和第三方系统消费。"
)
add_para("示例输出信号：", bold=True)
add_code_block('{"是否滞留": 1, "滞留人数": 3}')

# ═══════════════════════════════════════════════════════════════════════
# 二、业务范围
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("二、业务范围", level=1)
add_para("系统当前覆盖两个业务场景：")
add_table(
    ["场景", "场景 ID", "摄像头", "包含的检测规则"],
    [
        ["园区内部围栏检测", "campus_fence", "cam_fence",
         "翻越围栏检测 + 围栏区域滞留检测"],
        ["仓库 + 码头检测", "warehouse_dock", "cam_dock、cam_warehouse",
         "码头区域滞留检测（人）+ 仓库滞留检测"],
    ],
    col_widths=[1.5, 1.2, 1.5, 2.5],
)

# ═══════════════════════════════════════════════════════════════════════
# 三、技术栈一览
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("三、技术栈一览", level=1)
add_table(
    ["层面", "技术选型", "说明"],
    [
        ["Web 框架", "FastAPI（Python）", "高性能异步 Web 框架"],
        ["目标检测", "YOLOv8（ultralytics）", "实时目标检测模型"],
        ["视频处理", "OpenCV", "视频编解码与图像处理"],
        ["跟踪算法", "IoU Greedy Tracker", "自研帧间目标跟踪器"],
        ["规则引擎", "纯 Python", "射线法、有向距离法、几何算法"],
        ["数据库", "SQLite", "轻量级，零部署成本"],
        ["LLM 对话（可选）", "DeepSeek API", "Agent 智能问答"],
        ["视频理解（可选）", "MiMo 视频模型（小米）", "事件视频语义分析"],
        ["前端", "静态 HTML + Nginx", "纯前端，无需构建"],
        ["部署", "Docker Compose", "容器化一键部署"],
    ],
    col_widths=[1.5, 2.0, 2.8],
)

# ═══════════════════════════════════════════════════════════════════════
# 四、整体架构
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("四、整体架构", level=1)

doc.add_heading("4.1 数据流全景图", level=2)
add_code_block(
    "摄像头视频流\n"
    "    |\n"
    "    v\n"
    "+-----------------------------+\n"
    "|  OpenCV 视频帧采集           |  <-- 从摄像头/视频文件读取画面\n"
    "+-----------------------------+\n"
    "    |\n"
    "    v\n"
    "+-----------------------------+\n"
    "|  YOLO 目标检测               |  <-- 识别画面中有什么（人/车/动物）\n"
    "|  （或 MiMo 视频理解）        |      输出：位置 + 类别 + 置信度\n"
    "+-----------------------------+\n"
    "    |\n"
    "    v\n"
    "+-----------------------------+\n"
    "|  IoU 跟踪器                 |  <-- 认出\"第1帧的张三\"和\"第100帧的张三\"\n"
    "|                              |      是同一个人，分配唯一 track_id\n"
    "+-----------------------------+\n"
    "    |\n"
    "    v\n"
    "+-----------------------------+\n"
    "|  规则引擎                    |  <-- 判断是否违规（越界？滞留超时？）\n"
    "|  boundary + dwell + cooldown |      生成告警事件\n"
    "+-----------------------------+\n"
    "    |\n"
    "    v\n"
    "+-----------------------------+\n"
    "|  信号输出 + SQLite 持久化    |  <-- 生成标准信号 + 写入数据库\n"
    "+-----------------------------+\n"
    "    |\n"
    "    +--> 前端大屏（实时视频 + 告警弹窗）\n"
    "    +--> Agent 智能问答（自然语言对话）\n"
    "    +--> 第三方系统（标准 JSON 信号接口）"
)

doc.add_heading("4.2 目录结构", level=2)
add_code_block(
    "D:\\Project/\n"
    "+-- backend/                     # 后端 Python 代码\n"
    "|   +-- app/\n"
    "|       +-- main.py              # FastAPI 应用入口\n"
    "|       +-- core/config.py       # 配置加载（YAML 读取）\n"
    "|       +-- api/routes.py        # 所有 HTTP 接口路由\n"
    "|       +-- schemas/\n"
    "|       |   +-- detection.py     # 检测数据模型\n"
    "|       |   +-- vision.py        # 视觉分析数据模型\n"
    "|       +-- services/\n"
    "|           +-- yolo_service.py         # YOLO 模型加载与检测\n"
    "|           +-- tracking_service.py     # IoU 跟踪器\n"
    "|           +-- rules_engine.py         # 规则引擎（核心）\n"
    "|           +-- stream_service.py       # 实时视频流处理\n"
    "|           +-- storage_service.py      # SQLite 存储层\n"
    "|           +-- llm_client.py           # LLM 客户端\n"
    "|           +-- agent_orchestrator.py   # Agent 编排器\n"
    "|           +-- agent_tools.py          # Agent 工具集\n"
    "|           +-- agent_policy.py         # Agent 安全策略\n"
    "+-- frontend/static/             # 前端静态页面\n"
    "+-- config/\n"
    "|   +-- rules.yaml               # 规则配置\n"
    "|   +-- tracker.yaml             # 跟踪器参数\n"
    "|   +-- vision_backend.yaml      # 检测后端切换配置\n"
    "+-- models/                      # YOLO 预训练权重\n"
    "+-- data/runtime/                # SQLite 数据库\n"
    "+-- scripts/                     # 辅助脚本\n"
    "+-- requirements.txt             # Python 依赖\n"
    "+-- Dockerfile / docker-compose.yml"
)

# ═══════════════════════════════════════════════════════════════════════
# 五、核心模块详解
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("五、核心模块详解", level=1)

# ── 5.1 YOLO ───────────────────────────────────────────────────────────
doc.add_heading("5.1 YOLO 目标检测模块（yolo_service.py）", level=2)

doc.add_heading("功能", level=3)
add_para(
    "把视频画面翻译成结构化数据。每一帧图像经过 YOLO 模型后，输出每个物体的位置（矩形框 BBox）、"
    "类别（人/车/动物）、置信度（0.99 = 99% 确定）。"
)

doc.add_heading("数据结构", level=3)
add_code_block(
    "Detection(\n"
    '    camera_id="cam_fence",         # 来源摄像头\n'
    '    category="person",             # 归一化后的类别\n'
    '    display_category="pedestrian", # YOLO 原始类别名\n'
    '    confidence=0.92,               # 置信度 92%\n'
    '    bbox=BBox(x1=0.3, y1=0.1, x2=0.5, y2=0.9),\n'
    '    track_id=7                     # 跟踪编号\n'
    ")"
)

doc.add_heading("类别归一化算法", level=3)
add_para(
    "YOLO 输出的类别名非常多（pedestrian、man、woman 都算人），"
    "系统通过 CATEGORY_GROUPS 映射表统一归为3大类："
)
add_table(
    ["标准类别", "包含的原始类别名"],
    [
        ["person（人）", "person, human, pedestrian, man, woman, worker, head, person-like"],
        ["vehicle（车）", "bicycle, car, motorcycle, bus, truck, auto, vehicle, van"],
        ["animal（动物）", "bird, cat, dog, horse, sheep, cow, animal, pet"],
    ],
    col_widths=[1.5, 5.0],
)

doc.add_heading("模型自动选择机制", level=3)
add_para("系统按优先级自动搜索可用的 YOLO 权重文件：")
add_bullet("优先使用 .env 中 YOLO_WEIGHTS_PATH 指定的路径")
add_bullet("依次在 models/ 目录查找：yolo26s.pt -> yolov8s.pt -> yolov8n.pt -> best.pt")
add_bullet("还会检查模型与当前 ultralytics 版本的兼容性")

# ── 5.2 IoU Tracker ────────────────────────────────────────────────────
doc.add_heading("5.2 IoU 跟踪器（tracking_service.py）", level=2)

doc.add_heading("功能", level=3)
add_para(
    "让系统认出第1帧里的人 A 和第100帧里的人 A 是同一个人，"
    "给每个人分配唯一的 track_id，并在整个生命周期内保持不变。"
)

doc.add_heading("核心算法：IoU Greedy Tracker", level=3)
add_para("工作流程：")
add_code_block(
    "第1帧：检测到3个人 --> 分配 track_id: 1, 2, 3\n"
    "第2帧：检测到3个人 --> 计算每对框的 IoU，匹配最大 IoU 的配对\n"
    "第3帧：检测到2个人（有1人被柱子挡住）--> 使用\"记忆\"等待其重新出现\n"
    "...\n"
    "第N帧：被挡住的人重新出现 --> 如果 IoU 匹配成功，保持原 ID"
)

doc.add_heading("IoU（交并比）计算", level=3)
add_code_block(
    "IoU = 两个矩形框的重叠面积 / 两个矩形框的合并面积\n\n"
    "     +--------+\n"
    "     |  框A   |\n"
    "     |    +---+---+\n"
    "     |    |重叠|   |\n"
    "     +----+---+   |\n"
    "          |  框B  |\n"
    "          +-------+\n\n"
    "IoU = 重叠面积 / (框A面积 + 框B面积 - 重叠面积)\n"
    "取值范围：0（完全不重叠） ~ 1（完全重合）"
)

doc.add_heading("关键参数（config/tracker.yaml）", level=3)
add_table(
    ["参数", "默认值", "含义"],
    [
        ["match_thresh", "0.15", "IoU >= 0.15 就认为是同一目标（门槛很低，非常粘人）"],
        ["track_buffer", "60", "允许丢失最多60帧（约2.4秒 @25fps）"],
        ["frame_rate", "25", "帧率，用于计算时间"],
        ["max_age_seconds", "5.0", "目标消失超过5秒后，从记忆中清除"],
    ],
    col_widths=[1.5, 1.0, 4.0],
)

doc.add_heading("贪心匹配算法详细步骤", level=3)
steps = [
    "Step 1: 将当前帧的检测结果按 (camera_id, category) 分组",
    "Step 2: 每组内，计算所有新检测框与所有旧轨迹框的 IoU",
    "Step 3: 将所有 IoU >= 0.15 的配对按 IoU 从大到小排序",
    "Step 4: 依次处理排序后的配对：如果该检测框和该轨迹都没有被占用则匹配成功，否则跳过",
    "Step 5: 没有匹配上的检测框分配新的 track_id",
    "Step 6: 清理超过5秒未出现的旧轨迹",
]
for s in steps:
    add_bullet(s)

# ── 5.3 规则引擎 ───────────────────────────────────────────────────────
doc.add_heading("5.3 规则引擎（rules_engine.py）", level=2)

add_para(
    "规则引擎是整个系统的「大脑」，根据 config/rules.yaml 中的配置判断哪些行为算「违规」。"
)

doc.add_heading("类型一：Boundary（越界检测）", level=3)
add_para("场景：围栏翻越检测。")
add_para("原理：")
add_bullet("在视频画面中定义一条"警戒线"（两个点构成一条线段）")
add_bullet("每一帧，计算人的中心点相对于这条线的有向距离")
add_bullet("如果上一帧中心点在线左侧（负值），这一帧到了右侧（正值） -> 判定越界")

add_para("越线判断核心算法：有向距离法", bold=True)
add_code_block(
    "# 有向距离：通过叉积计算\n"
    "def signed_distance_to_line(point, line_start, line_end):\n"
    "    px, py = point\n"
    "    x1, y1 = line_start\n"
    "    x2, y2 = line_end\n"
    "    return (x2 - x1) * (py - y1) - (y2 - y1) * (px - x1)\n\n"
    "# 判断是否越线：前后两帧的有向距离符号相反\n"
    "def segment_crossed_line(prev_point, cur_point, line_start, line_end):\n"
    "    prev_side = signed_distance_to_line(prev_point, line_start, line_end)\n"
    "    cur_side = signed_distance_to_line(cur_point, line_start, line_end)\n"
    "    return (prev_side * cur_side) < 0  # 一正一负 -> 越线！"
)

add_para("方向控制：", bold=True)
add_bullet("any：任何方向越线都告警")
add_bullet("neg_to_pos：只检测从负侧到正侧的越线")
add_bullet("pos_to_neg：只检测从正侧到负侧的越线")

add_para("防抖机制：", bold=True)
add_para(
    "需要连续 confirm_frames 帧都确认越线，才正式触发告警，避免单帧误判。"
)

doc.add_heading("类型二：Dwell（滞留检测）", level=3)
add_para("场景：仓库、码头、围栏区域有人待太久。")
add_para("原理：")
add_bullet("在画面中定义一个多边形区域（dwell_zone）")
add_bullet("每帧检测人的中心点是否在区域内（使用射线法判断）")
add_bullet("如果在区域内，开始计时")
add_bullet("如果连续停留超过 threshold_seconds（默认5秒）-> 触发告警")

add_para("核心算法：射线法（Ray Casting）判断点在多边形内", bold=True)
add_code_block(
    "def point_in_polygon(x, y, polygon):\n"
    '    """\n'
    "    从点 (x, y) 向右发射一条水平射线，\n"
    "    数一数这条射线和多边形边界的交点数：\n"
    "    - 奇数 -> 点在多边形内部\n"
    "    - 偶数 -> 点在多边形外部\n"
    '    """\n'
    "    inside = False\n"
    "    j = len(polygon) - 1\n"
    "    for i in range(len(polygon)):\n"
    "        xi, yi = polygon[i]\n"
    "        xj, yj = polygon[j]\n"
    "        intersect = ((yi > y) != (yj > y)) and (\n"
    "            x < (xj - xi) * (y - yi) / ((yj - yi) + 1e-9) + xi\n"
    "        )\n"
    "        if intersect:\n"
    "            inside = not inside\n"
    "        j = i\n"
    "    return inside"
)

add_para("滞留检测完整流程：", bold=True)
add_code_block(
    "帧1: 人在区域内 --> 记录首次出现时间（dwell_first_seen）\n"
    "帧2: 人在区域内 --> 继续等待\n"
    "帧3: 人在区域内 --> 停留时间 < 5秒 --> 继续\n"
    "...\n"
    "帧N: 停留时间 >= 5秒 --> 开始确认计数\n"
    "帧N+1: 仍在区域内 --> 确认计数 +1\n"
    "帧N+confirm_frames: 确认计数达到阈值 --> 触发告警！\n"
    "帧M: 人离开区域 --> 重置计时器"
)

doc.add_heading("辅助几何算法", level=3)

add_para("线段相交检测（segments_intersect）：", bold=True)
add_para("用于判断检测框（矩形）是否与警戒线相交。使用方向测试（Orientation Test）判断两线段是否相交。")

add_para("检测框与线相交判断（bbox_intersects_line）：", bold=True)
add_para("将矩形框拆为4条边，依次检测每条边是否与警戒线相交。")

add_para("检测框膨胀（expand_bbox）：", bold=True)
add_para(
    "为检测框增加容差范围（line_touch_tolerance），"
    "允许人"接近"警戒线时就触发检测，而不必完全跨过。"
)

doc.add_heading("信号锁存与冷却机制", level=3)

add_para("信号锁存（Signal Latch）：", bold=True)
add_para(
    "告警触发后不会立刻消失。系统有一个"信号保持时间"（signal_hold_seconds），"
    "在此期间即使人已离开，信号仍保持为 1（活跃状态），避免信号频繁闪烁。"
)

add_para("冷却机制（Cooldown）：", bold=True)
add_para(
    "同一个目标触发同一条规则后，必须等待 cooldown_seconds（如5秒）"
    "才能再次触发，防止短时间内重复告警。"
)

# ── 5.4 Stream Service ─────────────────────────────────────────────────
doc.add_heading("5.4 实时视频流模块（stream_service.py）", level=2)

doc.add_heading("功能", level=3)
add_para(
    "把 AI 检测结果实时叠加到视频画面上，通过 MJPEG 流推送到前端浏览器。"
)

doc.add_heading("处理流程", level=3)
flow_steps = [
    "OpenCV 打开摄像头（支持 USB 摄像头、RTSP 流、HTTP 流、视频文件）",
    "逐帧读取画面",
    "每隔 N 帧（默认4帧）执行一次 AI 检测",
    "将检测结果叠加到画面上（检测框、告警红框、区域边界线、状态信息条）",
    "将画面编码为 JPEG 图片",
    "通过 HTTP multipart 流推送给前端",
]
for s in flow_steps:
    add_bullet(s)

doc.add_heading("性能优化策略", level=3)
add_table(
    ["策略", "说明"],
    [
        ["分辨率缩放", "检测前将图片缩到最大960像素，大幅加快推理速度"],
        ["间隔检测", "每4帧才跑一次AI，中间帧复用上一次的检测结果"],
        ["MJPEG 流", "用 HTTP multipart 方式推送，前端用 <img> 标签即可播放"],
        ["JPEG 压缩质量", "默认76，在画质和带宽间取平衡"],
    ],
    col_widths=[1.5, 5.0],
)

doc.add_heading("画面渲染细节", level=3)
add_table(
    ["元素", "颜色/样式"],
    [
        ["人员框", "蓝色 (104, 226, 255)"],
        ["车辆框", "橙色 (255, 176, 76)"],
        ["动物框", "紫色 (188, 132, 255)"],
        ["告警状态框", "红色粗框 (76, 122, 255)，带\"告警\"前缀"],
        ["警戒线", "半透明蓝色线条"],
        ["滞留区域", "半透明绿色多边形"],
        ["状态信息条", "左上角半透明卡片，显示摄像头ID、目标数、筛查模式"],
    ],
    col_widths=[1.5, 5.0],
)

# ── 5.5 Agent ──────────────────────────────────────────────────────────
doc.add_heading("5.5 Agent 智能问答模块", level=2)

doc.add_heading("功能", level=3)
add_para(
    '让安保人员可以用自然语言与系统对话，例如：「现在系统运行情况怎么样？」、'
    '「最近有什么告警？」、「帮我分析一下那段监控视频」。'
)

doc.add_heading("架构：混合式（本地工具 + 可选云端 LLM）", level=3)
add_code_block(
    "用户自然语言提问\n"
    "    |\n"
    "    v\n"
    "+------------------+\n"
    "|  意图识别         |  <-- 第一层：本地关键词匹配\n"
    "|                  |  <-- 第二层：LLM 分类（可选）\n"
    "+------------------+\n"
    "    |\n"
    "    v\n"
    "+------------------+\n"
    "|  工具选择与执行   |  <-- 根据意图选择工具并获取数据\n"
    "+------------------+\n"
    "    |\n"
    "    v\n"
    "+------------------+\n"
    "|  回答生成         |  <-- 本地模板拼接 / LLM 生成自然语言\n"
    "+------------------+"
)

doc.add_heading("意图识别", level=3)
add_table(
    ["关键词类别", "匹配的关键词", "识别出的意图"],
    [
        ["运行状态", "runtime, 状态, 健康, yolo, 视频理解", "runtime"],
        ["告警查询", "alert, 告警, 报警, 事件", "alerts"],
        ["视频回放", "replay, 回放, 录像, 视频, 发生了什么", "replay"],
        ["总览总结", "summary, 总览, 总结", "summary"],
    ],
    col_widths=[1.2, 2.8, 1.2],
)

doc.add_heading("工具集（只读）", level=3)
add_table(
    ["工具名", "功能"],
    [
        ["get_runtime_status", "获取系统运行状态（引擎、跟踪器、检测器状态）"],
        ["get_alert_summary", "获取告警统计摘要（按等级、摄像头、规则分类）"],
        ["get_replay_hint", "获取最新告警的回放定位参数"],
        ["analyze_replay_video", "截取事件短视频并使用 MiMo 进行语义分析"],
    ],
    col_widths=[2.0, 4.5],
)

add_para("安全策略：Agent 只允许调用只读工具，禁止任何写操作，确保系统安全。", bold=True)

doc.add_heading("运行模式", level=3)
add_table(
    ["模式", "说明"],
    [
        ["local_fallback", "LLM 未配置，完全使用本地模板生成回答"],
        ["local_fallback_answer", "LLM 已配置但调用失败，回退到本地回答"],
        ["hybrid_llm", "本地工具获取数据 + LLM 生成自然语言回答"],
    ],
    col_widths=[2.0, 4.5],
)

# ── 5.6 Storage ────────────────────────────────────────────────────────
doc.add_heading("5.6 存储层（storage_service.py）", level=2)

doc.add_heading("功能", level=3)
add_para("使用 SQLite 持久化所有重要数据，支持多表、索引、自动迁移。")

doc.add_heading("数据库表结构", level=3)
add_table(
    ["表名", "用途", "主要字段"],
    [
        ["alerts", "所有告警记录", "timestamp, rule_id, camera_id, severity"],
        ["signal_snapshots", "场景信号历史快照", "scene_id, payload_json"],
        ["users", "用户账号", "username, password_hash, role, status"],
        ["auth_sessions", "登录会话", "token, user_id, expires_at"],
        ["operation_logs", "操作审计日志", "module, action, operator"],
        ["video_analyses", "视频分析结果", "camera_id, summary, analysis_json"],
        ["alert_workflows", "告警工作流", "status, assignee, note"],
        ["ingest_frames", "接入帧统计", "camera_id, detection_count"],
        ["system_settings", "系统配置", "key, value_json"],
    ],
    col_widths=[1.5, 1.8, 3.2],
)

doc.add_heading("密码安全", level=3)
add_bullet("使用 PBKDF2-SHA256 加盐哈希")
add_bullet("迭代次数：120,000 次")
add_bullet("盐值：32字节随机生成（secrets.token_hex(16)）")

doc.add_heading("自动 Schema 迁移", level=3)
add_para("数据库升级时自动检测旧结构并迁移：")
add_bullet("旧版明文密码 -> 自动转为哈希存储")
add_bullet("缺失字段 -> 自动 ALTER TABLE ADD COLUMN")

# ── 5.7 配置驱动 ──────────────────────────────────────────────────────
doc.add_heading("5.7 配置驱动机制（config/rules.yaml）", level=2)

add_para(
    "整个系统的核心行为完全由 YAML 配置文件驱动，不需要修改任何代码即可调整。"
)

doc.add_heading("场景定义", level=3)
add_code_block(
    "scenes:\n"
    "- id: campus_fence\n"
    "  name: 园区内部围栏检测\n"
    "  cameras: [cam_fence]\n"
    "  rule_ids:\n"
    "  - fence_intrusion    # 翻越围栏规则\n"
    "  - fence_dwell        # 围栏滞留规则"
)

doc.add_heading("摄像头定义", level=3)
add_code_block(
    "cameras:\n"
    "- id: cam_fence\n"
    "  name: 园区内部围栏\n"
    "  stream: camera://0           # 视频源（USB摄像头编号0）\n"
    "  rois:                        # 感兴趣区域（用于越界检测）\n"
    "  - id: fence_top\n"
    "    type: boundary_line\n"
    "    line:                      # 警戒线坐标（归一化）\n"
    "    - [0.3844, 0.4237]\n"
    "    - [0.5711, 0.4421]\n"
    "    path_width: 0.08           # 警戒线容差宽度\n"
    "  dwell_zones:                 # 滞留检测区域\n"
    "  - id: fence_inside\n"
    "    polygon:                   # 多边形区域坐标（归一化）\n"
    "    - [0.2467, 0.1281]\n"
    "    - [0.6228, 0.1281]\n"
    "    - [0.6228, 0.9295]\n"
    "    - [0.2467, 0.9295]\n"
    "    threshold_seconds: 5       # 滞留阈值：5秒"
)

doc.add_heading("规则定义", level=3)
add_code_block(
    "rules:\n"
    "- id: fence_intrusion\n"
    "  type: boundary                 # 越界类型\n"
    "  camera_id: cam_fence\n"
    "  roi_id: fence_top\n"
    "  severity: high                 # 严重级别：high / medium / low\n"
    "  category_filter: [person]      # 仅检测人员\n"
    "  crossing_direction: any        # 任何方向越线\n"
    "  confirm_frames: 1              # 确认帧数\n"
    "  cooldown_seconds: 2            # 冷却时间：2秒\n"
    "  signal_key: is_fence_climb     # 输出信号键名\n"
    "  signal_cn: 是否翻越围栏         # 中文信号名"
)

doc.add_heading("检测后端切换（vision_backend.yaml）", level=3)
add_table(
    ["方案", "标签", "说明"],
    [
        ["yolo（方案一）", "YOLO 目标检测", "本地运行，低延迟，适合实时监控"],
        ["video_understanding（方案二）", "视频理解模型", "云端 VLM，语义理解能力强"],
    ],
    col_widths=[2.0, 1.5, 3.0],
)
add_para("可按摄像头灵活切换，例如：围栏和仓库使用视频理解，其他使用 YOLO。")

# ═══════════════════════════════════════════════════════════════════════
# 六、完整数据流
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("六、完整数据流（从摄像头到最终输出）", level=1)

flow = [
    ("Step 1: 摄像头采集", "OpenCV VideoCapture 从摄像头/视频文件读取一帧画面"),
    ("Step 2: 预处理", "如果画面分辨率过大，缩放到最大960像素"),
    ("Step 3: YOLO 检测", "将画面输入 YOLO 模型，输出：[person, 0.92, (x1=0.3, y1=0.1, x2=0.5, y2=0.9)]"),
    ("Step 4: 类别归一化", "\"pedestrian\" -> \"person\"，\"bicycle\" -> \"vehicle\""),
    ("Step 5: IoU 跟踪", "计算新检测框与已有轨迹的 IoU，匹配成功则保持 track_id"),
    ("Step 6: 规则引擎评估", "判断是否越界（有向距离法）+ 判断滞留时间（射线法 + 计时）+ 检查冷却期"),
    ("Step 7: 触发告警", "生成告警事件：{rule_id, camera_id, track_id, message, severity, timestamp}"),
    ("Step 8: 信号输出", "将告警聚合为标准信号：{\"是否翻越围栏\": 1, \"翻越围栏人数\": 1}"),
    ("Step 9: 多渠道输出", "前端大屏 + Agent 聊天 + SQLite 存储 + 第三方 API"),
]
for title, desc in flow:
    add_para(title, bold=True, space_before=6)
    add_para(desc)

# ═══════════════════════════════════════════════════════════════════════
# 七、关键 API 接口
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("七、关键 API 接口", level=1)

doc.add_heading("7.1 健康与配置", level=2)
add_table(
    ["方法", "路径", "功能"],
    [
        ["GET", "/health", "健康检查"],
        ["GET", "/config/rules", "获取规则配置"],
        ["GET", "/config/scenes", "获取场景列表"],
        ["POST", "/config/reload", "热重载规则配置（需调试token）"],
    ],
    col_widths=[0.8, 2.0, 3.7],
)

doc.add_heading("7.2 检测输入与事件", level=2)
add_table(
    ["方法", "路径", "功能"],
    [
        ["POST", "/ingest/detections", "接收外部检测帧（对接外部检测器）"],
        ["GET", "/alerts", "获取当前活跃告警"],
        ["GET", "/alerts/scene/{scene_id}", "获取指定场景的告警"],
        ["GET", "/alerts/history", "查询告警历史"],
    ],
    col_widths=[0.8, 2.2, 3.5],
)

doc.add_heading("7.3 信号输出", level=2)
add_table(
    ["方法", "路径", "功能"],
    [
        ["GET", "/signals/scenes", "获取所有场景信号"],
        ["GET", "/signals/scenes/{scene_id}", "获取指定场景信号"],
        ["GET", "/signals/output/{scene_id}", "获取标准化信号（支持中英文）"],
        ["GET", "/signals/history/{scene_id}", "信号历史快照"],
    ],
    col_widths=[0.8, 2.2, 3.5],
)

doc.add_heading("7.4 视频流与调试", level=2)
add_table(
    ["方法", "路径", "功能"],
    [
        ["GET", "/stream/{camera_id}", "实时 MJPEG 视频流（带检测叠加）"],
        ["POST", "/debug/simulate", "模拟告警事件（调试用）"],
        ["POST", "/debug/upload-video", "上传视频文件"],
        ["POST", "/agent/chat", "Agent 自然语言对话"],
        ["GET", "/agent/status", "Agent 运行状态"],
    ],
    col_widths=[0.8, 2.0, 3.7],
)

# ═══════════════════════════════════════════════════════════════════════
# 八、关键算法总结
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("八、关键算法总结", level=1)
add_table(
    ["算法", "所在模块", "用途"],
    [
        ["YOLOv8 目标检测", "yolo_service.py", "识别画面中的物体类别和位置"],
        ["IoU 贪心匹配", "tracking_service.py", "帧间目标关联，保持同一 track_id"],
        ["射线法（Ray Casting）", "rules_engine.py", "判断点是否在多边形区域内"],
        ["有向距离法", "rules_engine.py", "判断目标是否越过警戒线"],
        ["线段相交检测", "rules_engine.py", "判断检测框是否与警戒线相交"],
        ["信号锁存 + 冷却", "rules_engine.py", "防止信号闪烁和重复告警"],
        ["PBKDF2-SHA256", "storage_service.py", "密码安全存储"],
        ["关键词 + LLM 意图分类", "agent_orchestrator.py", "Agent 自然语言意图识别"],
    ],
    col_widths=[2.0, 1.8, 2.7],
)

# ═══════════════════════════════════════════════════════════════════════
# 九、定时维护任务
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("九、定时维护任务", level=1)
add_para("系统在每天凌晨 3:00 自动执行深度清理：")
add_bullet("清除24小时前的过期告警（内存 + SQLite）")
add_bullet("清除过期的跟踪轨迹")
add_bullet("清除过期的滞留计时状态")
add_bullet("清除过期的边界待确认状态")
add_bullet("清除过期的规则触发记录")
add_para("防止系统内存和数据库无限增长。")

# ═══════════════════════════════════════════════════════════════════════
# 十、启动方式与部署
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("十、启动方式与部署", level=1)

doc.add_heading("10.1 首次环境准备", level=2)
add_code_block("cd D:\\Project\n.\\setup_env.bat")

doc.add_heading("10.2 环境变量配置（.env）", level=2)
add_code_block(
    'AGENT_ENABLE_LLM="1"                    # 启用 LLM Agent（可选）\n'
    'API_KEY="sk-你的密钥"                     # DeepSeek API 密钥（可选）\n'
    'BASE_URL="https://api.deepseek.com/v1"  # LLM API 地址（可选）\n'
    'MODEL_NAME="deepseek-chat"              # 模型名称（可选）\n'
    'YOLO_WEIGHTS_PATH="models/yolo26s.pt"   # YOLO 权重路径\n'
    'MIMO_API_KEY="..."                       # MiMo 视频理解 API（可选）'
)

doc.add_heading("10.3 启动后端", level=2)
add_code_block(
    ".\\.venv\\Scripts\\python.exe -m uvicorn backend.app.main:app "
    "--host 0.0.0.0 --port 8000 --reload"
)

doc.add_heading("10.4 启动前端", level=2)
add_code_block(
    ".\\.venv\\Scripts\\python.exe -m http.server 5500 --directory frontend\\static"
)

doc.add_heading("10.5 访问地址", level=2)
add_table(
    ["页面", "地址"],
    [
        ["主页", "http://127.0.0.1:5500/index.html"],
        ["场景监控", "http://127.0.0.1:5500/module.html?scene=campus_fence"],
        ["调试页", "http://127.0.0.1:5500/debug.html"],
        ["后端 API 文档", "http://127.0.0.1:8000/docs（Swagger）"],
    ],
    col_widths=[1.5, 5.0],
)

doc.add_heading("10.6 Docker 部署", level=2)
add_code_block("docker compose up -d")
add_para("后端端口：8000 | 前端（Nginx）端口：5500")

# ═══════════════════════════════════════════════════════════════════════
# 十一、自动化测试
# ═══════════════════════════════════════════════════════════════════════
doc.add_heading("十一、自动化测试", level=1)
add_code_block(
    "python -m unittest discover -s backend\\tests -p \"test_*.py\" -v"
)
add_para("当前测试覆盖：")
test_items = [
    "IoU 跟踪器轨迹稳定性",
    "滞留规则触发与信号聚合",
    "认证路由",
    "管理员保护路由",
    "存储服务",
    "Agent 工具与对话",
    "MiMo 视频客户端",
    "视觉后端服务",
    "规则引擎",
]
for t in test_items:
    add_bullet(t)

# ── Footer ─────────────────────────────────────────────────────────────
doc.add_paragraph()
add_separator()
add_para("文档生成日期：2026年5月25日", size=9, color=(0x99, 0x99, 0x99),
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=12)

# ── Save ───────────────────────────────────────────────────────────────
output_path = os.path.expanduser("~/Desktop/项目讲解书.docx")
doc.save(output_path)
print(f"Document saved to: {output_path}")
