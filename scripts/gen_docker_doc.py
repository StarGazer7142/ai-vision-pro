from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── 样式设置 ──────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

def set_cell_shading(cell, color):
    shading_elm = cell._tc.get_or_add_tcPr()
    shd = shading_elm.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear',
    })
    shading_elm.append(shd)

def add_code_block(text):
    for line in text.strip().split('\n'):
        p = doc.add_paragraph(line)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.2
        for run in p.runs:
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        pPr = p._p.get_or_add_pPr()
        shd = pPr.makeelement(qn('w:shd'), {
            qn('w:fill'): 'F0F0F0',
            qn('w:val'): 'clear',
        })
        pPr.append(shd)

def make_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
        set_cell_shading(cell, '2B579A')
    # 数据行
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
            if r % 2 == 0:
                set_cell_shading(cell, 'E8EDF3')
    return table

# ── 标题 ──────────────────────────────────────────────────
title = doc.add_heading('Docker 部署', level=1)

# ── 前提条件 ──────────────────────────────────────────────
doc.add_heading('前提条件', level=2)
doc.add_paragraph('• 已安装 Docker 和 Docker Compose（Docker Desktop 自带）')
doc.add_paragraph('• 确认 Docker 服务已启动')

# ── 步骤 1 ────────────────────────────────────────────────
doc.add_heading('步骤 1：配置环境变量', level=2)
doc.add_paragraph('在项目根目录下，将 .env.example 复制为 .env，并根据实际需求填写关键配置：')
add_code_block('cp .env.example .env')
doc.add_paragraph('')
doc.add_paragraph('需要填写的关键项：')
make_table(
    ['配置项', '说明', '示例'],
    [
        ['API_KEY',        'LLM 服务的 API Key（如 DeepSeek）', 'sk-xxxx'],
        ['MIMO_API_KEY',   'MiMo 视频理解的 API Key（可选）',   'xxxx'],
        ['DEBUG_PASSWORD', '调试工具的登录密码',                 '建议修改默认值'],
    ]
)
doc.add_paragraph('')

# ── 步骤 2 ────────────────────────────────────────────────
doc.add_heading('步骤 2：构建并启动服务', level=2)
add_code_block('docker-compose up --build -d')
doc.add_paragraph('')
p = doc.add_paragraph('首次启动会拉取基础镜像并安装依赖，耗时较长（约 5-15 分钟，取决于网络），后续启动仅需数秒。')

# ── 步骤 3 ────────────────────────────────────────────────
doc.add_heading('步骤 3：访问服务', level=2)
make_table(
    ['服务', '地址'],
    [
        ['前端页面',  'http://localhost:5500'],
        ['后端 API',  'http://localhost:8000'],
        ['健康检查',  'http://localhost:8000/health'],
    ]
)
doc.add_paragraph('')

# ── 常用运维命令 ──────────────────────────────────────────
doc.add_heading('常用运维命令', level=2)
add_code_block(
    '# 查看运行状态\n'
    'docker-compose ps\n'
    '\n'
    '# 查看实时日志\n'
    'docker-compose logs -f\n'
    '\n'
    '# 停止服务\n'
    'docker-compose down\n'
    '\n'
    '# 重启服务（配置变更后）\n'
    'docker-compose down && docker-compose up -d\n'
    '\n'
    '# 重新构建并启动（代码更新后）\n'
    'docker-compose up --build -d'
)

# ── 保存 ──────────────────────────────────────────────────
out = os.path.join(os.path.expanduser('~'), 'Desktop', 'Docker部署说明.docx')
doc.save(out)
print(f'已保存到: {out}')
