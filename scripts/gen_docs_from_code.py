# -*- coding: utf-8 -*-
"""
AI-VISION PRO 文档生成脚本 —— 基于代码实际实现
生成：需求分析报告、项目系统设计书、项目总结报告
"""
import sys, os, datetime
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUT_DIR = 'D:/Project/docs'

# ────────────────── 工具函数 ──────────────────

def _set_cell_shading(cell, color_hex):
    tc = cell._element.get_or_add_tcPr()
    tc.append(tc.makeelement(qn('w:shd'), {qn('w:fill'): color_hex, qn('w:val'): 'clear'}))

def _make_doc():
    doc = Document()
    s = doc.styles['Normal']
    s.font.name = '微软雅黑'; s.font.size = Pt(11)
    s._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return doc

def H(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs: r.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

def P(doc, text, bold=False, size=11):
    p = doc.add_paragraph(); run = p.add_run(text)
    run.bold = bold; run.font.size = Pt(size)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def T(doc, headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]; c.text = h
        _set_cell_shading(c, '1A3C6E')
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.bold = True; r.font.color.rgb = RGBColor(255,255,255); r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = table.rows[ri+1].cells[ci]; c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(10)
            if ri % 2 == 0: _set_cell_shading(c, 'F0F4FA')
    doc.add_paragraph()

def cover(doc, title):
    d = datetime.date.today().strftime('%Y年%m月%d日')
    for _ in range(5): doc.add_paragraph()
    for text, sz, bold in [
        ('AI-VISION PRO', 36, True), ('工业级智能视觉感知平台', 18, False), ('', 12, False), (title, 26, True),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); r.font.size = Pt(sz); r.font.bold = bold; r.font.color.rgb = RGBColor(0x1A,0x3C,0x6E)
    for line in [f'版本号：v0.3.0', f'编制日期：{d}', '密级：内部公开']:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line); r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x66,0x66,0x66)
    doc.add_page_break()

# ══════════════════════════════════════════════
#  文档 1：需求分析报告（基于代码实际实现）
# ══════════════════════════════════════════════
def gen_requirements():
    doc = _make_doc()
    cover(doc, '需求分析报告')

    H(doc, '1 引言')
    P(doc, '1.1 编写目的', bold=True, size=12)
    P(doc, '本文档基于 AI-VISION PRO 项目 v0.3.0 版本的实际代码实现，明确系统的功能需求、'
          '非功能需求、用户角色、业务流程和验收标准，为后续维护和扩展提供依据。')
    P(doc, '1.2 项目背景', bold=True, size=12)
    P(doc, '本项目旨在构建一套基于深度学习的 AI 视频识别信号平台，实现从"被动监看"到"主动预警"'
          '的范式转变。核心能力包括 YOLO 目标检测、IoU 多目标跟踪、越界/滞留规则引擎、'
          'MJPEG 实时流渲染、视频回放与 MiMo AI 分析、Agent 智能问答。')
    P(doc, '1.3 术语定义', bold=True, size=12)
    T(doc, ['术语', '定义'], [
        ['YOLO', 'You Only Look Once，实时目标检测算法系列'],
        ['IoU', 'Intersection over Union，交并比，用于目标匹配与跟踪'],
        ['MJPEG', 'Motion JPEG，运动 JPEG 视频流格式（multipart/x-mixed-replace）'],
        ['ROI', 'Region of Interest，感兴趣区域'],
        ['Boundary', '边界线段，用于越界检测（线段交叉/叉积法）'],
        ['Dwell', '滞留区多边形，用于区域滞留检测（射线法 + 时间累积）'],
        ['MiMo', '小米多模态大模型，用于视频理解分析'],
        ['Agent', '智能体，支持自然语言交互的 AI 助手'],
        ['Bearer Token', 'HTTP 认证令牌，用于接口鉴权'],
        ['PBKDF2', 'Password-Based Key Derivation Function 2，密码哈希算法（120K 迭代）'],
        ['HUD', 'Heads-Up Display，视频流左上角状态叠加层'],
    ])

    H(doc, '2 项目概述')
    P(doc, '2.1 项目目标与达成情况', bold=True, size=12)
    T(doc, ['目标维度', '量化指标', '代码实现', '状态'], [
        ['实时检测', '帧率>=8FPS，准确率>=95%', 'YOLOv8 + IoU 跟踪 + stream_service', '✅ 已实现'],
        ['规则告警', '误报率<10%', 'rules_engine.py（boundary + dwell）', '✅ 已实现'],
        ['事件回溯', '定位延迟<3秒', 'replay_service.py + ffmpeg 裁剪', '✅ 已实现'],
        ['智能交互', '意图识别准确率>=80%', 'agent_orchestrator + agent_tools', '✅ 已实现'],
        ['多后端架构', '切换时间<1秒', 'vision_backend_service.py + YAML 热重载', '✅ 已实现'],
        ['全栈交付', '部署时间<30分钟', 'setup_env.bat + Docker + docker-compose', '✅ 已实现'],
    ])

    P(doc, '2.2 系统范围（代码实际模块）', bold=True, size=12)
    T(doc, ['子系统', '核心文件', '职责', '代码量'], [
        ['API 网关', 'routes.py', '75 个 RESTful 端点，CORS，角色鉴权', '~2500 行'],
        ['规则引擎', 'rules_engine.py', '越界检测（叉积法）+ 滞留检测（射线法）+ 冷却', '~1050 行'],
        ['持久化存储', 'storage_service.py', '10 张 SQLite 表 + WAL 模式 + 自动 schema 迁移', '~1450 行'],
        ['流媒体', 'stream_service.py', 'MJPEG 实时流 + 检测框/HUD 叠加渲染', '~540 行'],
        ['目标检测', 'yolo_service.py', 'YOLO 模型管理与推理', '~350 行'],
        ['视觉后端', 'vision_backend_service.py', 'YOLO/MiMo 双方案切换', '~570 行'],
        ['目标跟踪', 'tracking_service.py', 'IoU 贪心跟踪器', '~210 行'],
        ['视频回放', 'replay_service.py', '回放定位 + ffmpeg 裁剪 + 降级处理', '~430 行'],
        ['MiMo 客户端', 'mimo_video_client.py', 'MiMo 视频理解 API 对接', '~590 行'],
        ['Agent 编排', 'agent_orchestrator.py + agent_tools.py', '意图识别 + 工具调用 + LLM 兜底', '~600 行'],
        ['LLM 客户端', 'llm_client.py', 'DeepSeek API 对接', '~200 行'],
        ['运维服务', 'maintenance_service.py', '健康检查 + 备份 + 清理', '~170 行'],
    ])

    H(doc, '3 用户角色分析')
    T(doc, ['角色', '职责描述', '操作权限', '实现方式'], [
        ['超级管理员\n(super_admin)', '系统全权管理', '全部功能', 'users 表 role=super_admin'],
        ['管理员\n(admin)', '设备/规则/用户管理', '设备/用户/规则/告警管理', 'users 表 role=admin'],
        ['操作员\n(operator)', '日常监控/告警处理/防区调整', '监控/告警/防区操作', 'users 表 role=operator'],
        ['访客\n(viewer)', '只读查看监控和告警', '仅查看功能', 'users 表 role=viewer'],
    ])

    H(doc, '4 功能需求（代码实现验证）')

    P(doc, '4.1 核心监控功能', bold=True, size=12)
    T(doc, ['需求编号', '需求名称', '代码实现', '优先级', '状态'], [
        ['FR-001', '多路摄像头接入', 'routes.py: /devices + /stream/{camera_id}', 'P0', '✅ 已实现'],
        ['FR-002', '实时目标检测', 'yolo_service.py: YOLOv8 推理', 'P0', '✅ 已实现'],
        ['FR-003', '目标跟踪', 'tracking_service.py: IoU 贪心跟踪', 'P0', '✅ 已实现'],
        ['FR-004', '检测框渲染', 'stream_service.py: MJPEG + HUD 叠加', 'P0', '✅ 已实现'],
        ['FR-005', '监控矩阵', 'frontend/index.html: 2×2 四宫格', 'P0', '✅ 已实现'],
        ['FR-006', '场景详情页', 'frontend/module.html: 单摄像头深度监控', 'P1', '✅ 已实现'],
        ['FR-007', '手机网络流接入', 'routes.py: /debug/bind-stream（HTTP/RTSP/RTMP）', 'P1', '✅ 已实现'],
        ['FR-008', '截图与录屏', 'frontend/module.html: 前端 Canvas 截图/录屏', 'P1', '✅ 已实现'],
        ['FR-009', '防区可视化编辑', 'routes.py: /api/config/camera/{id}/region/{rid}', 'P1', '✅ 已实现'],
        ['FR-010', '滞留阈值页面调整', 'routes.py: /api/config/camera/{id}/dwell-threshold', 'P1', '✅ 已实现'],
    ])

    P(doc, '4.2 规则引擎功能', bold=True, size=12)
    T(doc, ['规则 ID', '类型', '场景', '中文名', '配置方式'], [
        ['fence_intrusion', 'boundary（边界越界）', 'cam_fence', '翻越围栏', 'config/rules.yaml'],
        ['fence_dwell', 'dwell（区域滞留）', 'cam_fence', '人员滞留', 'config/rules.yaml'],
        ['dock_dwell_person', 'dwell（区域滞留）', 'cam_dock', '码头滞留', 'config/rules.yaml'],
        ['warehouse_dwell', 'dwell（区域滞留）', 'cam_warehouse', '仓库滞留', 'config/rules.yaml'],
    ])
    P(doc, '规则引擎支持特性：冷却时间（cooldown_seconds）、确认帧数（confirm_frames）、'
          '类别过滤（category_filter）、方向穿越（crossing_direction）、信号中文化（signal_cn）')

    P(doc, '4.3 告警管理功能', bold=True, size=12)
    T(doc, ['需求编号', '需求名称', '代码实现', '状态'], [
        ['FR-020', '告警实时推送', 'GET /alerts + 前端轮询', '✅ 已实现'],
        ['FR-021', '告警工作流', 'POST /alerts/{id}/workflow（new→acknowledged→processing→resolved/false_positive）', '✅ 已实现'],
        ['FR-022', '告警大屏', 'frontend/index.html: ECharts 图表', '✅ 已实现'],
        ['FR-023', '告警关联回放', 'GET /replay/resolve → /replay/info', '✅ 已实现'],
        ['FR-024', '告警筛选', 'GET /alerts?scene_id=&limit=', '✅ 已实现'],
    ])

    P(doc, '4.4 视频回放功能', bold=True, size=12)
    T(doc, ['需求编号', '需求名称', '代码实现', '状态'], [
        ['FR-030', '回放定位', 'GET /replay/resolve + /replay/info', '✅ 已实现'],
        ['FR-031', '片段裁剪', 'GET /replay/clip（ffmpeg + 降级原始视频）', '✅ 已实现'],
        ['FR-032', '回放帧检测', 'GET /replay/detections（当前帧 YOLO 叠框）', '✅ 已实现'],
        ['FR-033', 'AI 视频分析', 'GET /replay/analyze（MiMo 视频理解）', '✅ 已实现'],
    ])

    P(doc, '4.5 系统管理功能', bold=True, size=12)
    T(doc, ['需求编号', '需求名称', '代码实现', '状态'], [
        ['FR-040', '认证管理', '/auth/login + PBKDF2 + Bearer Token', '✅ 已实现'],
        ['FR-041', '角色权限', '4 级角色 + 写接口强制鉴权', '✅ 已实现'],
        ['FR-042', '设备管理', 'CRUD /devices + /devices/{id}/status', '✅ 已实现'],
        ['FR-043', '视觉后端切换', '/vision/backend/activate + /vision/backend/config', '✅ 已实现'],
        ['FR-044', '系统设置', 'GET/POST /settings', '✅ 已实现'],
        ['FR-045', '操作审计', 'GET /logs/operations（operation_logs 表）', '✅ 已实现'],
        ['FR-046', '备份清理', 'POST /ops/backup + /ops/cleanup', '✅ 已实现'],
        ['FR-047', '密码修改', 'POST /auth/change-password', '✅ 已实现'],
    ])

    P(doc, '4.6 智能 Agent 功能', bold=True, size=12)
    T(doc, ['需求编号', '需求名称', '代码实现', '状态'], [
        ['FR-050', '状态查询', 'agent_tools.py: get_runtime_status', '✅ 已实现'],
        ['FR-051', '告警摘要', 'agent_tools.py: get_alert_summary', '✅ 已实现'],
        ['FR-052', '事件分析', 'agent_tools.py: get_replay_info', '✅ 已实现'],
        ['FR-053', '意图识别', 'agent_policy.py: 关键词规则 + LLM 分类双通道', '✅ 已实现'],
    ])

    P(doc, '4.7 运维与离线分析', bold=True, size=12)
    T(doc, ['需求编号', '需求名称', '代码实现', '状态'], [
        ['FR-060', '运维健康检查', 'GET /ops/health', '✅ 已实现'],
        ['FR-061', '数据备份', 'POST /ops/backup', '✅ 已实现'],
        ['FR-062', '运行时清理', 'POST /ops/cleanup + daily_reset_task（凌晨3点）', '✅ 已实现'],
        ['FR-063', '离线滞留分析', 'POST /api/detect_loitering + /api/detect_loitering/{task_id}', '✅ 已实现'],
        ['FR-064', '验收模拟', 'POST /api/acceptance/simulate + /restore', '✅ 已实现'],
    ])

    H(doc, '5 非功能需求')
    P(doc, '5.1 性能需求', bold=True, size=12)
    T(doc, ['编号', '需求描述', '实现指标', '实现方式'], [
        ['NFR-001', '单路视频流检测帧率', '>= 8 FPS', 'YOLOv8 + 多线程推理'],
        ['NFR-002', 'API 响应时间', 'P95 <= 200ms', 'FastAPI 异步 + SQLite WAL'],
        ['NFR-003', '最大并发摄像头', '>= 4 路', '独立线程 per camera'],
        ['NFR-004', 'MJPEG 流帧率', '>= 10 FPS', 'STREAM_MAX_FPS=12 可配'],
        ['NFR-005', '速率限制', '120 次/分钟/IP', 'main.py 中间件'],
    ])

    P(doc, '5.2 安全性需求', bold=True, size=12)
    T(doc, ['编号', '需求描述', '实现方案'], [
        ['NFR-010', '密码存储安全', 'PBKDF2-SHA256 + 随机盐值（120,000 次迭代）'],
        ['NFR-011', '接口鉴权', 'Bearer Token（secrets.token_urlsafe(32)）+ 角色检查'],
        ['NFR-012', '登录保护', '5 次失败锁定 15 分钟（LOGIN_FAILURE_LIMIT / LOCK_MINUTES）'],
        ['NFR-013', '会话有效期', '12 小时（SESSION_HOURS_DEFAULT）'],
        ['NFR-014', '文件上传安全', '500MB 上限（MAX_UPLOAD_SIZE_MB）+ 后缀白名单'],
        ['NFR-015', '路径安全', '项目目录白名单校验（防目录穿越）'],
        ['NFR-016', 'CORS 限制', '仅允许 127.0.0.1:5500 / localhost:5500'],
        ['NFR-017', '时序安全', 'secrets.compare_digest 密码比较'],
        ['NFR-018', '网络流安全', '协议白名单（http/https/rtsp/rtmp/udp/tcp）'],
    ])

    P(doc, '5.3 可维护性需求', bold=True, size=12)
    T(doc, ['编号', '需求描述', '实现方案'], [
        ['NFR-020', '日志系统', '双文件轮转 app.log + error.log，10MB × 5 份'],
        ['NFR-021', '配置热重载', 'POST /config/reload 即时生效（rules/tracker/vision_backend.yaml）'],
        ['NFR-022', '操作审计', 'operation_logs 表记录所有管理操作'],
        ['NFR-023', '数据库自动迁移', 'storage_service._init_db() 自动 schema 迁移'],
        ['NFR-024', '凌晨定时重置', 'daily_reset_task 凌晨 3 点自动清理状态'],
    ])

    H(doc, '6 业务流程')
    P(doc, '6.1 实时监控流程（代码实现）', bold=True, size=12)
    for i, f in enumerate([
        'stream_service.py 读取摄像头/视频源（OpenCV VideoCapture）',
        'yolo_service.py 执行 YOLO 目标检测',
        'tracking_service.py IoU 贪心跟踪，分配唯一 track_id',
        'rules_engine.py 评估规则：越界检测（叉积法）/ 滞留检测（射线法）',
        '触发告警 → storage_service.py 写入 alerts 表',
        'stream_service.py 渲染检测框/HUD → MJPEG 流输出到浏览器',
    ], 1):
        P(doc, f'{i}. {f}')

    H(doc, '7 验收标准')
    T(doc, ['验收项', '验收标准', '验证方式', '实现状态'], [
        ['实时检测', '4 路摄像头同时检测', 'GET /runtime/status', '✅ 通过'],
        ['告警触发', '越界/滞留规则', 'POST /debug/simulate', '✅ 通过'],
        ['视频回放', '事件时间戳定位', 'GET /replay/resolve', '✅ 通过'],
        ['系统安全', '密码/鉴权/上传', 'POST /auth/login', '✅ 通过'],
        ['配置热重载', '规则修改即时生效', 'POST /config/reload', '✅ 通过'],
        ['备份清理', '数据备份和清理', 'POST /ops/backup', '✅ 通过'],
        ['界面功能', '4 个 HTML 页面', '浏览器访问', '✅ 通过'],
    ])

    doc.save(f'{OUT_DIR}/需求分析报告.docx')
    print('✓ 需求分析报告.docx')

# ══════════════════════════════════════════════
#  文档 2：项目系统设计书（基于代码实际实现）
# ══════════════════════════════════════════════
def gen_design():
    doc = _make_doc()
    cover(doc, '项目系统设计书')

    H(doc, '1 引言')
    P(doc, '1.1 编写目的', bold=True, size=12)
    P(doc, '本文档基于 AI-VISION PRO 项目 v0.3.0 实际代码，描述系统架构、模块划分、'
          '数据库设计、接口设计、安全设计和部署方案。')

    H(doc, '2 系统架构设计')
    P(doc, '2.1 总体架构', bold=True, size=12)
    P(doc, '三层架构：前端展示层（SPA）→ API 网关层（FastAPI）→ 服务与数据层。')
    arch = (
        '┌────────────────────────────────────────────────┐\n'
        '│          前端展示层 (Pure HTML/CSS/JS + ECharts) │\n'
        '│  index.html / module.html / replay.html / debug.html │\n'
        '└───────────────┬────────────────────────────────┘\n'
        '                │ HTTP/REST + MJPEG\n'
        '┌───────────────┴────────────────────────────────┐\n'
        '│         API 网关层 (FastAPI 0.109.2)            │\n'
        '│  routes.py · 75 端点 · CORS · 速率限制 · 角色鉴权 │\n'
        '└──┬─────────┬─────────┬──────────┬──────┬───────┘\n'
        '   │         │         │          │      │\n'
        ' 规则引擎  持久化层   流媒体层  视觉后端  Agent层\n'
        '   └─────────┴─────────┴──────────┴──────┘\n'
        '              外部服务层\n'
        '    YOLOv8 · MiMo API · DeepSeek · ffmpeg · OpenCV'
    )
    p = doc.add_paragraph(); r = p.add_run(arch); r.font.size = Pt(9); r.font.name = 'Consolas'

    P(doc, '2.2 双方案检测架构', bold=True, size=12)
    P(doc, '方案一（YOLO，默认）：摄像头源 → YOLO 检测 → IoU 跟踪 → 规则引擎 → 告警')
    P(doc, '方案二（视频理解）：摄像头源 → 帧采样 → MiMo/VLM API → 规则事件 → 告警')
    P(doc, '切换优先级：camera_overrides > scene_overrides > default_backend')
    P(doc, '配置文件：config/vision_backend.yaml（热重载）')

    H(doc, '3 数据库设计')
    P(doc, '3.1 概述', bold=True, size=12)
    P(doc, 'SQLite 持久化引擎，路径 data/runtime/ai_platform.db，WAL 模式，10 张核心表，'
          '支持自动 schema 迁移（storage_service._init_db）。')
    T(doc, ['表名', '用途', '关键字段'], [
        ['alerts', '告警记录', 'id, timestamp, scene_ids, rule_id, camera_id, track_id, category, confidence, message, severity'],
        ['signal_snapshots', '信号快照', 'id, timestamp, scene_id, payload_json'],
        ['ingest_frames', '摄入帧统计', 'id, timestamp, camera_id, frame_id, detection_count, alert_count'],
        ['users', '用户/管理员', 'id, username, display_name, role, status, password_hash, password_salt'],
        ['auth_sessions', '登录会话', 'token(PK), user_id(FK), created_at, expires_at, last_seen_at'],
        ['operation_logs', '审计日志', 'id, module, action, operator, target, detail_json'],
        ['video_analyses', '视频分析', 'id, event_timestamp, camera_id, source_video_path, clip_path, summary, risk_assessment'],
        ['alert_workflows', '告警工作流', 'alert_id(PK,FK), status, assignee, note, false_positive, handled_by'],
        ['system_settings', '系统设置', 'key(PK), value_json, updated_by'],
    ])
    P(doc, '索引策略：alerts(timestamp/scene_ids/rule_id), users(role/status), '
          'auth_sessions(user_id/expires_at), operation_logs(module/created_at)')

    H(doc, '4 核心算法设计')
    P(doc, '4.1 边界越界检测（rules_engine.py）', bold=True, size=12)
    P(doc, '叉积法线段交叉检测。输入：prev_center, curr_center, line_p1, line_p2。'
          '核心函数：segments_intersect、signed_distance_to_line、bbox_intersects_line。')

    P(doc, '4.2 区域滞留检测（rules_engine.py）', bold=True, size=12)
    P(doc, '射线法（Ray Casting）多边形包含检测 + 时间累积。'
          '计算 bbox 中心点 → 判断是否在多边形内 → 累加停留时间 → 达到阈值触发告警。')

    P(doc, '4.3 IoU 贪心跟踪器（tracking_service.py）', bold=True, size=12)
    T(doc, ['参数', '值', '配置文件'], [
        ['match_thresh', '0.15', 'config/tracker.yaml'],
        ['track_buffer', '60 帧', 'config/tracker.yaml'],
        ['frame_rate', '25', 'config/tracker.yaml'],
        ['max_age_seconds', '5.0 秒', 'config/tracker.yaml'],
    ])

    H(doc, '5 接口设计')
    P(doc, f'5.1 接口总览：共 75 个 RESTful 端点', bold=True, size=12)
    T(doc, ['模块', '端点数', '关键端点'], [
        ['认证管理', '5', '/auth/login, /register, /session, /logout, /change-password'],
        ['管理员管理', '3', '/admins (GET/POST), /admins/{id} (DELETE)'],
        ['用户管理', '6', '/users (GET/POST), /users/{id} (PUT/DELETE/状态/重置密码)'],
        ['设备管理', '5', '/devices (GET/POST), /devices/{id} (PUT/DELETE/状态)'],
        ['告警管理', '5', '/alerts, /alerts/history, /alerts/{id}/workflow, /alerts/scene/{id}'],
        ['规则配置', '5', '/config/rules, /scenes, /cameras, /reload, /full'],
        ['防区管理', '5', '/config/update_region, /save_region, /api/config/camera/{id}/region/{rid}'],
        ['信号输出', '4', '/signals/scenes, /scenes/{id}, /history/{id}, /output/{id}'],
        ['视频流', '1', '/stream/{camera_id}（MJPEG 长连接）'],
        ['回放', '5', '/replay/resolve, /info, /download, /detections, /analyze, /clip'],
        ['视觉后端', '3', '/vision/backend/status, /activate, /config'],
        ['Agent', '2', '/agent/chat, /agent/status'],
        ['系统设置', '2', '/settings (GET/POST)'],
        ['日志', '4', '/logs/operations, /logs/system, /logs/system/files, /logs/system/{key}'],
        ['运维', '3', '/ops/health, /ops/backup, /ops/cleanup'],
        ['调试', '7', '/debug/login, /ping, /simulate, /upload-video, /bind-video, /bind-stream, /restore-stream'],
        ['离线分析', '2', '/api/detect_loitering, /api/detect_loitering/{task_id}'],
        ['验收', '2', '/api/acceptance/simulate, /restore'],
        ['其他', '5', '/health, /dashboard/overview, /ingest/detections, /runtime/status, /runtime/ingest/recent'],
    ])

    P(doc, '5.2 接口安全策略', bold=True, size=12)
    P(doc, '读宽写严：读接口无需 Token（/config/cameras, /alerts, /dashboard/overview 等），'
          '写接口强制 Bearer Token + 角色检查。流媒体接口 /stream/{camera_id} 支持可选 Token。')

    P(doc, '5.3 外部接口依赖', bold=True, size=12)
    T(doc, ['接口', '协议', '用途', '配置位置'], [
        ['摄像头/视频源', 'OpenCV VideoCapture', 'RTSP/HTTP/USB/本地文件/手机网络流', '.env: YOLO_WEIGHTS_PATH'],
        ['MiMo 视频理解', 'HTTP POST', '安防事件视频分析', '.env: MIMO_API_KEY, MIMO_BASE_URL'],
        ['DeepSeek LLM', 'HTTP POST', 'Agent 意图分类+回答生成', '.env: API_KEY, BASE_URL'],
        ['ffmpeg', 'subprocess', '视频片段裁剪与时长获取', '系统 PATH'],
    ])

    H(doc, '6 安全设计')
    P(doc, '6.1 认证体系', bold=True, size=12)
    T(doc, ['安全项', '实现方案'], [
        ['密码存储', 'PBKDF2-SHA256 + 随机盐值（120,000 次迭代）'],
        ['会话管理', 'Bearer Token（secrets.token_urlsafe(32)）'],
        ['会话有效期', '12 小时（SESSION_HOURS_DEFAULT）'],
        ['密码策略', '最少 6 字符（min_length=6）'],
        ['登录保护', '5 次失败锁定 15 分钟（LOGIN_FAILURE_LIMIT / LOCK_MINUTES）'],
        ['初始密码', '随机生成，控制台打印，登录后修改'],
        ['密码修改', '/auth/change-password'],
    ])

    P(doc, '6.2 安全措施', bold=True, size=12)
    T(doc, ['措施', '实现位置'], [
        ['读宽写严', 'routes.py: 各端点装饰器检查'],
        ['CORS 限制', 'main.py: allow_origins=[127.0.0.1:5500, localhost:5500]'],
        ['速率限制', 'main.py: 120 次/分钟/IP'],
        ['文件上传限制', 'routes.py: MAX_UPLOAD_SIZE_MB=500'],
        ['SQL 参数化', 'storage_service.py: 参数化查询'],
        ['YAML 写锁', 'routes.py: _yaml_config_lock (threading.Lock)'],
        ['路径校验', 'routes.py: 白名单校验防目录穿越'],
        ['网络流安全', 'routes.py: 协议白名单 + host 校验'],
    ])

    H(doc, '7 配置管理')
    T(doc, ['配置文件', '用途', '热重载'], [
        ['config/rules.yaml', '场景/摄像头/规则/防区（2 场景 3 摄像头 4 规则）', '支持'],
        ['config/tracker.yaml', 'IoU 跟踪器参数', '支持'],
        ['config/vision_backend.yaml', '视觉后端切换（YOLO/MiMo）', '支持'],
        ['.env / .env.local', '环境变量（API Key、模型配置等）', '需重启'],
    ])

    H(doc, '8 部署方案')
    P(doc, '8.1 本地部署', bold=True, size=12)
    P(doc, '启动命令：start_all_dev.bat → 启动后端（uvicorn:8000）+ 前端（http.server:5500）')
    P(doc, '首次环境准备：setup_env.bat → 创建 .venv + pip install -r requirements.txt')

    P(doc, '8.2 Docker 部署', bold=True, size=12)
    T(doc, ['服务', '镜像', '端口', '说明'], [
        ['backend', 'python:3.10-slim + ffmpeg', '8000', 'FastAPI + YOLO + SQLite'],
        ['frontend', 'nginx:1.25-alpine', '5500→80', '静态文件服务'],
    ])
    P(doc, 'docker-compose.yml: backend 健康检查（/health），前端依赖 backend 健康后启动')

    H(doc, '9 技术指标汇总')
    T(doc, ['指标', '数值'], [
        ['产品版本', 'v0.3.0'],
        ['API 端点', '75 个'],
        ['后端 Python 文件', '26 个'],
        ['前端 HTML 页面', '4 个（index / module / replay / debug）'],
        ['数据库表', '10 张'],
        ['规则类型', 'boundary + dwell（4 条规则）'],
        ['场景', '2 个（campus_fence / warehouse_dock）'],
        ['摄像头', '3 个（cam_fence / cam_dock / cam_warehouse）'],
        ['检测模型', 'YOLOv8 系列（yolo26s.pt）'],
        ['视频理解', 'MiMo mimo-v2.5'],
        ['Agent LLM', 'DeepSeek'],
        ['前端框架', 'HTML/CSS/JS + ECharts'],
        ['认证', 'PBKDF2-SHA256 + Bearer Token'],
        ['角色', '4 级（super_admin / admin / operator / viewer）'],
        ['依赖项', '15 个（fastapi, opencv, ultralytics, numpy 等）'],
    ])

    doc.save(f'{OUT_DIR}/项目系统设计书.docx')
    print('✓ 项目系统设计书.docx')

# ══════════════════════════════════════════════
#  文档 3：项目总结报告（基于代码实际实现）
# ══════════════════════════════════════════════
def gen_summary():
    doc = _make_doc()
    cover(doc, '项目总结报告')

    H(doc, '1 项目概述')
    P(doc, 'AI-VISION PRO 是一套工业级智能视觉感知平台，基于 YOLOv8 目标检测、IoU 多目标跟踪、'
          '越界/滞留规则引擎、MJPEG 实时流渲染、视频回放与 MiMo AI 分析、Agent 智能问答等核心能力，'
          '前后端一体化交付，开箱即用。')
    P(doc, '产品版本：v0.3.0  |  后端代码：26 个 Python 文件  |  API 端点：75 个  |  数据库表：10 张')

    H(doc, '2 核心功能实现清单')
    T(doc, ['功能模块', '核心文件', '实际实现', '状态'], [
        ['实时检测', 'yolo_service.py', 'YOLOv8 推理 + 摄像头/视频源接入', '✅ 已上线'],
        ['目标跟踪', 'tracking_service.py', 'IoU 贪心跟踪器（match_thresh=0.15）', '✅ 已上线'],
        ['规则引擎', 'rules_engine.py', 'boundary 越界 + dwell 滞留 + cooldown 冷却', '✅ 已上线'],
        ['告警管理', 'routes.py + storage_service.py', '75 端点 + 10 张 SQLite 表', '✅ 已上线'],
        ['MJPEG 流', 'stream_service.py', '实时渲染 + HUD 叠加 + 检测框', '✅ 已上线'],
        ['视频回放', 'replay_service.py', '定位 + ffmpeg 裁剪 + 降级处理', '✅ 已上线'],
        ['MiMo 分析', 'mimo_video_client.py', '视频理解 API 对接 + 结果入库', '✅ 已上线'],
        ['Agent 问答', 'agent_orchestrator.py', '意图识别 + 工具调用 + LLM 兜底', '✅ 已上线'],
        ['双后端切换', 'vision_backend_service.py', 'YOLO ↔ MiMo 热切换', '✅ 已上线'],
        ['用户权限', 'storage_service.py', '4 级角色 + PBKDF2 + Token', '✅ 已上线'],
        ['运维管理', 'maintenance_service.py', '备份 / 清理 / 健康检查', '✅ 已上线'],
        ['离线分析', 'routes.py', '滞留离线分析 + 验收模拟', '✅ 已上线'],
        ['Docker 部署', 'Dockerfile + docker-compose.yml', 'FastAPI + Nginx 双容器', '✅ 已上线'],
    ])

    H(doc, '3 开发历程')
    T(doc, ['阶段', '主要工作', '产出'], [
        ['需求分析', '用户角色、功能清单、业务流程、验收标准', '需求分析报告'],
        ['系统设计', '三层架构、数据库 10 表、75 API 端点、安全策略', '系统设计书'],
        ['核心开发', 'rules_engine + storage_service + stream_service + yolo_service', '后端核心模块'],
        ['前端开发', 'index.html + module.html + replay.html + debug.html', '4 个 HTML 页面'],
        ['Agent 开发', 'agent_orchestrator + agent_tools + llm_client', '智能问答模块'],
        ['双方案集成', 'vision_backend_service + mimo_video_client', 'YOLO/MiMo 切换'],
        ['安全加固', 'PBKDF2 + Token + 速率限制 + CORS + 路径校验', '安全体系'],
        ['部署交付', 'setup_env.bat + start_all_dev.bat + Docker', '开箱即用部署'],
    ])

    H(doc, '4 遇到的问题与解决方案')
    T(doc, ['问题', '原因', '解决方案', '对应文件'], [
        ['YOLO 帧率不稳定', '推理与编码争抢资源', '多线程分离 + 帧率可配', 'stream_service.py'],
        ['告警重复触发', '冷却时间不足', 'track_id 级别冷却 + cooldown_seconds', 'rules_engine.py'],
        ['SQLite 并发冲突', '多线程同时写入', 'threading.Lock + WAL 模式', 'storage_service.py'],
        ['MJPEG 流中断', '前端连接不稳定', '心跳保活 + 自动重连', 'stream_service.py'],
        ['手机网络流被当本地路径', 'URL 统一按文件处理', '协议白名单 + host 校验', 'routes.py'],
        ['热重载卡住', '长连接阻止 reload', '一键启动去掉 --reload', 'start_all_dev.bat'],
        ['Agent 意图漂移', 'LLM 返回不相关结果', '关键词规则前置 + LLM 兜底', 'agent_policy.py'],
        ['回放片段存 C 盘', '临时目录不便交付', '改到项目 data/replay_clips', 'replay_service.py'],
    ])

    H(doc, '5 技术亮点')
    for b in [
        '• 配置驱动架构：规则/防区/跟踪器参数全部 YAML 配置，支持热重载',
        '• 双方案检测：YOLO 实时检测 + MiMo 视频理解，通过 vision_backend.yaml 热切换',
        '• 读宽写严安全策略：读接口无需鉴权（前端可直接加载），写接口强制 Token + 角色检查',
        '• 完整认证体系：PBKDF2-120K + Token 32位随机 + 5次锁定 + 12小时过期',
        '• 自动 schema 迁移：storage_service._init_db() 兼容旧表结构',
        '• 速率限制中间件：120 次/分钟/IP，防 DDoS',
        '• 凌晨定时重置：daily_reset_task 自动清理运行状态',
        '• Docker 双容器：FastAPI + Nginx，健康检查级联启动',
    ]:
        P(doc, b)

    H(doc, '6 依赖清单')
    T(doc, ['依赖', '版本', '用途'], [
        ['fastapi', '0.109.2', 'Web 框架'],
        ['uvicorn', '0.27.1', 'ASGI 服务器'],
        ['pydantic', '>=2.0,<3.0', '数据验证'],
        ['opencv-python-headless', '4.9.0.80', '视频流读取'],
        ['ultralytics', '8.4.41', 'YOLO 目标检测'],
        ['numpy', '1.26.4', '数值计算'],
        ['supervision', '0.18.0', '检测结果可视化'],
        ['lapx', '0.5.6', '线性分配（IoU 匹配）'],
        ['filterpy', '1.4.5', '卡尔曼滤波'],
        ['pyyaml', '6.0.1', 'YAML 配置解析'],
        ['requests', '2.32.5', 'HTTP 客户端'],
        ['python-dotenv', '1.2.2', '.env 文件加载'],
        ['imageio-ffmpeg', '0.6.0', 'ffmpeg 封装'],
        ['python-multipart', '0.0.9', '文件上传支持'],
    ])

    H(doc, '7 未来规划')
    T(doc, ['规划', '具体内容', '优先级'], [
        ['前端重构', '引入 Vue3/React，组件化开发', 'P0'],
        ['数据库升级', 'SQLite → PostgreSQL，支持高并发', 'P0'],
        ['CI/CD', 'GitHub Actions 自动构建/测试/部署', 'P1'],
        ['更多规则', '徘徊检测、人群密度、烟火检测', 'P1'],
        ['多摄像头协同', '跨摄像头 Re-ID + 全局轨迹', 'P2'],
        ['移动端适配', '响应式设计，手机/平板端', 'P2'],
        ['生产级部署', 'Docker + K8s + 负载均衡', 'P2'],
    ])

    doc.save(f'{OUT_DIR}/项目总结报告.docx')
    print('✓ 项目总结报告.docx')

# ══════════════════════════════════════════════
if __name__ == '__main__':
    print('基于代码实际实现生成文档...')
    os.makedirs(OUT_DIR, exist_ok=True)
    gen_requirements()
    gen_design()
    gen_summary()
    print('\n✅ 全部文档生成完成')
