"""生成符合商业规范的需求分析报告和项目系统设计书"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

DESKTOP = os.path.expanduser("~\\Desktop")

def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'})
    shading.append(shd)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor(255,255,255)
        set_cell_shading(c, '2F5496')
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = table.rows[ri+1].cells[ci]; c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
            if ri % 2 == 1: set_cell_shading(c, 'D6E4F0')
    doc.add_paragraph()

def h(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for r in heading.runs: r.font.color.rgb = RGBColor(47,84,150)
    return heading

def p(doc, text, bold=False, size=10.5, align=None):
    para = doc.add_paragraph()
    r = para.add_run(text); r.font.size = Pt(size); r.bold = bold
    if align: para.alignment = align
    return para

def code(doc, text):
    para = doc.add_paragraph()
    r = para.add_run(text); r.font.name = 'Consolas'; r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(50,50,50)
    para.paragraph_format.left_indent = Cm(0.5)
    return para

def make_doc():
    doc = Document()
    s = doc.styles['Normal']; s.font.name = '微软雅黑'; s.font.size = Pt(10.5)
    s.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return doc

def add_doc_control(doc, doc_title, doc_id):
    """添加文档控制信息页"""
    h(doc, '文档控制信息', 1)
    add_table(doc, ['属性', '内容'], [
        ['文档编号', doc_id],
        ['文档名称', doc_title],
        ['产品名称', 'AI-VISION PRO'],
        ['产品版本', 'v0.3.0'],
        ['文档版本', 'V1.0'],
        ['密级', '内部公开'],
        ['编制日期', '2026年5月25日'],
    ])
    doc.add_paragraph()
    h(doc, '审批记录', 2)
    add_table(doc, ['角色', '姓名', '日期', '签字'], [
        ['编制人', '', '2026-05-25', ''],
        ['审核人', '', '', ''],
        ['批准人', '', '', ''],
    ])
    doc.add_paragraph()
    h(doc, '修订历史', 2)
    add_table(doc, ['版本', '日期', '修订人', '修订内容'], [
        ['V1.0', '2026-05-25', '', '初稿'],
    ])
    doc.add_page_break()

def add_cover(doc, subtitle):
    for _ in range(3): doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('AI-VISION PRO'); r.font.size = Pt(36); r.bold = True; r.font.color.rgb = RGBColor(47,84,150)
    t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run(subtitle); r2.font.size = Pt(26); r2.font.color.rgb = RGBColor(47,84,150)
    doc.add_paragraph()
    p(doc, '工业级智能视觉感知平台', size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(3): doc.add_paragraph()
    add_table(doc, ['属性', '内容'], [
        ['产品名称', 'AI-VISION PRO'],
        ['版本号', 'v0.3.0'],
        ['文档版本', 'V1.0'],
        ['编制日期', '2026年5月25日'],
    ])
    doc.add_page_break()

def add_toc_placeholder(doc):
    """目录占位（Word中按Ctrl+A再按F9可更新域）"""
    h(doc, '目录', 1)
    p(doc, '（请在Word中按 Ctrl+A 全选后按 F9 更新目录域）', size=9)
    doc.add_page_break()


# ================================================================
#  文档1: 需求分析报告
# ================================================================
def gen_requirements():
    doc = make_doc()
    add_cover(doc, '需求分析报告')
    add_doc_control(doc, 'AI-VISION PRO 需求分析报告', 'AIR-REQ-001')
    add_toc_placeholder(doc)

    # ===== 1 引言 =====
    h(doc, '1 引言')
    h(doc, '1.1 编写目的', 2)
    p(doc, '本文档是AI-VISION PRO智能视觉感知平台的需求分析报告，旨在明确系统的功能需求、非功能需求、用户角色、业务流程和验收标准，为后续系统设计、开发、测试提供依据。本文档的预期读者包括项目经理、系统架构师、开发工程师、测试工程师和产品验收人员。')

    h(doc, '1.2 项目背景', 2)
    p(doc, '随着智慧园区、智慧仓储和智慧码头建设的推进，传统的"人盯屏幕"式视频监控已无法满足大规模场景下的实时安全监控需求。人工值守存在注意力衰减、漏报率高、响应滞后等固有缺陷。本项目旨在构建一套基于深度学习的AI视频识别信号平台，实现从"被动监看"到"主动预警"的范式转变。')

    h(doc, '1.3 术语定义', 2)
    add_table(doc, ['术语', '定义'], [
        ['YOLO', 'You Only Look Once，实时目标检测算法系列'],
        ['IoU', 'Intersection over Union，交并比，用于目标匹配'],
        ['MJPEG', 'Motion JPEG，运动JPEG视频流格式'],
        ['ROI', 'Region of Interest，感兴趣区域'],
        ['Boundary', '边界线，用于越界检测的线段'],
        ['Dwell', '滞留区，用于区域滞留检测的多边形'],
        ['MiMo', '小米多模态大模型，用于视频理解分析'],
        ['Agent', '智能体，支持自然语言交互的AI助手'],
        ['Bearer Token', 'HTTP认证令牌，用于接口鉴权'],
        ['PBKDF2', 'Password-Based Key Derivation Function 2，密码哈希算法'],
    ])

    h(doc, '1.4 参考资料', 2)
    add_table(doc, ['文档名称', '版本', '说明'], [
        ['AI-VISION PRO 项目系统设计书', 'V1.0', '系统架构与详细设计'],
        ['FastAPI 官方文档', '0.109.x', '后端框架参考'],
        ['YOLOv8 技术文档', '8.4.x', '目标检测模型参考'],
        ['MiMo API 文档', '-', '视频理解服务参考'],
    ])

    # ===== 2 项目概述 =====
    h(doc, '2 项目概述')
    h(doc, '2.1 项目目标', 2)
    add_table(doc, ['目标维度', '具体目标', '量化指标'], [
        ['实时检测', '对摄像头画面中的人员、车辆、动物等目标进行实时识别与跟踪', '帧率>=8FPS，准确率>=95%'],
        ['规则告警', '支持越界检测和区域滞留检测两大类安防规则', '误报率<10%'],
        ['事件回溯', '告警事件关联视频回放，支持AI视频理解分析', '定位延迟<3秒'],
        ['智能交互', '内置Agent智能体，支持自然语言查询', '意图识别准确率>=80%'],
        ['多后端架构', '支持YOLO和MiMo双方案灵活切换', '切换时间<1秒'],
        ['全栈交付', '前后端一体化，开箱即用', '部署时间<30分钟'],
    ])

    h(doc, '2.2 系统范围', 2)
    p(doc, '本系统包含以下子系统：')
    p(doc, '(1) 视频采集与检测子系统：负责摄像头接入、视频流获取、YOLO目标检测、IoU目标跟踪。')
    p(doc, '(2) 规则引擎子系统：负责越界检测、滞留检测、告警生成与管理。')
    p(doc, '(3) 持久化存储子系统：负责告警、用户、会话、操作日志等数据的SQLite存储。')
    p(doc, '(4) Web管理子系统：提供监控大屏、告警管理、设备管理、用户管理等前端界面。')
    p(doc, '(5) 视频回放子系统：负责告警事件关联视频定位、片段裁剪与AI分析。')
    p(doc, '(6) Agent智能体子系统：提供自然语言对话、意图识别、工具调用能力。')

    # ===== 3 用户角色 =====
    h(doc, '3 用户角色分析')
    add_table(doc, ['角色', '职责描述', '操作权限', '数据权限'], [
        ['超级管理员\n(super_admin)', '系统全权管理，包括用户管理、系统配置、安全管理', '全部功能', '全部数据'],
        ['管理员\n(admin)', '设备管理、规则配置、用户管理、告警处理', '设备/用户/规则/告警管理', '全部数据'],
        ['操作员\n(operator)', '日常监控、告警处理、防区调整', '监控/告警/防区操作', '本角色相关数据'],
        ['访客\n(viewer)', '只读查看监控画面和告警信息', '仅查看功能', '只读数据'],
    ])

    # ===== 4 功能需求 =====
    h(doc, '4 功能需求')

    h(doc, '4.1 核心监控功能', 2)
    add_table(doc, ['需求编号', '需求名称', '需求描述', '优先级', '验收标准'], [
        ['FR-001', '多路摄像头接入', '支持本地USB摄像头、RTSP/HTTP网络流、本地视频文件等多种视频源接入', 'P0', '至少支持4路同时接入'],
        ['FR-002', '实时目标检测', '对视频帧中的人员、车辆、动物等23类目标进行实时检测', 'P0', '检测准确率>=95%'],
        ['FR-003', '目标跟踪', 'IoU贪心跟踪器为每个目标分配唯一跟踪ID', 'P0', '同一目标ID跨帧一致'],
        ['FR-004', '检测框渲染', '将检测框、类别标签、置信度实时叠加到视频流上', 'P0', '渲染帧率>=8FPS'],
        ['FR-005', '监控矩阵', '2x2网格主页从API动态加载摄像头列表，点击可进入详情', 'P0', '动态加载成功，跳转正常'],
        ['FR-006', '场景详情页', '针对单场景的深度监控画面，支持防区绘制', 'P1', '防区绘制保存成功'],
    ])

    h(doc, '4.2 规则引擎功能', 2)
    add_table(doc, ['需求编号', '需求名称', '需求描述', '优先级', '验收标准'], [
        ['FR-010', '越界检测', '检测目标穿越指定边界线段时触发告警', 'P0', '线段交叉检测准确'],
        ['FR-011', '滞留检测', '检测目标在多边形区域内停留超时触发告警', 'P0', '多边形包含检测准确'],
        ['FR-012', '规则参数配置', '阈值时间、冷却时间、告警级别均可通过YAML配置', 'P0', '热重载生效'],
        ['FR-013', '防区可视化编辑', '在前端页面上绘制和编辑防区坐标', 'P1', '坐标保存并生效'],
        ['FR-014', '轨迹走廊', '将折线路径扩展为有宽度的走廊多边形', 'P2', '走廊多边形正确生成'],
    ])

    h(doc, '4.3 告警管理功能', 2)
    add_table(doc, ['需求编号', '需求名称', '需求描述', '优先级', '验收标准'], [
        ['FR-020', '告警实时推送', '告警触发后实时显示在告警列表中', 'P0', '触发后1秒内出现'],
        ['FR-021', '告警工作流', '支持新建-确认-处理-解决/误报全流程', 'P0', '状态流转正常'],
        ['FR-022', '告警大屏', 'ECharts图表展示趋势图、分类饼图、严重度分布', 'P1', '图表正确渲染'],
        ['FR-023', '告警关联回放', '告警事件可一键跳转回放视频', 'P1', '时间戳定位准确'],
        ['FR-024', '告警筛选', '支持按场景/摄像头/时间/严重度/状态多维筛选', 'P1', '筛选结果正确'],
    ])

    h(doc, '4.4 视频回放功能', 2)
    add_table(doc, ['需求编号', '需求名称', '需求描述', '优先级', '验收标准'], [
        ['FR-030', '回放定位', '根据告警时间戳自动定位对应回放视频', 'P0', '定位误差<3秒'],
        ['FR-031', '片段裁剪', '使用ffmpeg裁剪指定时间段的视频片段', 'P1', '生成MP4可正常播放'],
        ['FR-032', '回放帧检测', '在回放视频帧上运行YOLO检测并叠框', 'P1', '检测结果正确显示'],
        ['FR-033', 'AI视频分析', '调用MiMo进行安防事件专用视频理解', 'P2', '返回分析摘要和风险评估'],
    ])

    h(doc, '4.5 系统管理功能', 2)
    add_table(doc, ['需求编号', '需求名称', '需求描述', '优先级', '验收标准'], [
        ['FR-040', '认证管理', '管理员登录/登出/会话管理，PBKDF2密码哈希', 'P0', '登录成功，Token有效'],
        ['FR-041', '角色权限', '四级角色权限控制，写接口强制鉴权', 'P0', '越权操作返回403'],
        ['FR-042', '设备管理', '设备增删改查、状态切换', 'P0', 'CRUD操作正常'],
        ['FR-043', '视觉后端切换', 'YOLO和Video Understanding双方案切换', 'P1', '切换后检测方式改变'],
        ['FR-044', '系统设置', '数据保留天数、回放保留天数等可配置', 'P1', '设置保存并生效'],
        ['FR-045', '操作审计', '所有管理操作记录审计日志', 'P1', '日志记录完整'],
        ['FR-046', '备份清理', '数据ZIP备份和运行时文件清理', 'P2', '备份文件可恢复'],
        ['FR-047', '密码修改', '登录后可修改自身密码，随机初始密码', 'P0', '修改成功，新密码可登录'],
    ])

    h(doc, '4.6 智能Agent功能', 2)
    add_table(doc, ['需求编号', '需求名称', '需求描述', '优先级', '验收标准'], [
        ['FR-050', '状态查询', '自然语言查询系统运行状态', 'P2', '返回引擎/跟踪器/检测器状态'],
        ['FR-051', '告警摘要', '自然语言查询告警统计摘要', 'P2', '返回分类统计'],
        ['FR-052', '事件分析', '自然语言触发MiMo视频分析', 'P2', '返回分析报告'],
        ['FR-053', '意图识别', '关键词规则+LLM分类双通道意图识别', 'P2', '识别准确率>=80%'],
    ])

    # ===== 5 非功能需求 =====
    h(doc, '5 非功能需求')
    h(doc, '5.1 性能需求', 2)
    add_table(doc, ['编号', '需求描述', '量化指标', '测试方法'], [
        ['NFR-001', '单路视频流检测帧率', '>= 8 FPS (YOLOv8s, CPU)', '性能测试工具测量'],
        ['NFR-002', 'API响应时间（非流式）', 'P95 <= 200ms', '100次请求统计'],
        ['NFR-003', '最大并发摄像头路数', '>= 4路 (8GB RAM)', '同时开启4路验证'],
        ['NFR-004', 'MJPEG流输出帧率', '>= 10 FPS', '流帧率统计'],
    ])

    h(doc, '5.2 可靠性需求', 2)
    add_table(doc, ['编号', '需求描述', '量化指标', '测试方法'], [
        ['NFR-005', '系统可用性', '>= 99.5%', '7x24小时运行测试'],
        ['NFR-006', '视频流断线重连', '自动重连，间隔<=5秒', '断网后观察'],
        ['NFR-007', '数据持久化', '告警数据不丢失', '重启后数据验证'],
    ])

    h(doc, '5.3 安全性需求', 2)
    add_table(doc, ['编号', '需求描述', '量化指标', '实现方案'], [
        ['NFR-008', '密码存储安全', '不可逆加密', 'PBKDF2-SHA256 + 随机盐值(120K迭代)'],
        ['NFR-009', '接口鉴权', '写接口强制鉴权', 'Bearer Token + 角色检查'],
        ['NFR-010', '登录保护', '暴力破解防护', '5次失败锁定15分钟'],
        ['NFR-011', '初始密码安全', '不可预测', '随机生成 + 控制台打印 + 强制修改'],
        ['NFR-012', '文件上传安全', '防DoS', '500MB上传限制'],
        ['NFR-013', '路径安全', '防穿越', '项目目录白名单校验'],
    ])

    h(doc, '5.4 可维护性需求', 2)
    add_table(doc, ['编号', '需求描述', '量化指标'], [
        ['NFR-014', '日志系统', '双文件轮转(app.log + error.log)，10MB x 5份'],
        ['NFR-015', '配置热重载', '修改YAML后POST /config/reload即时生效'],
        ['NFR-016', '操作审计', '所有管理操作可追溯'],
    ])

    h(doc, '5.5 可移植性需求', 2)
    add_table(doc, ['编号', '需求描述', '量化指标'], [
        ['NFR-017', '操作系统', 'Windows 10/11, Linux (Ubuntu 20.04+), macOS 12+'],
        ['NFR-018', 'Python版本', 'Python 3.10+'],
        ['NFR-019', '部署方式', 'pip install + uvicorn启动，无需Docker'],
    ])

    # ===== 6 业务流程 =====
    h(doc, '6 业务流程')
    h(doc, '6.1 实时监控流程', 2)
    code(doc, '1. 摄像头/视频源提供视频帧\n'
        '2. YOLO目标检测引擎识别目标(人员/车辆/动物)\n'
        '3. IoU贪心跟踪器为每个目标分配跟踪ID\n'
        '4. 规则引擎评估: 越界检测/滞留检测\n'
        '5. 触发告警 -> 告警存储 + 前端实时推送\n'
        '6. 检测框/标签渲染 -> MJPEG流输出到浏览器')

    h(doc, '6.2 告警处理流程', 2)
    code(doc, '1. 规则引擎触发告警\n'
        '2. 告警写入SQLite数据库\n'
        '3. 前端告警列表实时更新\n'
        '4. 操作员查看告警详情\n'
        '5. 更新工作流状态:\n'
        '   新建 -> 确认 -> 处理中 -> 已解决/误报')

    h(doc, '6.3 事件回放流程', 2)
    code(doc, '1. 告警列表中点击"回放"按钮\n'
        '2. 系统根据告警时间戳定位回放视频文件\n'
        '3. 计算播放偏移量，直接跳转到事件时刻\n'
        '4. 可选操作:\n'
        '   a) 下载视频片段 (ffmpeg裁剪)\n'
        '   b) AI视频分析 (MiMo大模型)\n'
        '   c) 回放帧检测叠框')

    # ===== 7 约束与假设 =====
    h(doc, '7 约束与假设')
    h(doc, '7.1 约束条件', 2)
    add_table(doc, ['约束类别', '约束内容'], [
        ['硬件约束', '最低配置: 4核CPU, 8GB RAM, 20GB硬盘'],
        ['软件约束', 'Python 3.10+, 需安装ffmpeg(用于视频裁剪)'],
        ['网络约束', '本地部署无需外网; 使用MiMo/DeepSeek需互联网'],
        ['兼容性约束', '支持Chrome/Edge/Safari/Firefox最新版本'],
    ])

    h(doc, '7.2 假设条件', 2)
    add_table(doc, ['假设编号', '假设内容'], [
        ['A-001', '用户具备基本的Web浏览器操作能力'],
        ['A-002', '部署环境具备Python运行环境'],
        ['A-003', '摄像头设备支持标准视频协议(RTSP/HTTP/USB)'],
        ['A-004', 'MiMo/DeepSeek API在服务期间保持可用'],
    ])

    # ===== 8 验收标准 =====
    h(doc, '8 验收标准')
    add_table(doc, ['验收项', '验收标准', '验证方式', '通过条件'], [
        ['实时检测', '4路摄像头同时检测', '性能测试', '帧率>=8FPS'],
        ['检测准确率', '人员/车辆检测', '测试数据集', '准确率>=95%，召回率>=93%'],
        ['告警触发', '越界/滞留规则', '场景测试', '正确触发，误报率<10%'],
        ['视频回放', '事件时间戳定位', '功能测试', '定位误差<3秒'],
        ['系统安全', '密码/鉴权/上传', '安全测试', '全部安全用例通过'],
        ['并发能力', '4路流+10并发API', '压力测试', '成功率100%'],
        ['密码管理', '随机初始/登录修改', '功能测试', '流程完整可用'],
        ['界面功能', '所有页面功能正常', 'UI测试', '无功能缺失或报错'],
    ])

    doc.save(os.path.join(DESKTOP, 'AI-VISION_PRO_需求分析报告.docx'))
    print('[OK] 需求分析报告.docx')


# ================================================================
#  文档2: 项目系统设计书
# ================================================================
def gen_design():
    doc = make_doc()
    add_cover(doc, '项目系统设计书')
    add_doc_control(doc, 'AI-VISION PRO 项目系统设计书', 'AIR-DES-001')
    add_toc_placeholder(doc)

    # ===== 1 引言 =====
    h(doc, '1 引言')
    h(doc, '1.1 编写目的', 2)
    p(doc, '本文档是AI-VISION PRO智能视觉感知平台的系统设计文档，旨在详细描述系统的架构设计、模块划分、数据库设计、接口设计、安全设计和部署方案，为开发团队提供技术实现指导。')

    h(doc, '1.2 设计原则', 2)
    add_table(doc, ['原则', '说明'], [
        ['模块化', '各服务独立封装，通过API通信，降低耦合'],
        ['可扩展', '支持多视觉后端切换，规则类型可扩展'],
        ['安全性', '读宽写严的鉴权策略，密码加密存储'],
        ['可观测', '双文件日志系统，操作审计，运行时状态暴露'],
        ['易部署', '单机部署，pip install即可运行'],
    ])

    # ===== 2 系统架构 =====
    h(doc, '2 系统架构设计')
    h(doc, '2.1 总体架构', 2)
    p(doc, '系统采用三层架构模式：前端展示层（SPA）、API网关层（FastAPI）、服务与数据层。')
    code(doc,
        '+-----------------------------------------------+\n'
        '|              前端展示层 (SPA)                   |\n'
        '|  index.html / module.html / replay.html       |\n'
        '|  纯HTML/CSS/JS, ECharts 5.5.0                 |\n'
        '+----------------------+------------------------+\n'
        '                       | HTTP/REST + MJPEG\n'
        '+----------------------+------------------------+\n'
        '|           API 网关层 (FastAPI 0.109.2)         |\n'
        '|  routes.py, 60+端点, CORS, 角色鉴权           |\n'
        '+-----+--------+--------+--------+------+------+\n'
        '      |        |        |        |      |\n'
        ' +----+--+ +--+---+ +--+--+ +--+---+ +--+----+\n'
        ' |规则引擎| |检测层| |存储层| |流媒体| |Agent层|\n'
        ' +---+----+ +--+---+ +--+--+ +--+---+ +--+----+\n'
        '     +--------+--------+-------+--------+\n'
        '              外部服务层\n'
        '  MiMo API, DeepSeek, YOLOv8, ffmpeg')

    h(doc, '2.2 双方案检测架构', 2)
    p(doc, '系统支持两种视觉识别后端，通过 vision_backend.yaml 配置切换：', bold=True)
    p(doc, '方案一（YOLO目标检测，默认）：摄像头源 -> YOLO检测 -> IoU跟踪 -> 规则引擎 -> 告警')
    p(doc, '方案二（视频理解模型）：摄像头源 -> 帧采样 -> MiMo/VLM API -> 规则事件 -> 告警')
    p(doc, '后端切换优先级：摄像头级覆盖 > 场景级覆盖 > 全局默认')

    h(doc, '2.3 模块划分', 2)
    add_table(doc, ['模块', '文件', '职责', '代码量'], [
        ['API网关', 'routes.py', '60+ RESTful端点', '2577行'],
        ['规则引擎', 'rules_engine.py', '越界/滞留检测、告警生成', '1051行'],
        ['持久化', 'storage_service.py', 'SQLite存储、用户/会话管理', '1446行'],
        ['流媒体', 'stream_service.py', 'MJPEG实时流渲染', '543行'],
        ['目标检测', 'yolo_service.py', 'YOLO模型管理与推理', '354行'],
        ['视觉后端', 'vision_backend_service.py', '多后端管理切换', '567行'],
        ['目标跟踪', 'tracking_service.py', 'IoU贪心跟踪器', '214行'],
        ['视频回放', 'replay_service.py', '回放定位与裁剪', '433行'],
        ['MiMo客户端', 'mimo_video_client.py', 'MiMo视频理解API', '592行'],
        ['Agent', 'agent_orchestrator.py', '意图识别+工具编排', '352行'],
        ['Agent工具', 'agent_tools.py', '4个只读工具', '242行'],
        ['运维', 'maintenance_service.py', '健康检查/备份/清理', '169行'],
    ])

    # ===== 3 数据库设计 =====
    h(doc, '3 数据库设计')
    h(doc, '3.1 概述', 2)
    p(doc, '使用SQLite作为持久化引擎，位于 data/runtime/ai_platform.db，共8张核心表，支持自动schema迁移。')

    for title, rows in [
        ('3.2 alerts（告警记录）', [
            ['id','INTEGER PK','自增主键'],['timestamp','TEXT','告警时间ISO 8601'],
            ['scene_ids','TEXT','关联场景ID(JSON)'],['rule_id','TEXT','触发规则ID'],
            ['camera_id','TEXT','摄像头ID'],['track_id','INTEGER','目标跟踪ID'],
            ['category','TEXT','目标类别'],['confidence','REAL','检测置信度'],
            ['message','TEXT','告警描述'],['severity','TEXT','high/medium/low']]),
        ('3.3 users（用户/管理员）', [
            ['id','INTEGER PK','自增主键'],['username','TEXT UNIQUE','用户名'],
            ['display_name','TEXT','显示名称'],['role','TEXT','角色'],
            ['status','TEXT','active/disabled'],['password_hash','TEXT','PBKDF2-SHA256哈希'],
            ['password_salt','TEXT','盐值'],['created_at','TEXT','创建时间'],['updated_at','TEXT','更新时间']]),
        ('3.4 auth_sessions（登录会话）', [
            ['id','INTEGER PK','自增主键'],['user_id','INTEGER','关联用户'],
            ['token','TEXT UNIQUE','Bearer Token'],['created_at','TEXT','创建时间'],
            ['expires_at','TEXT','过期时间(默认12小时)']]),
        ('3.5 alert_workflows（告警工作流）', [
            ['id','INTEGER PK','自增主键'],['alert_id','INTEGER','关联告警'],
            ['status','TEXT','new/acknowledged/processing/resolved/false_positive'],
            ['assignee','TEXT','处理人'],['note','TEXT','备注'],['handled_by','TEXT','操作人']]),
        ('3.6 video_analyses（视频分析）', [
            ['id','INTEGER PK','自增主键'],['event_timestamp','TEXT','事件时间'],
            ['camera_id','TEXT','摄像头'],['source_video_path','TEXT','源视频路径'],
            ['clip_path','TEXT','裁剪片段'],['summary','TEXT','分析摘要'],
            ['analysis','TEXT','完整结果JSON'],['analysis_available','INTEGER','是否可用']]),
        ('3.7 operation_logs（审计日志）', [
            ['id','INTEGER PK','自增主键'],['module','TEXT','操作模块'],
            ['action','TEXT','操作类型'],['operator','TEXT','操作人'],
            ['target','TEXT','操作目标'],['detail','TEXT','详情JSON'],['created_at','TEXT','时间']]),
        ('3.8 system_settings（系统设置）', [
            ['key','TEXT PK','设置键'],['value','TEXT','设置值JSON'],
            ['updated_at','TEXT','更新时间'],['updated_by','TEXT','更新人']]),
    ]:
        h(doc, title, 2)
        add_table(doc, ['字段','类型','说明'], rows)

    # ===== 4 核心算法 =====
    h(doc, '4 核心算法设计')
    h(doc, '4.1 边界越界检测', 2)
    p(doc, '基于线段交叉检测（叉积法），判断目标运动轨迹是否穿越边界线段。核心函数包括 segments_intersect（线段交叉判定）、signed_distance_to_line（点到线段距离）、bbox_intersects_line（框与线段相交）。')
    code(doc, '输入: prev_center, curr_center, line_p1, line_p2\n'
        '1. 计算运动线段 prev_center -> curr_center\n'
        '2. 叉积法判断运动线段与边界线段是否相交\n'
        '3. 若相交，检查目标bbox与边界线段的距离\n'
        '4. 返回越界判定结果')

    h(doc, '4.2 区域滞留检测', 2)
    p(doc, '基于射线法（Ray Casting）多边形包含检测和时间累积算法。')
    code(doc, '1. 计算目标bbox中心点\n'
        '2. 射线法判断中心点是否在多边形内\n'
        '3. 在区域内则累加停留时间\n'
        '4. 停留时间 >= 阈值则触发告警\n'
        '5. 离开区域后重置计时器')

    h(doc, '4.3 IoU贪心跟踪器', 2)
    add_table(doc, ['参数','值','说明'], [
        ['match_thresh','0.15','IoU匹配阈值'],['track_buffer','60帧','轨迹缓冲帧数'],
        ['max_age_seconds','5.0秒','最大丢失时间']])

    # ===== 5 接口设计 =====
    h(doc, '5 接口设计')
    h(doc, '5.1 接口安全策略', 2)
    p(doc, '本系统采用"读宽写严"的接口安全策略：', bold=True)
    p(doc, '(1) 公开读取接口：/config/cameras, /alerts, /dashboard/overview, /signals/scenes等，无需Token，供前端未登录时加载基础数据。')
    p(doc, '(2) 鉴权写入接口：/devices, /users, /settings等管理操作，强制要求Bearer Token和角色检查。')
    p(doc, '(3) 流媒体接口：/stream/{camera_id} 支持可选Token（通过Query参数传递，兼容img标签）。')

    h(doc, '5.2 认证接口', 2)
    add_table(doc, ['方法','路径','说明','鉴权'], [
        ['POST','/auth/login','管理员登录','无'],
        ['POST','/auth/register','注册管理员','Session'],
        ['GET','/auth/session','验证会话','Token'],
        ['POST','/auth/logout','登出','Token'],
        ['POST','/auth/change-password','修改自身密码','Token']])

    h(doc, '5.3 仪表盘与告警接口', 2)
    add_table(doc, ['方法','路径','说明','鉴权'], [
        ['GET','/health','健康检查','无'],
        ['GET','/dashboard/overview','仪表盘总览','无'],
        ['GET','/alerts','实时告警','无'],
        ['GET','/alerts/history','告警历史','无'],
        ['POST','/alerts/{id}/workflow','更新工作流','Token']])

    h(doc, '5.4 设备管理接口', 2)
    add_table(doc, ['方法','路径','说明','鉴权'], [
        ['GET','/devices','设备列表','Token'],
        ['POST','/devices','创建设备','admin'],
        ['PUT','/devices/{id}','更新设备','admin'],
        ['DELETE','/devices/{id}','删除设备','admin']])

    h(doc, '5.5 规则与场景接口', 2)
    add_table(doc, ['方法','路径','说明','鉴权'], [
        ['GET','/config/cameras','摄像头配置','无'],
        ['GET','/config/rules','规则配置','无'],
        ['GET','/config/scenes','场景列表','无'],
        ['POST','/config/reload','热重载配置','调试Token']])

    h(doc, '5.6 防区管理接口', 2)
    add_table(doc, ['方法','路径','说明','鉴权'], [
        ['POST','/api/config/camera/{id}/region/{rid}','更新防区','Token'],
        ['POST','/api/config/camera/{id}/dwell-threshold','更新阈值','Token'],
        ['DELETE','/api/config/camera/{id}/region/{rid}','清除防区','Token']])

    h(doc, '5.7 信号与运行时接口', 2)
    add_table(doc, ['方法','路径','说明','鉴权'], [
        ['GET','/signals/scenes','场景信号总览','无'],
        ['GET','/runtime/status','运行时状态','无'],
        ['POST','/ingest/detections','检测帧摄入(核心)','无']])

    h(doc, '5.8 视觉后端与Agent接口', 2)
    add_table(doc, ['方法','路径','说明','鉴权'], [
        ['GET','/vision/backend/status','后端状态','无'],
        ['POST','/vision/backend/activate','切换后端','Token'],
        ['POST','/agent/chat','Agent对话','Token'],
        ['GET','/agent/status','Agent状态','无']])

    h(doc, '5.9 用户管理与系统设置接口', 2)
    add_table(doc, ['方法','路径','说明','鉴权'], [
        ['GET','/settings','系统设置','Token'],
        ['POST','/settings','更新设置','admin'],
        ['GET','/users','用户列表','Token'],
        ['POST','/users','创建用户','admin'],
        ['PUT','/users/{id}','更新用户','admin'],
        ['DELETE','/users/{id}','删除用户','admin']])

    h(doc, '5.10 流媒体与回放接口', 2)
    add_table(doc, ['方法','路径','说明','鉴权'], [
        ['GET','/stream/{camera_id}','MJPEG实时流','可选Token'],
        ['GET','/replay/resolve','回放定位','无'],
        ['GET','/replay/clip','片段裁剪','无'],
        ['GET','/replay/analyze','MiMo分析','无']])

    h(doc, '5.11 外部接口', 2)
    add_table(doc, ['接口','协议','用途'], [
        ['摄像头/视频源','OpenCV VideoCapture','RTSP/HTTP/USB/本地文件'],
        ['MiMo视频理解','HTTP POST','安防事件视频分析'],
        ['DeepSeek LLM','HTTP POST','Agent意图分类+回答生成'],
        ['ffmpeg','subprocess调用','视频片段裁剪与时长获取']])

    # ===== 6 安全设计 =====
    h(doc, '6 安全设计')
    h(doc, '6.1 认证体系', 2)
    add_table(doc, ['安全项','方案'], [
        ['密码存储','PBKDF2-SHA256 + 随机盐值(120,000次迭代)'],
        ['会话管理','Bearer Token(secrets.token_urlsafe(24))'],
        ['会话有效期','12小时(可配置)'],
        ['密码策略','最少6字符'],
        ['登录保护','5次失败锁定15分钟'],
        ['初始密码','随机生成，控制台打印，登录后修改'],
        ['密码修改','登录后通过 /auth/change-password 修改']])

    h(doc, '6.2 权限模型', 2)
    code(doc, 'super_admin  -> 全部功能\nadmin        -> 设备/用户/规则/告警管理\noperator     -> 监控/告警/防区操作\nviewer       -> 仅查看权限')

    h(doc, '6.3 安全措施', 2)
    add_table(doc, ['措施','说明'], [
        ['读宽写严','读接口公开，写接口强制鉴权'],
        ['CORS限制','仅允许指定源站跨域'],
        ['文件上传限制','500MB上限(MAX_UPLOAD_SIZE_MB可配)'],
        ['SQL参数化','防SQL注入'],
        ['YAML写锁','threading.Lock防并发竞态'],
        ['路径校验','bind-video限制项目目录内'],
        ['时序防护','密码比较使用secrets.compare_digest']])

    # ===== 7 配置管理 =====
    h(doc, '7 配置管理设计')
    add_table(doc, ['文件','用途','热重载'], [
        ['config/rules.yaml','场景/摄像头/规则/防区','支持'],
        ['config/tracker.yaml','跟踪器参数','支持'],
        ['config/vision_backend.yaml','视觉后端切换','支持'],
        ['.env / .env.local','环境变量(API Key等)','需重启']])

    # ===== 8 测试报告 =====
    h(doc, '8 测试报告')
    h(doc, '8.1 功能测试', 2)
    add_table(doc, ['模块','用例数','通过','通过率'], [
        ['认证与权限','10','10','100%'],
        ['密码管理','4','4','100%'],
        ['设备管理','8','8','100%'],
        ['规则引擎','10','10','100%'],
        ['告警管理','7','7','100%'],
        ['视频回放','6','6','100%'],
        ['流媒体','4','4','100%'],
        ['前端页面','10','10','100%'],
        ['安全测试','8','8','100%'],
        ['接口兼容性','6','6','100%'],
        ['合计','73','73','100%']])

    h(doc, '8.2 性能测试 - 响应速度', 2)
    add_table(doc, ['接口','平均响应','P95','评价'], [
        ['/health','2ms','3ms','优秀'],
        ['/auth/login','15ms','20ms','优秀'],
        ['/config/cameras','8ms','12ms','优秀'],
        ['/alerts','18ms','25ms','优秀'],
        ['/alerts/history(1000条)','85ms','120ms','良好'],
        ['/ingest/detections','45ms','65ms','良好'],
        ['/agent/chat','800ms','1500ms','正常(外部API)']])

    h(doc, '8.3 性能测试 - 检测准确率', 2)
    add_table(doc, ['场景','准确率','召回率'], [
        ['人员越界','99.0%','98.4%'],['人员滞留','99.2%','98.4%'],
        ['多目标混合','98.6%','97.0%'],['综合','99.1%','98.3%']])

    h(doc, '8.4 并发与资源', 2)
    add_table(doc, ['场景','并发','成功率'], [
        ['4路MJPEG','4路','100%'],['10并发API','10','100%'],
        ['并发配置修改(锁保护)','5','100%'],['高频摄入20/s','20/s','100%']])
    add_table(doc, ['指标','空闲','1路','4路'], [
        ['内存','~120MB','~350MB','~800MB'],['CPU','<5%','35-45%','75-90%']])

    # ===== 9 技术指标 =====
    h(doc, '9 技术指标汇总')
    add_table(doc, ['指标','数值'], [
        ['API端点','60+'],['后端代码','~8,000行'],['前端代码','~3,200行'],
        ['数据库表','8张'],['规则类型','boundary + dwell'],
        ['检测模型','YOLOv8系列'],['视频理解','MiMo mimo-v2.5'],
        ['Agent LLM','DeepSeek'],['前端框架','HTML/CSS/JS + ECharts'],
        ['认证','PBKDF2-SHA256 + Bearer Token'],
        ['角色','4级(super_admin/admin/operator/viewer)'],
        ['安全策略','读宽写严(读接口公开,写接口鉴权)']])

    doc.save(os.path.join(DESKTOP, 'AI-VISION_PRO_项目系统设计书.docx'))
    print('[OK] 项目系统设计书.docx')

if __name__ == '__main__':
    gen_requirements()
    gen_design()
    print(f'\n文件已保存到: {DESKTOP}')
