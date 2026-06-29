# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── 设置默认字体 ──────────────────────────────────────────────
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5

# ── 标题 ──────────────────────────────────────────────────────
title = doc.add_heading('', level=0)
run = title.add_run('AI推演模拟：20世纪50年代中国不进行三大改造的后果分析')
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6D)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 副标题
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('——基于历史逻辑与经济学原理的反事实推演（1956—1961）')
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
r.font.name = '楷体'
r.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

doc.add_paragraph()  # 空行

# ── 辅助函数 ──────────────────────────────────────────────────
def add_section_heading(text, level=1):
    h = doc.add_heading('', level=level)
    run = h.add_run(text)
    run.font.name = '黑体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6D)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x2E, 0x5E, 0x8E)
    return h

def add_body(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = '宋体'
        r.font.size = Pt(12)
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        r2 = p.add_run(text)
        r2.font.name = '宋体'
        r2.font.size = Pt(12)
        r2.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    else:
        run = p.add_run(text)
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

# ── 一、历史背景 ──────────────────────────────────────────────
add_section_heading('一、历史背景：三大改造的核心内容')

add_body(
    '三大改造是指1953年至1956年间，中国共产党领导的对农业、手工业和资本主义工商业'
    '的社会主义改造。其核心目标是将生产资料私有制转变为社会主义公有制，为全面建设'
    '社会主义奠定制度基础。'
)

add_section_heading('三大改造的基本进程', level=2)

# 表格
table = doc.add_table(rows=4, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['改造领域', '主要形式', '基本完成时间', '核心内容']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(11)

data = [
    ['农业', '农业合作化（互助组→初级社→高级社）', '1956年底', '土地与主要生产资料入社，统一经营分配'],
    ['手工业', '手工业合作化（供销小组→供销合作社→生产合作社）', '1956年底', '个体手工业者加入合作社，集体生产'],
    ['资本主义工商业', '公私合营（加工订货→统购包销→全行业公私合营）', '1956年初（全行业）', '国家资本主义→赎买政策→国有化'],
]
for r_idx, row_data in enumerate(data):
    for c_idx, val in enumerate(row_data):
        table.rows[r_idx + 1].cells[c_idx].text = val

doc.add_paragraph()

# ── 二、推演前提假设 ──────────────────────────────────────────
add_section_heading('二、推演前提与方法')

add_body(
    '本推演采用反事实历史分析法（Counterfactual Historical Analysis），设定以下核心假设：'
    '从1953年起，中央政府放弃三大改造计划，维持新民主主义经济体制（即公有制经济、'
    '合作社经济、个体经济、私人资本主义经济、国家资本主义经济并存的"五种经济成分"'
    '格局），延续过渡时期总路线中"逐步过渡"的温和路径。'
)

add_section_heading('关键约束条件', level=2)

constraints = [
    ('时间跨度：', '推演至1961年（即不改造条件下五年后），与实际历史形成对照。'),
    ('外部环境：', '假设国际冷战格局不变，西方对华封锁持续，苏联援助条件基本维持。'),
    ('人口与资源：', '假设自然条件、人口规模与实际历史一致。'),
    ('政治体制：', '假设中国共产党执政地位不变，但经济决策路径不同。'),
]
for bold, text in constraints:
    add_bullet(text, bold)

# ── 三、农业领域推演 ──────────────────────────────────────────
add_section_heading('三、农业领域的推演后果')

add_section_heading('3.1 农村经济结构的分化', level=2)
add_body(
    '如果不进行农业合作化运动，农村将维持以家庭为单位的小农经济格局。到1958年左右，'
    '农村社会将出现显著的阶层分化：'
)
add_bullet('一部分勤劳且经营能力强的农户通过土地买卖、雇工经营等方式逐步扩大规模，形成新型富农阶层。')
add_bullet('另一部分农户因天灾、疾病或经营不善而丧失土地，沦为雇农或半雇农。')
add_bullet('农村贫富分化加剧，可能出现土地兼并现象，部分地区出现地主经济的变相复归。')

add_section_heading('3.2 农业生产效率的两面性', level=2)
add_body(
    '短期来看（1953—1956年），由于农民拥有生产自主权，且国家通过统购统销政策维持了'
    '基本的粮食价格稳定，农业产出可能高于实际历史中的合作化初期水平。个体农民在自有'
    '土地上的精耕细作，使粮食单产维持在较高水平。'
)
add_body(
    '但中长期来看（1957—1961年），小农经济面临严重的结构性瓶颈：'
)
add_bullet('水利设施：', '以家庭为单位无法组织大规模农田水利建设，旱涝灾害应对能力极弱。')
add_bullet('技术推广：', '良种、化肥、农药等现代农业要素难以在分散的小农中有效推广。')
add_bullet('规模经济：', '无法实现机械化耕作，劳动生产率长期停滞。')

add_section_heading('3.3 农业产出预测', level=2)

table2 = doc.add_table(rows=6, cols=3)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
h2 = ['年份', '实际历史粮食产量（万吨）', '推演预测产量（万吨）']
for i, h in enumerate(h2):
    table2.rows[0].cells[i].text = h
    for p in table2.rows[0].cells[i].paragraphs:
        for run in p.runs:
            run.bold = True

pred_data = [
    ['1953', '16,683', '16,683（基准年）'],
    ['1955', '18,394', '18,800（略高，个体积极性）'],
    ['1957', '19,505', '18,200（水利瓶颈显现）'],
    ['1959', '16,968', '16,500（自然灾害+无组织抗灾）'],
    ['1961', '14,750', '14,200（持续下滑，缺乏集体保障）'],
]
for r_idx, row in enumerate(pred_data):
    for c_idx, val in enumerate(row):
        table2.rows[r_idx + 1].cells[c_idx].text = val

doc.add_paragraph()

# ── 四、手工业领域推演 ────────────────────────────────────────
add_section_heading('四、手工业领域的推演后果')

add_body(
    '手工业在新中国成立初期是国民经济的重要组成部分，从业人员约一千万人，产值占'
    '全国工业总产值的20%以上。如果不进行手工业合作化，将出现以下变化：'
)

add_section_heading('4.1 市场竞争与分化', level=2)
add_bullet('经营灵活的个体手工业者在市场竞争中脱颖而出，部分发展为小型私营工场甚至现代工厂。')
add_bullet('技术落后、经营不善的手工业者逐步被市场淘汰，沦为产业工人或失业者。')
add_bullet('手工业内部分化加速，形成新的社会阶层结构。')

add_section_heading('4.2 与工业化的关系', level=2)
add_body(
    '国家无法通过手工业合作社进行统一的生产调度和原料分配。一方面，个体手工业者的'
    '灵活性使其能更快适应市场需求变化；另一方面，国家难以将其纳入计划经济体系，'
    '导致工业化所需的配套零部件、日用消费品供应更加市场化但更不稳定。'
)

# ── 五、资本主义工商业领域推演 ────────────────────────────────
add_section_heading('五、资本主义工商业领域的推演后果')

add_body(
    '这是三大改造中影响最为深远的领域。如果不实行全行业公私合营和国有化，中国的'
    '工业经济将呈现截然不同的面貌。'
)

add_section_heading('5.1 私营经济的恢复与发展', level=2)
add_bullet('民族资本家获得更大的经营空间，私营工商业迎来新一轮发展期。')
add_bullet('市场竞争机制在工业领域发挥作用，企业效率可能高于国营企业的平均水平。')
add_bullet('私人资本投入扩大再生产，工业投资来源更加多元化。')

add_section_heading('5.2 国家调控能力的削弱', level=2)
add_body(
    '然而，国家将面临严重的宏观调控困境：'
)
add_bullet('财政收入：', '无法通过国有化获取国有企业的利润上缴，中央财政收入大幅减少。')
add_bullet('计划调配：', '难以对关系国计民生的重要物资实行统一调配，可能出现市场投机和价格波动。')
add_bullet('重工业建设：', '缺乏集中力量办大事的制度基础，"一五"计划中的156项重点工程难以按期推进。')
add_bullet('资本外流：', '部分民族资本家可能将资金转移至香港或海外，造成资本外逃。')

add_section_heading('5.3 工业产出影响', level=2)

table3 = doc.add_table(rows=5, cols=3)
table3.style = 'Light Grid Accent 1'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
h3 = ['指标', '实际历史（1957年）', '推演预测（1957年）']
for i, h in enumerate(h3):
    table3.rows[0].cells[i].text = h
    for p in table3.rows[0].cells[i].paragraphs:
        for run in p.runs:
            run.bold = True

ind_data = [
    ['钢产量', '535万吨', '约350—400万吨'],
    ['工业总产值', '783.9亿元', '约550—600亿元'],
    ['重工业占比', '约45%', '约30—35%'],
    ['国企产值占比', '约54%', '约25—30%'],
]
for r_idx, row in enumerate(ind_data):
    for c_idx, val in enumerate(row):
        table3.rows[r_idx + 1].cells[c_idx].text = val

doc.add_paragraph()

# ── 六、宏观经济与社会发展 ────────────────────────────────────
add_section_heading('六、宏观经济与社会发展的综合推演')

add_section_heading('6.1 经济体制：新民主主义的延续', level=2)
add_body(
    '如果不进行三大改造，中国将长期维持"新民主主义经济体制"，即多种经济成分并存的'
    '混合经济模式。这一体制具有以下特征：'
)
add_bullet('所有制结构：国有经济、合作社经济、个体经济、私人资本主义经济、国家资本主义经济并存。')
add_bullet('资源配置：市场机制与国家计划并行，但以市场为主导。')
add_bullet('收入分配：按劳分配与按资分配并存，社会收入差距可能较大。')

add_section_heading('6.2 GDP与经济增长', level=2)

table4 = doc.add_table(rows=5, cols=3)
table4.style = 'Light Grid Accent 1'
table4.alignment = WD_TABLE_ALIGNMENT.CENTER
h4 = ['年份', '实际GDP增速（%）', '推演GDP增速（%）']
for i, h in enumerate(h4):
    table4.rows[0].cells[i].text = h
    for p in table4.rows[0].cells[i].paragraphs:
        for run in p.runs:
            run.bold = True

gdp_data = [
    ['1953—1955', '平均约9%', '平均约7—8%（工业化投入减少）'],
    ['1956—1957', '平均约6%', '平均约5—6%（市场活跃但投资不足）'],
    ['1958—1960', '大跃进期间剧烈波动', '平均约4—5%（平稳但缓慢）'],
    ['1961', '严重负增长', '约2—3%（无大跃进冲击，但增长乏力）'],
]
for r_idx, row in enumerate(gdp_data):
    for c_idx, val in enumerate(row):
        table4.rows[r_idx + 1].cells[c_idx].text = val

doc.add_paragraph()

add_section_heading('6.3 社会阶层与社会矛盾', level=2)
add_body('不进行三大改造将导致社会阶层结构发生显著变化：')
add_bullet('新的资产阶级：', '民族资本家群体壮大，形成有影响力的资产阶级。')
add_bullet('农村两极分化：', '富裕农民与贫困农民之间的矛盾加剧。')
add_bullet('工人阶级处境：', '私营企业工人缺乏组织保障，劳动条件和工资待遇取决于资本家。')
add_bullet('知识分子：', '在更加多元的社会中，知识分子可能享有更大的思想和学术自由。')
add_bullet('社会矛盾性质变化：', '主要矛盾可能从阶级斗争转向劳资矛盾、贫富分化等更为复杂的社会问题。')

add_section_heading('6.4 对外关系与国防', level=2)
add_body(
    '没有社会主义公有制的制度基础，中国在冷战格局中的定位将更加模糊：'
)
add_bullet('中苏关系：', '苏联可能减少对华援助，因为中国未走上社会主义道路。')
add_bullet('国防建设：', '缺乏集中力量建设国防工业的制度条件，国防现代化进程放缓。')
add_bullet('第三世界影响力：', '中国作为社会主义国家的国际号召力减弱，在亚非拉地区的影响力下降。')

# ── 七、五类后果汇总 ──────────────────────────────────────────
add_section_heading('七、五年后（1958—1961年）核心后果汇总')

table5 = doc.add_table(rows=8, cols=3)
table5.style = 'Light Grid Accent 1'
table5.alignment = WD_TABLE_ALIGNMENT.CENTER
h5 = ['领域', '实际历史后果', '推演后果']
for i, h in enumerate(h5):
    table5.rows[0].cells[i].text = h
    for p in table5.rows[0].cells[i].paragraphs:
        for run in p.runs:
            run.bold = True

summary = [
    ['经济体制', '单一公有制确立', '多种所有制并存，市场调节为主'],
    ['农业', '合作化→人民公社（大跃进破坏）', '小农经济为主，产量缓慢增长但缺乏抗灾能力'],
    ['工业', '国有化完成，重工业优先', '私营经济活跃，轻工业发展快但重工业薄弱'],
    ['社会结构', '消灭剥削阶级，阶层扁平化', '贫富分化加剧，新资产阶级出现'],
    ['财政能力', '国家掌握经济命脉，财政集中', '财政分散，国家调控能力弱'],
    ['对外关系', '一边倒向苏联，社会主义阵营', '外交定位模糊，苏联援助减少'],
    ['政治风险', '大跃进、人民公社化运动', '避免大跃进但面临阶级矛盾激化风险'],
]
for r_idx, row in enumerate(summary):
    for c_idx, val in enumerate(row):
        table5.rows[r_idx + 1].cells[c_idx].text = val

doc.add_paragraph()

# ── 八、辩证评价 ──────────────────────────────────────────────
add_section_heading('八、辩证评价与历史启示')

add_section_heading('8.1 不进行三大改造的"得"', level=2)
add_bullet('避免了后期"左"倾错误的制度基础（人民公社化、大跃进的制度根源部分在于高度集中的公有制）。')
add_bullet('市场经济的活力得以释放，微观经济效率可能更高。')
add_bullet('社会阶层流动性增强，个人经济自由度提高。')
add_bullet('避免了对民族资本家的过度改造，有利于保存经营管理人才。')

add_section_heading('8.2 不进行三大改造的"失"', level=2)
add_bullet('国家无法集中力量推进重工业和国防现代化，在冷战中处于被动。')
add_bullet('农村两极分化可能导致严重的社会不稳定，甚至引发农民运动。')
add_bullet('缺乏公有制基础，社会主义制度的经济根基不牢固。')
add_bullet('无法建立计划经济体制，国家对经济的宏观调控能力大幅削弱。')
add_bullet('面对自然灾害时，分散的小农经济缺乏集体抗灾能力。')

add_section_heading('8.3 历史启示', level=2)
add_body(
    '通过这一反事实推演，我们可以得出以下历史启示：三大改造是中国共产党在特定历史'
    '条件下做出的重大战略选择，它深刻地塑造了此后数十年中国的经济体制和社会结构。'
    '虽然三大改造在后期存在要求过急、工作过粗、改变过快等问题，但从总体上看，它为'
    '中国建立独立完整的工业体系和国民经济体系奠定了制度基础。'
)
add_body(
    '同时，这一推演也提示我们：任何经济体制的选择都有其历史合理性与局限性。改革开放'
    '后中国逐步建立起社会主义市场经济体制，允许多种所有制经济共同发展，正是在总结'
    '历史经验教训基础上做出的正确选择。这段历史告诉我们：经济制度必须与生产力发展'
    '水平相适应，既不能脱离实际盲目求纯，也不能放任自流忽视宏观调控。'
)

# ── 尾注 ──────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('— 本文由AI推演模拟生成，仅供学术参考 —')
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
r.italic = True

# ── 保存 ──────────────────────────────────────────────────────
desktop = os.path.join(os.path.expanduser('~'), 'Desktop')
path = os.path.join(desktop, 'AI推演_不进行三大改造的后果分析.docx')
doc.save(path)
print(f'文档已保存至: {path}')
