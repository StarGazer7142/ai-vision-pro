from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ==================== 页面设置（紧凑） ====================
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)

BLUE   = RGBColor(0x1A, 0x47, 0x8A)
DARK   = RGBColor(0x22, 0x22, 0x22)
GRAY   = RGBColor(0x55, 0x55, 0x55)
LGRAY  = RGBColor(0xBB, 0xBB, 0xBB)

# ==================== 工具函数 ====================
def run(p, text, size=10, bold=False, color=DARK, cn="宋体"):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.font.color.rgb = color
    r.font.name = cn
    r._element.rPr.rFonts.set(qn('w:eastAsia'), cn)
    return r

def set_no_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top    w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:left   w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:right  w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)

def set_col_width(cell, cm):
    cell.width = Cm(cm)

def vertical_align(cell, align="center"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    v = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{align}"/>')
    tcPr.append(v)

def section_bar(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.line_spacing = Pt(14)
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="18" w:space="4" w:color="1A478A"/>'
        f'</w:pBdr>'
    )
    p._element.get_or_add_pPr().append(pBdr)
    run(p, f"  {title}", size=11, bold=True, color=BLUE, cn="微软雅黑")
    return p

def item_line(doc, left, right, left_bold=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.line_spacing = Pt(15)
    run(p, left,  size=10, bold=left_bold, color=GRAY)
    run(p, right, size=10, color=DARK)
    return p

def exp_header(doc, title, time_str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.line_spacing = Pt(14)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(17), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    run(p, title,       size=10.5, bold=True,  color=DARK)
    run(p, f"\t{time_str}", size=9,  color=GRAY)

def body_line(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0.5)
    p.paragraph_format.space_after  = Pt(0.5)
    p.paragraph_format.line_spacing = Pt(14.5)
    if indent:
        p.paragraph_format.left_indent = Cm(0.4)
    run(p, text, size=10, color=DARK)

def thin_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.line_spacing = Pt(4)
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
        f'</w:pBdr>'
    )
    p._element.get_or_add_pPr().append(pBdr)

# ==================== 标题 ====================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_after  = Pt(2)
title_p.paragraph_format.space_before = Pt(0)
run(title_p, "个 人 简 历", size=20, bold=True, color=BLUE, cn="微软雅黑")
thin_line(doc)

# ==================== 基本信息（3列无边框表格） ====================
# 布局：左列标签 | 中列内容 | 右列照片占位（跨4行居中）
info_table = doc.add_table(rows=4, cols=3)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# 清除全部单元格边框
for row in info_table.rows:
    for cell in row.cells:
        set_no_border(cell)

# 列宽
COL_W = [3.2, 7.5, 3.5]
for row in info_table.rows:
    for j, w in enumerate(COL_W):
        row.cells[j].width = Cm(w)

# --- 照片占位：合并右侧 4 行，内容居中 ---
photo_cell = info_table.cell(0, 2).merge(info_table.cell(3, 2))
photo_cell.width = Cm(3.5)
vertical_align(photo_cell, "center")

# 清空默认段落
photo_cell.paragraphs[0].text = ""

# 照片框：用段落 + 边框
ph_p = photo_cell.paragraphs[0]
ph_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
ph_p.paragraph_format.space_before = Pt(0)
ph_p.paragraph_format.space_after  = Pt(0)
# 添加内容使边框可见
run(ph_p, "​", size=1, color=LGRAY)  # 零宽空格撑开

# 给段落加四边框（模拟照片占位框）
pPr = ph_p._element.get_or_add_pPr()
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'<w:top    w:val="single" w:sz="6" w:space="1" w:color="BBBBBB"/>'
    f'<w:left   w:val="single" w:sz="6" w:space="1" w:color="BBBBBB"/>'
    f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="BBBBBB"/>'
    f'<w:right  w:val="single" w:sz="6" w:space="1" w:color="BBBBBB"/>'
    f'</w:pBdr>'
)
pPr.append(pBdr)

# 文字提示（照片框下方）
hint_p = photo_cell.add_paragraph()
hint_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
hint_p.paragraph_format.space_before = Pt(4)
hint_p.paragraph_format.space_after  = Pt(0)
run(hint_p, "上传照片", size=8, color=LGRAY)

# --- 信息行 ---
labels = ["姓    名：", "性    别：", "求职意向：", "联系方式："]
values = ["",           "",           "AI算法工程师 / 数据分析师", "电话：              邮箱："]

for i in range(4):
    c0 = info_table.rows[i].cells[0]
    c1 = info_table.rows[i].cells[1]
    c0.width = Cm(3.2)
    c1.width = Cm(7.5)
    vertical_align(c0, "center")
    vertical_align(c1, "center")

    p0 = c0.paragraphs[0]
    p0.paragraph_format.space_before = Pt(3)
    p0.paragraph_format.space_after  = Pt(3)
    p0.paragraph_format.line_spacing = Pt(16)
    run(p0, labels[i], size=10.5, bold=True, color=DARK)

    p1 = c1.paragraphs[0]
    p1.paragraph_format.space_before = Pt(3)
    p1.paragraph_format.space_after  = Pt(3)
    p1.paragraph_format.line_spacing = Pt(16)
    run(p1, values[i], size=10.5, color=DARK)

thin_line(doc)

# ==================== 教育背景 ====================
section_bar(doc, "教育背景")

p_edu = doc.add_paragraph()
p_edu.paragraph_format.space_before = Pt(3)
p_edu.paragraph_format.space_after  = Pt(1)
p_edu.paragraph_format.line_spacing = Pt(14)
run(p_edu, "人工智能", size=10.5, bold=True, color=DARK)
run(p_edu, "        郑州西亚斯学院（本科）", size=10.5, bold=True, color=DARK)
p_edu.paragraph_format.tab_stops.add_tab_stop(Cm(17), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
run(p_edu, "\t2023.09 — 2027.07", size=9, color=GRAY)

body_line(doc,
    "本科期间，我深入学习了人工智能的各项核心课程，包括机器学习、深度学习、数据结构与算法、"
    "Python编程、计算机视觉、自然语言处理、数据库系统等。我积极参与课程实践项目，"
    "努力掌握各种AI框架和开发工具，如TensorFlow、PyTorch、NumPy、Pandas等。",
    indent=True
)

# ==================== 项目经历 ====================
section_bar(doc, "项目经历")
exp_header(doc, "AI视频识别信号平台（团队项目）", "2025.03 — 至今")
body_line(doc, "项目简介：面向园区围栏、仓库、码头等安防场景的AI视频识别与告警闭环平台")
body_line(doc, "技术栈：Python / FastAPI / YOLOv8 / OpenCV / SQLite / Docker / HTML + JavaScript")
body_line(doc, "负责AI视觉检测模块：基于YOLOv8实现人员目标实时检测，集成IoU贪心跟踪器实现目标轨迹追踪", indent=True)
body_line(doc, "参与前后端开发：FastAPI构建RESTful API，HTML/CSS/JS实现监控大屏与告警管理页面", indent=True)
body_line(doc, "参与规则引擎：实现围栏翻越检测、区域滞留检测等安防规则，支持配置驱动的场景化部署", indent=True)
body_line(doc, "完成Docker容器化部署，系统经44项后端单元测试与16项黑盒测试均全部通过", indent=True)

# ==================== 竞赛经历 ====================
section_bar(doc, "竞赛经历")
exp_header(doc, "全国大学生统计建模大赛 · 河南赛区", "2025.04 — 2025.08")
body_line(doc, "参赛作品：《中国老龄化背景下农村家庭结构变迁与医疗支出趋势分析》")
body_line(doc, "负责数据采集与清洗，运用统计建模方法对农村老龄化医疗支出进行定量分析", indent=True)
body_line(doc, "基于Python完成数据可视化与回归模型构建，产出完整分析报告，荣获本科生组省级二等奖", indent=True)

# ==================== 荣誉与技能 ====================
section_bar(doc, "荣誉与技能")
honors = [
    "2025年  全国大学生统计建模大赛 河南赛区 本科生组二等奖",
    "2024年  全国大学英语四级证书（CET-4）",
    "2024年  普通话水平测试等级证书",
]
for h in honors:
    body_line(doc, h, indent=True)

p_skill = doc.add_paragraph()
p_skill.paragraph_format.space_before = Pt(4)
p_skill.paragraph_format.space_after  = Pt(1)
p_skill.paragraph_format.line_spacing = Pt(14.5)
run(p_skill, "专业技能：", size=10, bold=True, color=DARK)
run(p_skill, "Python（熟练）、YOLO目标检测、FastAPI后端开发、OpenCV视频处理、"
             "NumPy/Pandas数据分析、Docker容器化部署、Nginx反向代理", size=10, color=DARK)

# ==================== 自我评价 ====================
section_bar(doc, "自我评价")
evals = [
    "具备AI视觉检测与全栈开发能力，独立完成基于YOLOv8的视频识别平台开发，涵盖目标检测、规则引擎、前后端页面与容器化部署",
    "学习能力强，能够快速掌握新技术和新工具，在项目中自学YOLO模型部署、FastAPI框架与Docker容器化等技术栈",
    "具有良好的团队协作精神，在竞赛与项目中与团队成员紧密配合，高效完成分工任务",
]
for e in evals:
    body_line(doc, e)

# ==================== 保存 ====================
out = r"C:\Users\fy\Desktop\个人简历.docx"
doc.save(out)
print(f"Done → {out}")
