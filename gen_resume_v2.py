from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

# ==================== 页面设置（紧凑，一页） ====================
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)

BLUE = RGBColor(0x1A, 0x47, 0x8A)
DARK = RGBColor(0x22, 0x22, 0x22)
GRAY = RGBColor(0x55, 0x55, 0x55)

def make_run(p, text, size=10, bold=False, color=DARK, font_cn="宋体", spacing=0):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.font.color.rgb = color
    r.font.name = font_cn
    r._element.rPr.rFonts.set(qn('w:eastAsia'), font_cn)
    if spacing:
        r.font.spacing = Pt(spacing)
    return r

def section_bar(doc, title):
    """蓝色左边框章节标题"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(14)
    # 用底边框模拟左侧蓝线
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="18" w:space="4" w:color="1A478A"/>'
        f'</w:pBdr>'
    )
    p._element.get_or_add_pPr().append(pBdr)
    make_run(p, f"  {title}", size=11, bold=True, color=BLUE, font_cn="微软雅黑")
    return p

def item_line(doc, left, right, left_bold=True):
    """一行两列：左侧标签，右侧内容"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = Pt(15)
    make_run(p, left, size=10, bold=left_bold, color=GRAY)
    make_run(p, right, size=10, bold=False, color=DARK)
    return p

def exp_header(doc, title, time_str):
    """经历标题行：左标题，右时间（灰色）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = Pt(14)
    make_run(p, title, size=10.5, bold=True, color=DARK)
    # 右对齐时间：用tab
    tab_stop = p.paragraph_format.tab_stops
    tab_stop.add_tab_stop(Cm(17), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    make_run(p, f"\t{time_str}", size=9, bold=False, color=GRAY)
    return p

def body_line(doc, text, indent=False):
    """正文行，无符号"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0.5)
    p.paragraph_format.space_after = Pt(0.5)
    p.paragraph_format.line_spacing = Pt(14.5)
    if indent:
        p.paragraph_format.left_indent = Cm(0.4)
    make_run(p, text, size=10, color=DARK)
    return p

def thin_line(doc):
    """浅灰细分隔线"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(4)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p

# ==================== 标题 ====================
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_after = Pt(2)
title_p.paragraph_format.space_before = Pt(0)
make_run(title_p, "个 人 简 历", size=20, bold=True, color=BLUE, font_cn="微软雅黑")

thin_line(doc)

# ==================== 基本信息（表格 + 照片） ====================
# 4行3列：[信息列1] [信息列2] [照片]
t = doc.add_table(rows=4, cols=3)
t.alignment = WD_TABLE_ALIGNMENT.CENTER

# 去掉所有边框
for row in t.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(tcBorders)

# 列宽
for row in t.rows:
    row.cells[0].width = Cm(3)
    row.cells[1].width = Cm(6.5)
    row.cells[2].width = Cm(3)

# 照片占位（合并右侧4行）
photo_cell = t.cell(0, 2)
photo_end = t.cell(3, 2)
photo_merged = photo_cell.merge(photo_end)
photo_merged.width = Cm(3.5)
photo_p = photo_merged.paragraphs[0]
photo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
photo_p.paragraph_format.space_before = Pt(12)
# 照片占位框
box_p = photo_merged.add_paragraph()
box_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
box_p.paragraph_format.space_before = Pt(4)
make_run(box_p, "─────", size=14, color=RGBColor(0xCC, 0xCC, 0xCC), font_cn="宋体")
# 边框模拟照片框
pPr = box_p._element.get_or_add_pPr()
pBdr = parse_xml(
    f'<w:pBdr {nsdecls("w")}>'
    f'<w:top w:val="single" w:sz="4" w:space="2" w:color="BBBBBB"/>'
    f'<w:left w:val="single" w:sz="4" w:space="2" w:color="BBBBBB"/>'
    f'<w:bottom w:val="single" w:sz="4" w:space="2" w:color="BBBBBB"/>'
    f'<w:right w:val="single" w:sz="4" w:space="2" w:color="BBBBBB"/>'
    f'</w:pBdr>'
)
pPr.append(pBdr)
upload_p = photo_merged.add_paragraph()
upload_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
upload_p.paragraph_format.space_before = Pt(2)
make_run(upload_p, "上传照片", size=8, color=RGBColor(0xAA, 0xAA, 0xAA), font_cn="宋体")

# 表格内容
data = [
    ("姓    名：", ""),
    ("性    别：", ""),
    ("求职意向：", "AI算法工程师 / 数据分析师"),
    ("联系方式：", "电话：              邮箱："),
]
for i, (label, value) in enumerate(data):
    cell0 = t.rows[i].cells[0]
    cell1 = t.rows[i].cells[1]
    cell0.width = Cm(2.8)
    cell1.width = Cm(6.7)
    p0 = cell0.paragraphs[0]
    p0.paragraph_format.space_before = Pt(2)
    p0.paragraph_format.space_after = Pt(2)
    make_run(p0, label, size=10.5, bold=True, color=DARK)
    p1 = cell1.paragraphs[0]
    p1.paragraph_format.space_before = Pt(2)
    p1.paragraph_format.space_after = Pt(2)
    make_run(p1, value, size=10.5, color=DARK)

thin_line(doc)

# ==================== 教育背景 ====================
section_bar(doc, "教育背景")

item_line(doc, "郑州西亚斯学院", "2023.09 — 2026.06（预计）")
item_line(doc, "专业：", "人工智能（本科）", left_bold=False)
item_line(doc, "主修课程：", "深度学习、机器学习、Python程序设计、操作系统、计算机组成原理", left_bold=False)

# ==================== 项目经历 ====================
section_bar(doc, "项目经历")

exp_header(doc, "AI视频识别信号平台（团队项目）", "2025.03 — 至今")
body_line(doc, "项目简介：面向园区围栏、仓库、码头等安防场景的AI视频识别与告警闭环平台", indent=False)
body_line(doc, "技术栈：Python / FastAPI / YOLOv8 / OpenCV / SQLite / Docker / HTML + JavaScript", indent=False)
body_line(doc, "负责AI视觉检测模块：基于YOLOv8实现人员目标实时检测，集成IoU贪心跟踪器实现目标轨迹追踪", indent=True)
body_line(doc, "参与前后端开发：FastAPI构建RESTful API，HTML/CSS/JS实现监控大屏与告警管理页面", indent=True)
body_line(doc, "参与规则引擎：实现围栏翻越检测、区域滞留检测等安防规则，支持配置驱动的场景化部署", indent=True)
body_line(doc, "完成Docker容器化部署，系统经44项后端单元测试与16项黑盒测试均全部通过", indent=True)

# ==================== 竞赛经历 ====================
section_bar(doc, "竞赛经历")

exp_header(doc, "全国大学生统计建模大赛 · 河南赛区", "2025.04 — 2025.08")
body_line(doc, "参赛作品：《中国老龄化背景下农村家庭结构变迁与医疗支出趋势分析》", indent=False)
body_line(doc, "负责数据采集与清洗，运用统计建模方法对农村老龄化医疗支出进行定量分析", indent=True)
body_line(doc, "基于Python完成数据可视化与回归模型构建，产出完整分析报告，荣获本科生组省级二等奖", indent=True)

# ==================== 荣誉与技能 ====================
section_bar(doc, "荣誉与技能")

honors = [
    "2025年  全国大学生统计建模大赛 河南赛区 本科生组二等奖",
    "2024年  全国大学英语四级证书（CET-4）",
    "2024年  普通话水平测试等级证书",
]
for h in honors:
    body_line(doc, h, indent=True)

p_skill = doc.add_paragraph()
p_skill.paragraph_format.space_before = Pt(4)
p_skill.paragraph_format.space_after = Pt(1)
make_run(p_skill, "专业技能：", size=10, bold=True, color=DARK)
skills_text = (
    "Python（熟练）、YOLO目标检测、FastAPI后端开发、OpenCV视频处理、"
    "NumPy/Pandas数据分析、Docker容器化部署、Nginx反向代理"
)
make_run(p_skill, skills_text, size=10, color=DARK)

# ==================== 自我评价 ====================
section_bar(doc, "自我评价")

evals = [
    "具备AI视觉检测与全栈开发能力，独立完成基于YOLOv8的视频识别平台开发，涵盖目标检测、规则引擎、前后端页面与容器化部署",
    "学习能力强，能够快速掌握新技术和新工具，在项目中自学YOLO模型部署、FastAPI框架与Docker容器化等技术栈",
    "具有良好的团队协作精神，在竞赛与项目中与团队成员紧密配合，高效完成分工任务",
]
for e in evals:
    body_line(doc, e, indent=False)

# ==================== 保存 ====================
out_path = r"D:\Project\个人简历_v2.docx"
doc.save(out_path)
print(f"Done: {out_path}")
